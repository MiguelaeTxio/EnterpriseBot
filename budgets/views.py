# /home/MiguelAeTxio/PROJECTS/EnterpriseBot/budgets/views.py




# /home/MiguelAeTxio/PROJECTS/EnterpriseBot/budgets/views.py
"""
View definitions for the budgets application.
Implements the sequential budget wizard, HTMX partial endpoints,
result display, status update and ADMIN-only history and detail views.
---
Definiciones de vistas para la aplicacion budgets.
Implementa el asistente de presupuesto secuencial, endpoints parciales HTMX,
visualizacion de resultado, actualizacion de estado y vistas de historial
y detalle exclusivas para ADMIN.
"""

from django import forms as django_forms
from django.contrib import messages
from django.db import transaction
from django.db.models import DecimalField, ExpressionWrapper, F
from django.db.models import Count, Q
from django.http import HttpResponseBadRequest
from django.shortcuts import get_object_or_404, redirect, render
from django.utils.timezone import now
from django.views import View

import csv
import io
import logging
import os

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import weasyprint
from docx import Document
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

from budgets.models import (
    Base,
    Budget,
    BudgetLine,
    Insurer,
    InsurerBase,
    InsurerTariff,
    NightSchedule,
    TariffConcept,
    TariffLine,
    TariffPdfImport,
    TollSegment,
    VehicleType,
)
from budgets.services import calculate_budget, _is_holiday, _resolve_toll_pricing
from ivr_config.models import CompanyUser, PresenceStatus
from panel.mixins import (
    AdminRoleRequiredMixin,
    AssistanceRequiredMixin,
    BudgetAuditAccessMixin,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_company_user(request):
    """
    Return the CompanyUser linked to the authenticated request user.
    ---
    Devuelve el CompanyUser vinculado al usuario autenticado en la peticion.
    """
    return getattr(request.user, "company_user", None)


def _tariff_has_concept(insurer_id, vehicle_type_id, concept_code):
    """
    Return True if the active tariff for the given insurer contains a
    TariffLine for the given concept and vehicle type (or as a generic line).
    ---
    Devuelve True si la tarifa activa de la aseguradora contiene una TariffLine
    para el concepto y tipo de vehiculo dados (o como linea generica).
    """
    from budgets.models import InsurerTariff
    try:
        tariff = InsurerTariff.objects.get(
            insurer_id=insurer_id,
            valid_to__isnull=True,
        )
    except InsurerTariff.DoesNotExist:
        return False
    return tariff.lines.filter(
        concept__code=concept_code,
    ).filter(
        models_q_vehicle(vehicle_type_id)
    ).exists()


def models_q_vehicle(vehicle_type_id):
    """
    Return a Q object matching tariff lines that are either generic
    (vehicle_type=None) or specific to the given vehicle_type_id.
    ---
    Devuelve un objeto Q que coincide con lineas de tarifa que son genericas
    (vehicle_type=None) o especificas para el vehicle_type_id dado.
    """
    from django.db.models import Q
    return Q(vehicle_type__isnull=True) | Q(vehicle_type_id=vehicle_type_id)


def _get_optional_concepts(insurer_id, vehicle_type_id):
    """
    Return a list of concept codes that are available as optional inputs
    for the given insurer and vehicle type combination.
    Only includes concepts that the operator must actively fill in
    (rescue, wait, worker, assistant, custody). Surcharges and departures
    are handled separately and are never shown as optional inputs.
    ---
    Devuelve una lista de codigos de concepto disponibles como entradas
    opcionales para la combinacion de aseguradora y tipo de vehiculo dada.
    Solo incluye conceptos que el operario debe rellenar activamente
    (rescate, espera, mano de obra, ayudante, custodia). Los recargos y
    las salidas se gestionan por separado y nunca se muestran como opcionales.
    """
    from budgets.models import InsurerTariff
    optional_codes = [
        TariffConcept.CODE_RESCUE_HOUR,
        TariffConcept.CODE_WAIT_HOUR,
        TariffConcept.CODE_WORKER_HOUR,
        TariffConcept.CODE_ASSISTANT_HOUR,
        TariffConcept.CODE_CUSTODY_DAY,
    ]
    try:
        tariff = InsurerTariff.objects.get(
            insurer_id=insurer_id,
            valid_to__isnull=True,
        )
    except InsurerTariff.DoesNotExist:
        return []
    available = []
    for code in optional_codes:
        exists = tariff.lines.filter(
            concept__code=code,
        ).filter(
            models_q_vehicle(vehicle_type_id)
        ).exists()
        if exists:
            available.append(code)
    return available


def _has_loaded_surcharge(insurer_id):
    """
    Return True if the active tariff for the insurer includes a
    LOADED_PERCENT surcharge line.
    ---
    Devuelve True si la tarifa activa de la aseguradora incluye una linea
    de recargo LOADED_PERCENT.
    """
    from budgets.models import InsurerTariff
    try:
        tariff = InsurerTariff.objects.get(
            insurer_id=insurer_id,
            valid_to__isnull=True,
        )
    except InsurerTariff.DoesNotExist:
        return False
    return tariff.lines.filter(
        concept__code=TariffConcept.CODE_LOADED_PERCENT
    ).exists()


# ---------------------------------------------------------------------------
# Base context helpers — inject company, company_user, own_presence
# Helpers de contexto base — inyectan company, company_user, own_presence
# ---------------------------------------------------------------------------

def _get_own_presence(company_user):
    """
    Return the current active PresenceStatus for the given CompanyUser.
    ---
    Devuelve el PresenceStatus activo actual para el CompanyUser dado.
    """
    return (
        PresenceStatus.objects.filter(
            company_user=company_user,
            starts_at__lte=now(),
        )
        .filter(
            Q(ends_at__isnull=True) | Q(ends_at__gt=now())
        )
        .order_by("-starts_at")
        .first()
    )


def _build_base_context(request, extra=None):
    """
    Build the base template context required by panel/base.html.
    Injects company, company_user and own_presence so the sidebar
    renders correctly for all roles.
    ---
    Construye el contexto base requerido por panel/base.html.
    Inyecta company, company_user y own_presence para que el sidebar
    se renderice correctamente para todos los roles.
    """
    cu = request.user.company_user
    ctx = {
        "company": cu.company,
        "company_user": cu,
        "own_presence": _get_own_presence(cu),
    }
    if extra:
        ctx.update(extra)
    return ctx


# ---------------------------------------------------------------------------
# Optional concept metadata for form rendering
# ---------------------------------------------------------------------------

OPTIONAL_CONCEPT_META = {
    TariffConcept.CODE_RESCUE_HOUR: {
        "label": "Horas de rescate",
        "field": "rescue_hours",
        "input_type": "number",
        "step": "0.5",
    },
    TariffConcept.CODE_WAIT_HOUR: {
        "label": "Horas de espera",
        "field": "wait_hours",
        "input_type": "number",
        "step": "0.5",
    },
    TariffConcept.CODE_WORKER_HOUR: {
        "label": "Horas de mano de obra",
        "field": "worker_hours",
        "input_type": "number",
        "step": "0.5",
    },
    TariffConcept.CODE_ASSISTANT_HOUR: {
        "label": "Horas de ayudante",
        "field": "assistant_hours",
        "input_type": "number",
        "step": "0.5",
    },
    TariffConcept.CODE_CUSTODY_DAY: {
        "label": "Dias de custodia",
        "field": "custody_days",
        "input_type": "number",
        "step": "1",
    },
}


# ---------------------------------------------------------------------------
# Budget wizard — sequential form (ASSISTANCE + ADMIN)
# ---------------------------------------------------------------------------

class BudgetWizardView(AssistanceRequiredMixin, View):
    """
    Sequential budget wizard. Renders the step-by-step form for the operator.
    On GET: renders the wizard with the list of active insurers.
    On POST: validates all inputs, calculates the budget via the engine,
    persists Budget + BudgetLine records in a single transaction and
    redirects to the result view.
    ---
    Asistente de presupuesto secuencial. Renderiza el formulario paso a paso
    para el operario. En GET: renderiza el asistente con la lista de
    aseguradoras activas. En POST: valida todas las entradas, calcula el
    presupuesto via el motor, persiste Budget + BudgetLine en una sola
    transaccion y redirige a la vista de resultado.
    """

    template_name = "budgets/wizard.html"

    def get(self, request):
        """
        Render the wizard with active insurers for the operator's company.
        Calculates insurer_label dynamically based on the mix of
        is_insurance_company flags in the active insurer queryset.
        ---
        Renderiza el asistente con las aseguradoras activas de la empresa
        del operario. Calcula insurer_label de forma dinámica según
        la mezcla de flags is_insurance_company en el queryset activo.
        """
        company_user = _get_company_user(request)
        insurers = Insurer.objects.filter(
            company=company_user.company,
            is_active=True,
        ).order_by("name")
        # Resolve the label for the insurer dropdown based on the mix of
        # is_insurance_company flags in the active queryset.
        # Resuelve el label del desplegable según la mezcla de flags
        # is_insurance_company en el queryset activo.
        insurer_flags = list(
            insurers.values_list("is_insurance_company", flat=True)
        )
        has_insurance = any(insurer_flags)
        has_particular = not all(insurer_flags)
        if has_insurance and has_particular:
            insurer_label = "Aseguradora / Cliente"
        elif has_particular:
            insurer_label = "Cliente particular"
        else:
            insurer_label = "Aseguradora"
        # Build a pk->always_apply_iva map for the JS layer.
        # The template uses it to auto-check and lock the IVA checkbox
        # when the operator selects an insurer with always_apply_iva=True.
        # Construye un mapa pk->always_apply_iva para la capa JS.
        # El template lo usa para marcar y bloquear el checkbox de IVA
        # cuando el operario selecciona una aseguradora con always_apply_iva=True.
        # Build pk->always_apply_iva map for the JS IVA enforcement.
        # Construye mapa pk->always_apply_iva para el bloqueo de IVA en JS.
        import json
        always_iva_map = json.dumps({
            str(ins.pk): ins.always_apply_iva
            for ins in insurers
        })
        ctx = _build_base_context(request, {
            "insurers": insurers,
            "insurer_label": insurer_label,
            "always_iva_map": always_iva_map,
            "active_nav": "budgets_wizard",
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
        })
        return render(request, self.template_name, ctx)

    def post(self, request):
        """
        Process the wizard form submission, run the calculation engine and
        persist the budget. Redirects to BudgetResultView on success.
        ---
        Procesa el envio del formulario del asistente, ejecuta el motor de
        calculo y persiste el presupuesto. Redirige a BudgetResultView
        en caso de exito.
        """
        company_user = _get_company_user(request)
        data = request.POST

        # --- Validate required fields ---
        # --- Validar campos obligatorios ---
        import datetime as _dt
        from decimal import Decimal as _Dec
        try:
            insurer_id      = int(data["insurer_id"])
            vehicle_type_id = int(data["vehicle_type_id"])
            service_date_str = data["service_date"]
            service_date = _dt.date.fromisoformat(service_date_str)
            is_overnight = data.get("is_overnight") == "1"
            has_unlock   = data.get("has_unlock") == "1"
            is_loaded       = data.get("is_loaded") == "1"
            base_id = data.get("base_id") or None
            if base_id:
                base_id = int(base_id)
        except (KeyError, ValueError):
            messages.error(
                request,
                "Faltan datos obligatorios. Revisa el formulario.",
            )
            return redirect("budgets:wizard")

        # --- Read route fields and resolve km ---
        # --- Leer campos de ruta y resolver km ---
        road_name              = data.get("road_name", "").strip()
        dest_location          = data.get("dest_location", "").strip()
        pk_km_raw              = data.get("pk_km", "").strip()
        service_date_raw       = data.get("service_date", "").strip()
        # Resolve is_night automatically from the insurer's NightSchedule.
        # The operator no longer has a manual checkbox — the system decides.
        # Resolution: insurer.night_schedule if set, else company default.
        # Resolver is_night automáticamente desde el NightSchedule de la
        # aseguradora. El operario ya no tiene checkbox manual — el sistema
        # decide. Resolución: insurer.night_schedule si tiene, si no el
        # horario por defecto de la empresa.
        is_night = False
        route_distance_km_raw  = data.get("route_distance_km", "").strip()
        route_toll_cost_raw    = data.get("route_toll_cost", "").strip()
        route_calculation_mode = data.get(
            "route_calculation_mode", "MANUAL"
        ).strip()

        # --- Read multi-stop planner fields (Modo B / route_calculation_mode=API) ---
        # --- Leer campos del planificador multi-parada (Modo B / mode=API) ---
        # waypoints_json: JSON string serialised by wizard_map.js confirmRoute().
        # is_overnight_route: "true"/"false" string set by the modal, overrides
        #   the manual is_overnight checkbox when mode=API.
        # km_phase2_route: phase-2 distance set by the modal for overnight routes.
        # route_toll_budget_cost_raw: budget toll cost returned by BudgetWaypointView.
        import json as _json
        waypoints_json_raw = data.get("waypoints_json", "").strip()
        waypoints_json = None
        if route_calculation_mode == "API" and waypoints_json_raw:
            try:
                waypoints_json = _json.loads(waypoints_json_raw)
            except _json.JSONDecodeError:
                waypoints_json = None
            # Override is_overnight from the modal hidden input.
            _is_overnight_api = data.get("is_overnight_route", "false").strip().lower()
            is_overnight = _is_overnight_api == "true"

        # encoded_polyline: Google Encoded Polyline of the confirmed route.
        # Persisted so calculate_budget() can cross-reference TollSegment.
        # Overnight services: two polylines joined by \x01.
        # ---
        # Polyline codificada Google de la ruta confirmada.
        # Se persiste para que calculate_budget() cruce con TollSegment.
        # Servicios de pernocta: dos polylines unidas por \x01.
        encoded_polyline_raw = data.get("encoded_polyline", "").strip()
        encoded_polyline = encoded_polyline_raw if encoded_polyline_raw else None

        route_toll_budget_cost_raw = data.get(
            "route_toll_budget_cost", ""
        ).strip()
        route_toll_budget_cost = (
            _Dec(route_toll_budget_cost_raw.replace(",", "."))
            if route_toll_budget_cost_raw else None
        )

        # --- Nocturno/Festivo forzado — compartido por ambos modos (S005) ---
        # --- Forced nocturno/festivo — shared by both calculation modes (S005) ---
        # manual_is_night_holiday: operator checkbox, shown once in the
        #   wizard regardless of calculation mode (API or MANUAL). FORCES
        #   is_night_or_holiday to True when checked — the automatic
        #   calendar-based calculation still always runs underneath and is
        #   used whenever this checkbox is left unchecked (see
        #   calculate_budget()). Covers services agreed with the insurer as
        #   nocturno/festivo (e.g. an out-of-zone job) even when using route
        #   planning, not just Manual mode.
        # ---
        # manual_is_night_holiday: checkbox del operario, mostrado una sola
        #   vez en el wizard independientemente del modo de cálculo (API o
        #   MANUAL). FUERZA is_night_or_holiday a True cuando está marcado —
        #   el cálculo automático por calendario sigue ejecutándose siempre
        #   por debajo y es el que se usa cuando esta casilla queda sin
        #   marcar (ver calculate_budget()). Cubre servicios acordados con la
        #   aseguradora como nocturno/festivo (p. ej. un servicio fuera de
        #   zona) también cuando se usa planificación de ruta, no solo en
        #   modo Manual.
        is_night_or_holiday_manual_override = (
            data.get("manual_is_night_holiday") == "1"
        )

        # --- Manual calculation mode fields (route_calculation_mode=MANUAL) ---
        # --- Campos del modo de cálculo Manual (route_calculation_mode=MANUAL) ---
        # manual_toll_total: operator-entered "Otros peajes" catch-all amount,
        #   fed into route_toll_budget_cost. Additive with the itemized toll
        #   segment table below — not a replacement for it.
        # toll_pases_<segment_id>: itemized pass count per TollSegment row
        #   shown in the wizard's toll table (troncal/salida, typically
        #   AP-7/AP-46). Collected into manual_toll_segments as
        #   {segment_id: passes}, ignoring blank/zero entries.
        # ---
        # manual_toll_total: importe catch-all "Otros peajes" introducido
        #   por el operario, se vuelca en route_toll_budget_cost. Es aditivo
        #   con la tabla itemizada de tramos de peaje de abajo — no la
        #   sustituye.
        # toll_pases_<segment_id>: número de pases por cada fila TollSegment
        #   de la tabla de peajes del wizard (troncal/salida, típicamente
        #   AP-7/AP-46). Se recogen en manual_toll_segments como
        #   {id_tramo: pases}, ignorando entradas vacías o a cero.
        manual_toll_total_raw = data.get("manual_toll_total", "").strip()
        manual_toll_total = (
            _Dec(manual_toll_total_raw.replace(",", "."))
            if manual_toll_total_raw else None
        )
        manual_toll_segments = None
        if route_calculation_mode == "MANUAL":
            _toll_segments_dict = {}
            for _key in data:
                if not _key.startswith("toll_pases_"):
                    continue
                _sid = _key[len("toll_pases_"):].strip()
                if not _sid.isdigit():
                    continue
                try:
                    _passes = int(str(data.get(_key, "")).strip() or 0)
                except ValueError:
                    continue
                if _passes > 0:
                    _toll_segments_dict[_sid] = _passes
            manual_toll_segments = _toll_segments_dict or None

        service_time_raw = data.get("service_time", "").strip()
        # Normalise decimal separator: JS may write '.' or ','
        # depending on the browser locale.
        # Normalizar separador decimal: el JS puede escribir '.' o ','
        # segun el locale del navegador.
        route_distance_km = (
            _Dec(route_distance_km_raw.replace(",", "."))
            if route_distance_km_raw else None
        )
        route_toll_cost = (
            _Dec(route_toll_cost_raw.replace(",", "."))
            if route_toll_cost_raw else None
        )
        service_time_obj = (
            _dt.time.fromisoformat(service_time_raw)
            if service_time_raw else None
        )

        # --- Resolve is_night from NightSchedule ---
        # --- Resolver is_night desde NightSchedule ---
        # service_time is now required, so service_time_obj should always
        # be set at this point. Defensive None check kept for safety.
        # service_time es ahora obligatorio, por lo que service_time_obj
        # debería estar siempre disponible. Se mantiene el None check defensivo.
        if service_time_obj is not None:
            from budgets.services import _resolve_night_schedule
            _insurer_for_night = get_object_or_404(Insurer, pk=insurer_id)
            _ns_obj = _resolve_night_schedule(_insurer_for_night)
            if _ns_obj is not None:
                _ns = _ns_obj.night_start
                _ne = _ns_obj.night_end
                if _ns <= _ne:
                    is_night = _ns <= service_time_obj <= _ne
                else:
                    # Interval crosses midnight.
                    # El intervalo cruza la medianoche.
                    is_night = (
                        service_time_obj >= _ns
                        or service_time_obj <= _ne
                    )

        # When Modo B is active, use route_distance_km as km_phase1.
        # Route API returns a Decimal with up to 2 decimal places.
        # km_phase1 is an IntegerField — round using ROUND_HALF_UP
        # (standard arithmetic rounding: 33.5 → 34, 33.4 → 33).
        # Python's built-in round() uses banker's rounding — avoid it.
        # En Modo B activo, usar route_distance_km como km_phase1.
        # La Routes API devuelve un Decimal con hasta 2 decimales.
        # km_phase1 es un IntegerField — redondear con ROUND_HALF_UP
        # (redondeo aritmético estándar: 33.5 → 34, 33.4 → 33).
        # El round() de Python usa banker's rounding — no usar.
        if route_calculation_mode == "API" and route_distance_km:
            from decimal import ROUND_HALF_UP
            km_phase1 = int(
                _Dec(str(route_distance_km)).quantize(
                    _Dec("1"), rounding=ROUND_HALF_UP
                )
            )
        else:
            try:
                # Normalise decimal separator before conversion.
                # Normalizar separador decimal antes de la conversion.
                _km1_raw = str(data.get("km_phase1", "")).strip()
                km_phase1 = float(_km1_raw.replace(",", "."))
            except (KeyError, ValueError):
                messages.error(
                    request,
                    "Introduce los kilómetros de ida y vuelta (fase 1).",
                )
                return redirect("budgets:wizard")

        km_phase2 = None
        if is_overnight:
            if route_calculation_mode == "API":
                # km_phase2 is set by the modal into the hidden input km_phase2_route.
                # km_phase2 lo vuelca el modal en el hidden input km_phase2_route.
                try:
                    _km2_raw = str(data.get("km_phase2_route", "")).strip()
                    km_phase2 = float(_km2_raw.replace(",", ".")) if _km2_raw else None
                except (KeyError, ValueError):
                    km_phase2 = None
            else:
                try:
                    # Normalise decimal separator before conversion.
                    # Normalizar separador decimal antes de la conversion.
                    _km2_raw = str(data.get("km_phase2", "")).strip()
                    km_phase2 = float(_km2_raw.replace(",", "."))
                except (KeyError, ValueError):
                    messages.error(
                        request,
                        "Introduce los kilómetros de la fase 2 "
                        "para servicios de pernocta.",
                    )
                    return redirect("budgets:wizard")

        # --- Resolve FK objects ---
        # --- Resolver objetos FK ---
        insurer = get_object_or_404(
            Insurer,
            pk=insurer_id,
            company=company_user.company,
            is_active=True,
        )
        vehicle_type = get_object_or_404(
            VehicleType,
            pk=vehicle_type_id,
            insurer=insurer,
            is_active=True,
        )

        # --- Build Budget instance (unsaved) ---
        # --- Construir instancia Budget (sin guardar) ---
        budget = Budget(
            company=company_user.company,
            insurer=insurer,
            operator=company_user,
            service_date=service_date,
            vehicle_type=vehicle_type,
            is_overnight=is_overnight,
            km_phase1=km_phase1,
            km_phase2=km_phase2,
            has_unlock=has_unlock,
            is_night=is_night,
            is_night_or_holiday_manual_override=(
                is_night_or_holiday_manual_override
            ),
            manual_toll_segments=manual_toll_segments,
            is_loaded=is_loaded,
            base=(
                InsurerBase.objects.filter(
                    base_id=base_id,
                    insurer=insurer,
                    is_active=True,
                    base__is_active=True,
                ).select_related("base").first().base
                if base_id else None
            ),
            wait_hours=_parse_optional(data.get("wait_hours")),
            rescue_hours=_parse_optional(data.get("rescue_hours")),
            assistant_hours=_parse_optional(data.get("assistant_hours")),
            worker_hours=_parse_optional(data.get("worker_hours")),
            custody_days=_parse_optional(data.get("custody_days")),
            extra_notes=data.get("extra_notes", "").strip(),
            apply_iva=data.get("apply_iva") == "1",
            road_name=road_name,
            dest_location=dest_location,
            pk_km=_Dec(pk_km_raw.replace(",", ".")) if pk_km_raw else None,
            route_distance_km=route_distance_km,
            route_toll_cost=route_toll_cost,
            route_toll_budget_cost=(
                route_toll_budget_cost
                if route_calculation_mode == "API"
                else manual_toll_total
            ),
            route_calculation_mode=route_calculation_mode,
            service_time=service_time_obj,
            waypoints_json=waypoints_json,
            encoded_polyline=encoded_polyline,
            # Placeholders — set by the engine before save.
            # Valores provisionales — el motor los establece antes de guardar.
            total_amount=0,
            total_amount_with_iva=None,
            tariff_id=None,
        )

        # --- Run calculation engine and persist atomically ---
        # --- Ejecutar motor de calculo y persistir atomicamente ---
        try:
            with transaction.atomic():
                budget_lines = calculate_budget(budget)
                budget.save()
                for line in budget_lines:
                    line.budget = budget
                BudgetLine = budget.lines.model
                BudgetLine.objects.bulk_create(budget_lines)
        except ValueError as exc:
            messages.error(request, str(exc))
            return redirect("budgets:wizard")

        return redirect("budgets:result", pk=budget.pk)


def _get_concept_choices(company):
    """
    Return a list of (code, label) tuples for TariffConcept available for the
    given company: system concepts (company=None) plus concepts owned by this
    company. Used to populate the concept dropdown when adding or editing lines.
    ---
    Devuelve una lista de tuplas (code, label) de TariffConcept disponibles
    para la empresa dada: conceptos de sistema (company=None) más los propios.
    Se usa para poblar el desplegable de concepto al añadir o editar líneas.
    """
    from django.db.models import Q
    qs = TariffConcept.objects.filter(
        Q(company__isnull=True) | Q(company=company)
    ).order_by("sort_order", "label")
    return [(tc.code, tc.label) for tc in qs]


def _parse_decimal(raw, default=None):
    """
    Normalise a raw decimal string from POST, replacing comma with period
    before conversion. Returns the normalised string for DB assignment,
    or ``default`` if the input is empty or None.
    ---
    Normaliza una cadena decimal cruda del POST, sustituyendo la coma por
    punto antes de la conversion. Devuelve la cadena normalizada para BD,
    o ``default`` si la entrada esta vacia o es None.
    """
    if not raw:
        return default
    return str(raw).strip().replace(",", ".")


def _parse_optional(raw):
    """
    Parse an optional numeric form field. Returns None if empty or invalid.
    ---
    Parsea un campo de formulario numerico opcional.
    Devuelve None si esta vacio o es invalido.
    """
    if not raw:
        return None
    try:
        value = float(raw)
        return value if value > 0 else None
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# HTMX endpoint — route calculation via Routes API
# Endpoint HTMX — cálculo de ruta via Routes API
# ---------------------------------------------------------------------------

class BudgetRouteCalcView(AssistanceRequiredMixin, View):
    """
    HTMX POST endpoint. Receives base_id, road_name, pk_km, service_date
    and service_time. Calls calculate_route() and returns the route result
    fragment. Returns error fragment on RouteCalculationError.
    ---
    Endpoint POST HTMX. Recibe base_id, road_name, pk_km, service_date
    y service_time. Llama a calculate_route() y devuelve el fragmento de
    resultado de ruta. Devuelve fragmento de error en RouteCalculationError.
    """

    template_name = "budgets/_route_calc_fragment.html"

    def post(self, request):
        """
        Validate inputs, call calculate_route() and return the result fragment.
        ---
        Valida los campos, llama a calculate_route() y devuelve el fragmento
        de resultado.
        """
        import datetime as _dt
        from decimal import Decimal

        company_user = _get_company_user(request)
        data = request.POST

        base_id          = data.get("base_id", "").strip()
        road_name        = data.get("road_name", "").strip()
        dest_location    = data.get("dest_location", "").strip()
        pk_km_raw        = data.get("pk_km", "").strip()
        service_date_str = data.get("service_date", "").strip()
        service_time_str = data.get("service_time", "").strip()

        if not all([base_id, road_name, pk_km_raw, service_date_str, service_time_str]):
            return render(request, self.template_name, {
                "error": "Rellena carretera, punto kilómetrico, fecha y hora del servicio.",
            })

        try:
            base             = Base.objects.get(pk=int(base_id), company=company_user.company)
            pk_km            = Decimal(pk_km_raw.replace(",", "."))
            service_date     = _dt.date.fromisoformat(service_date_str)
            service_time_obj = _dt.time.fromisoformat(service_time_str)
            service_datetime = _dt.datetime.combine(service_date, service_time_obj)
        except Exception as exc:
            logger.warning(
                "# [budgets] Datos inválidos en formulario de ruta. "
                "base_id=%r pk_km_raw=%r service_date_str=%r "
                "service_time_str=%r: %s",
                base_id, pk_km_raw, service_date_str, service_time_str, exc,
                exc_info=True,
            )
            return render(request, self.template_name, {
                "error": (
                    "Datos inválidos. Revisa la base, el punto kilométrico, "
                    "la fecha y la hora del servicio."
                ),
            })

        # Any date (past, present or future) is now accepted — the service
        # layer omits departureTime for non-future dates since Google's
        # Routes API rejects a past departureTime for DRIVE mode.
        # Cualquier fecha (pasada, presente o futura) es aceptada ahora —
        # la capa de servicios omite departureTime para fechas no futuras,
        # ya que la Routes API de Google rechaza un departureTime pasado
        # en modo DRIVE.

        from budgets.services import calculate_route, RouteCalculationError
        try:
            result = calculate_route(
                base,
                road_name,
                pk_km,
                service_datetime,
                dest_location=dest_location,
            )
        except RouteCalculationError as exc:
            return render(request, self.template_name, {
                "error": str(exc),
            })

        # Extract primary route from the new dual contract.
        # Extraer la ruta primaria del nuevo contrato dual.
        primary = result["route_with_tolls"]
        return render(request, self.template_name, {
            "distance_km":  primary["distance_km"],
            "has_tolls":    primary["has_tolls"],
            "road_name":    road_name,
            "pk_km":        pk_km,
            "service_time": service_time_str,
            "base_name":    base.name,
            "base_lat":     base.latitude,
            "base_lng":     base.longitude,
        })


# ---------------------------------------------------------------------------
# HTMX endpoint — dual route visualisation (with tolls vs without tolls)
# Endpoint HTMX — visualizacion dual de rutas (con peajes vs sin peajes)
# ---------------------------------------------------------------------------

class BudgetRouteDualView(AssistanceRequiredMixin, View):
    """
    HTMX GET endpoint. Receives base_id, road_name, pk_km, dest_location
    and service_datetime as query parameters. Calls calculate_route() and
    returns the dual route fragment _route_dual_fragment.html.

    Context flags computed in the view (Dumb Template):
      - route_with_tolls: primary route dict.
      - route_without_tolls: secondary route dict or None.
      - show_dual: True only when route_without_tolls is not None.
    ---
    Endpoint GET HTMX. Recibe base_id, road_name, pk_km, dest_location
    y service_datetime como parametros de consulta. Llama a calculate_route()
    y devuelve el fragmento de ruta dual _route_dual_fragment.html.

    Flags de contexto calculados en la vista (Dumb Template):
      - route_with_tolls: dict de la ruta primaria.
      - route_without_tolls: dict de la ruta secundaria o None.
      - show_dual: True solo cuando route_without_tolls no es None.
    """

    template_name = "budgets/_route_dual_fragment.html"

    def get(self, request):
        """
        Validate inputs, call calculate_route() and return the dual
        route fragment with all presentation flags resolved.
        ---
        Valida los parametros, llama a calculate_route() y devuelve el
        fragmento de ruta dual con todos los flags de presentacion resueltos.
        """
        import datetime as _dt
        import json
        from decimal import Decimal

        company_user = _get_company_user(request)
        params = request.GET

        base_id           = params.get("base_id", "").strip()
        road_name         = params.get("road_name", "").strip()
        dest_location     = params.get("dest_location", "").strip()
        pk_km_raw         = params.get("pk_km", "").strip()
        service_date_str  = params.get("service_date", "").strip()
        service_time_str  = params.get("service_time", "").strip()

        if not all([
            base_id,
            road_name,
            pk_km_raw,
            service_date_str,
            service_time_str,
        ]):
            return render(request, self.template_name, {
                "error": (
                    "Rellena carretera, punto kilométrico, "
                    "fecha y hora del servicio."
                ),
                "show_dual": False,
                "route_with_tolls": None,
                "route_without_tolls": None,
            })

        try:
            base = Base.objects.get(
                pk=int(base_id),
                company=company_user.company,
            )
            pk_km = Decimal(pk_km_raw.replace(",", "."))
            service_date = _dt.date.fromisoformat(service_date_str)
            service_time_obj = _dt.time.fromisoformat(service_time_str)
            service_datetime = _dt.datetime.combine(
                service_date, service_time_obj,
            )
        except Exception as exc:
            logger.warning(
                "# [budgets] Datos inválidos en formulario de ruta (dual). "
                "base_id=%r pk_km_raw=%r service_date_str=%r "
                "service_time_str=%r: %s",
                base_id, pk_km_raw, service_date_str, service_time_str, exc,
                exc_info=True,
            )
            return render(request, self.template_name, {
                "error": (
                    "Datos inválidos. Revisa la base, el punto kilométrico, "
                    "la fecha y la hora del servicio."
                ),
                "show_dual": False,
                "route_with_tolls": None,
                "route_without_tolls": None,
            })

        # Any date (past, present or future) is now accepted — the
        # service layer omits departureTime for non-future dates since
        # Google's Routes API rejects a past departureTime for DRIVE mode.
        # Cualquier fecha es aceptada ahora — la capa de servicios omite
        # departureTime para fechas no futuras, ya que la Routes API de
        # Google rechaza un departureTime pasado en modo DRIVE.

        from budgets.services import calculate_route, RouteCalculationError
        try:
            result = calculate_route(
                base,
                road_name,
                pk_km,
                service_datetime,
                dest_location=dest_location,
            )
        except RouteCalculationError as exc:
            return render(request, self.template_name, {
                "error": str(exc),
                "show_dual": False,
                "route_with_tolls": None,
                "route_without_tolls": None,
            })

        route_with_tolls = result["route_with_tolls"]
        route_without_tolls = result["route_without_tolls"]
        show_dual = route_without_tolls is not None

        # Compute calendar_warning: True when the stored labor_calendar year
        # does not match the requested service_date year. Passed to the
        # template so the operator sees a visible alert in the UI.
        # Mobile holidays (Easter) may be inaccurate when calendar is stale.
        # Calcular calendar_warning: True cuando el anio del labor_calendar
        # almacenado no coincide con el anio de service_date. Se pasa al
        # template para que el operario vea un aviso visible en la UI.
        # Los festivos moviles (Semana Santa) pueden ser inexactos si caducado.
        calendar_warning = False
        calendar_warning_year = None
        _cal_json = (base.labor_calendar or "").strip()
        if _cal_json:
            try:
                import json as _json_cal
                _holidays_cal = _json_cal.loads(_cal_json)
                if _holidays_cal:
                    _stored_year = int(_holidays_cal[0][:4])
                    if _stored_year != service_date.year:
                        calendar_warning = True
                        calendar_warning_year = _stored_year
            except (ValueError, IndexError, Exception):
                pass

        # Serialise route dicts to JSON for data attributes in the template.
        # The JS layer reads these to avoid Django/JS interpolation conflicts.
        # Serializar los dicts de ruta a JSON para data attributes en el template.
        # La capa JS los lee para evitar conflictos de interpolacion Django/JS.
        # Use float() for distance_km so json.dumps always serialises
        # with a decimal point regardless of Django locale settings.
        # Usar float() para distance_km para que json.dumps serialice
        # siempre con punto decimal independientemente del locale Django.
        route_with_tolls_json = json.dumps({
            "distance_km": float(route_with_tolls["distance_km"]),
            "has_tolls": route_with_tolls["has_tolls"],
            "encoded_polyline": route_with_tolls["encoded_polyline"],
        })
        route_without_tolls_json = json.dumps(
            {
                "distance_km": float(
                    route_without_tolls["distance_km"]
                ),
                "has_tolls": route_without_tolls["has_tolls"],
                "encoded_polyline": (
                    route_without_tolls["encoded_polyline"]
                ),
            }
            if route_without_tolls else None
        )

        return render(request, self.template_name, {
            "route_with_tolls":         route_with_tolls,
            "route_without_tolls":      route_without_tolls,
            "show_dual":                show_dual,
            "route_with_tolls_json":    route_with_tolls_json,
            "route_without_tolls_json": route_without_tolls_json,
            "calendar_warning":         calendar_warning,
            "calendar_warning_year":    calendar_warning_year,
            "service_year":             service_date.year,
            "error":                    None,
        })


# ---------------------------------------------------------------------------
# HTMX endpoint — multi-leg route planner init (GET)
# Endpoint HTMX — inicializacion del planificador multi-parada (GET)
# ---------------------------------------------------------------------------


class BudgetRouteMultilegInitView(AssistanceRequiredMixin, View):
    """
    HTMX GET endpoint. Returns the multi-leg route planner fragment
    initialised with the base coordinates for the given base_id.
    The fragment renders the interactive Google Maps panel, the waypoint
    list and the Calcular ruta button.

    Query params: base_id (int), insurer_id (int, optional for context).
    ---
    Endpoint GET HTMX. Devuelve el fragmento del planificador multi-parada
    inicializado con las coordenadas de la base para el base_id dado.
    El fragmento renderiza el panel Google Maps interactivo, la lista de
    paradas y el boton Calcular ruta.

    Parametros de consulta: base_id (int), insurer_id (int, opcional).
    """

    template_name = "budgets/_route_multileg_fragment.html"

    def get(self, request):
        """
        Return the planner fragment with base coordinates and Maps API key.
        ---
        Devuelve el fragmento del planificador con coordenadas de base y
        clave de Maps API.
        """
        company_user = _get_company_user(request)
        base_id = request.GET.get("base_id", "").strip()

        if not base_id:
            return render(request, self.template_name, {
                "error": "base_id requerido.",
                "base": None,
            })

        try:
            base = Base.objects.get(
                pk=int(base_id),
                company=company_user.company,
            )
        except (Base.DoesNotExist, ValueError):
            return render(request, self.template_name, {
                "error": "Base no encontrada.",
                "base": None,
            })

        if not base.latitude or not base.longitude:
            return render(request, self.template_name, {
                "error": (
                    f"La base '{base.name}' no tiene coordenadas. "
                    "Geocodifícala desde el panel antes de usar el "
                    "planificador de ruta."
                ),
                "base": base,
            })

        google_maps_api_key = os.environ.get("GOOGLE_MAPS_API_KEY", "")
        google_maps_map_id = os.environ.get("GOOGLE_MAPS_MAP_ID", "")

        return render(request, self.template_name, {
            "base": base,
            "base_lat": float(base.latitude),
            "base_lng": float(base.longitude),
            "google_maps_api_key": google_maps_api_key,
            "google_maps_map_id": google_maps_map_id,
            "error": None,
        })


# ---------------------------------------------------------------------------
# HTMX endpoint — multi-leg route calculation (POST)
# Endpoint HTMX — calculo de ruta multi-parada (POST)
# ---------------------------------------------------------------------------


class BudgetWaypointView(AssistanceRequiredMixin, View):
    """
    HTMX POST endpoint. Receives base_id, waypoints_json, service_date and
    service_time. Calls calculate_route_multileg() and returns the result
    fragment with distance, phase km, overnight flag and encoded polyline.
    Returns an error context on RouteCalculationError or validation failure.

    POST params:
      - base_id:        int
      - waypoints_json: JSON string — list of {lat, lng, label, address,
                        is_base_return}
      - service_date:   ISO date string (YYYY-MM-DD)
      - service_time:   ISO time string (HH:MM)
    ---
    Endpoint POST HTMX. Recibe base_id, waypoints_json, service_date y
    service_time. Llama a calculate_route_multileg() y devuelve el fragmento
    de resultado con distancia, km por fase, flag de pernocta y polyline.
    Devuelve contexto de error en RouteCalculationError o fallo de validacion.
    """

    template_name = "budgets/_route_multileg_fragment.html"

    def post(self, request):
        """
        Validate inputs, call calculate_route_multileg() and return a
        JsonResponse with all route fields for wizard_map.js confirmRoute().
        ---
        Valida los parametros, llama a calculate_route_multileg() y devuelve
        un JsonResponse con todos los campos de ruta para confirmRoute().
        """
        import datetime as _dt
        import json as _json
        from django.http import JsonResponse

        company_user = _get_company_user(request)

        # Request body is JSON (sent by wizard_map.js fetch with
        # Content-Type: application/json). request.POST is empty.
        # El body es JSON enviado por wizard_map.js — request.POST está vacío.
        try:
            body = _json.loads(request.body)
        except (_json.JSONDecodeError, ValueError):
            return JsonResponse(
                {"error": "Cuerpo de la petición inválido."}, status=400
            )

        base_id            = str(body.get("base_id", "")).strip()
        waypoints_json_str = body.get("waypoints_json", "")
        service_date_str   = str(body.get("service_date", "")).strip()
        service_time_str   = str(body.get("service_time", "")).strip()

        # Validate required fields.
        # Validar campos obligatorios.
        if not all([base_id, waypoints_json_str,
                    service_date_str, service_time_str]):
            return JsonResponse(
                {
                    "error": (
                        "Faltan datos: base, paradas, fecha y hora "
                        "son obligatorios."
                    )
                },
                status=400,
            )

        # Resolve base.
        # Resolver la base.
        try:
            base = Base.objects.get(
                pk=int(base_id),
                company=company_user.company,
            )
        except (Base.DoesNotExist, ValueError):
            return JsonResponse({"error": "Base no encontrada."}, status=404)

        # Parse waypoints JSON.
        # Parsear JSON de waypoints.
        try:
            waypoints: list[dict] = _json.loads(waypoints_json_str)
            if not isinstance(waypoints, list) or not waypoints:
                raise ValueError("La lista de paradas esta vacia.")
        except (ValueError, _json.JSONDecodeError) as exc:
            logger.warning(
                "# [budgets] Paradas inválidas. waypoints_json_str=%r: %s",
                waypoints_json_str, exc, exc_info=True,
            )
            return JsonResponse(
                {"error": "La lista de paradas no es válida."}, status=400
            )

        # Parse service datetime.
        # Parsear fecha y hora del servicio.
        try:
            service_date     = _dt.date.fromisoformat(service_date_str)
            service_time_obj = _dt.time.fromisoformat(service_time_str)
            service_datetime = _dt.datetime.combine(
                service_date, service_time_obj,
            )
        except Exception as exc:
            logger.warning(
                "# [budgets] Fecha u hora inválidas. "
                "service_date_str=%r service_time_str=%r: %s",
                service_date_str, service_time_str, exc, exc_info=True,
            )
            return JsonResponse(
                {"error": "La fecha o la hora del servicio no son válidas."},
                status=400,
            )

        # Any date (past, present or future) is now accepted — the
        # service layer omits departureTime for non-future dates since
        # Google's Routes API rejects a past departureTime for DRIVE mode.
        # Cualquier fecha es aceptada ahora — la capa de servicios omite
        # departureTime para fechas no futuras, ya que la Routes API de
        # Google rechaza un departureTime pasado en modo DRIVE.

        api_key = os.environ.get("GOOGLE_MAPS_API_KEY", "")

        from budgets.services import calculate_route_multileg
        result = calculate_route_multileg(
            base=base,
            waypoints=waypoints,
            service_datetime=service_datetime,
            api_key=api_key,
            company=company_user.company,
        )

        if result.get("error"):
            return JsonResponse({"error": result["error"]}, status=502)

        return JsonResponse({
            "route_distance_km":      result["distance_km"],
            "km_phase1":              result["km_phase1"],
            "km_phase2":              result["km_phase2"],
            "is_overnight":           result["is_overnight"],
            "has_tolls":              result["has_tolls"],
            "route_toll_budget_cost": result.get("route_toll_budget_cost"),
            "encoded_polyline":       result.get("encoded_polyline"),
            "encoded_polyline_alt":   result.get("encoded_polyline_alt", ""),
            "distance_km_alt":        result.get("distance_km_alt"),
        })


# ---------------------------------------------------------------------------
# HTMX partial — vehicle types for selected insurer
# ---------------------------------------------------------------------------

class BudgetVehicleTypesView(AssistanceRequiredMixin, View):
    """
    HTMX partial view. Returns the vehicle type dropdown fragment for the
    insurer selected in step 1 of the wizard.
    ---
    Vista parcial HTMX. Devuelve el fragmento de desplegable de tipos de
    vehiculo para la aseguradora seleccionada en el paso 1 del asistente.
    """

    template_name = "budgets/_vehicle_types_fragment.html"

    def get(self, request):
        """
        Return vehicle types for the given insurer_id query parameter.
        ---
        Devuelve los tipos de vehiculo para el parametro insurer_id.
        """
        company_user = _get_company_user(request)
        insurer_id = request.GET.get("insurer_id")
        if not insurer_id:
            return HttpResponseBadRequest("insurer_id requerido.")
        vehicle_types = VehicleType.objects.filter(
            insurer_id=insurer_id,
            insurer__company=company_user.company,
            is_active=True,
        ).order_by("sort_order", "name")
        return render(request, self.template_name, {
            "vehicle_types": vehicle_types,
        })


# ---------------------------------------------------------------------------
# HTMX partial — base selector for selected insurer
# ---------------------------------------------------------------------------

class BudgetBasesView(AssistanceRequiredMixin, View):
    """
    HTMX partial view. Returns the base selector fragment for the insurer
    selected in step 1 of the wizard. If the insurer has exactly one active
    base, returns a hidden input with that base pk. If it has more than one,
    returns a visible dropdown. If it has none, returns an empty fragment.
    ---
    Vista parcial HTMX. Devuelve el fragmento de selector de base para la
    aseguradora seleccionada en el paso 1 del asistente. Si la aseguradora
    tiene exactamente una base activa, devuelve un input oculto con ese pk.
    Si tiene mas de una, devuelve un desplegable visible. Si no tiene ninguna,
    devuelve un fragmento vacio.
    """

    template_name = "budgets/_base_selector_fragment.html"

    def get(self, request):
        """
        Return base selector fragment for the given insurer_id query parameter.
        ---
        Devuelve el fragmento de selector de base para el parametro insurer_id.
        """
        company_user = _get_company_user(request)
        insurer_id = request.GET.get("insurer_id")
        if not insurer_id:
            return HttpResponseBadRequest("insurer_id requerido.")
        base_ids = InsurerBase.objects.filter(
            insurer_id=insurer_id,
            insurer__company=company_user.company,
            is_active=True,
            base__is_active=True,
        ).values_list("base_id", flat=True)
        bases = Base.objects.filter(pk__in=base_ids).order_by("name")
        return render(request, self.template_name, {
            "bases": bases,
        })


# ---------------------------------------------------------------------------
# HTMX partial — optional concepts for selected insurer + vehicle type
# ---------------------------------------------------------------------------

class BudgetOptionalConceptsView(AssistanceRequiredMixin, View):
    """
    HTMX partial view. Returns the optional concepts form fragment for
    the selected insurer + vehicle type combination.
    ---
    Vista parcial HTMX. Devuelve el fragmento de formulario de conceptos
    opcionales para la combinacion de aseguradora y tipo de vehiculo seleccionados.
    """

    template_name = "budgets/_optional_concepts_fragment.html"

    def get(self, request):
        """
        Return optional concept inputs for the given insurer_id and
        vehicle_type_id query parameters.
        ---
        Devuelve los campos opcionales para los parametros insurer_id
        y vehicle_type_id dados.
        """
        insurer_id = request.GET.get("insurer_id")
        vehicle_type_id = request.GET.get("vehicle_type_id")
        if not insurer_id or not vehicle_type_id:
            return HttpResponseBadRequest(
                "insurer_id y vehicle_type_id requeridos."
            )
        available_codes = _get_optional_concepts(
            int(insurer_id),
            int(vehicle_type_id),
        )
        has_loaded = _has_loaded_surcharge(int(insurer_id))
        concepts = [
            OPTIONAL_CONCEPT_META[code]
            for code in available_codes
            if code in OPTIONAL_CONCEPT_META
        ]
        return render(request, self.template_name, {
            "concepts": concepts,
            "has_loaded_surcharge": has_loaded,
        })


# ---------------------------------------------------------------------------
# HTMX endpoint — wizard steps 3-9 + submit (server-side step resolution)
# Endpoint HTMX — pasos 3-9 del wizard + submit (resolucion server-side)
# ---------------------------------------------------------------------------

class BudgetStepsView(AssistanceRequiredMixin, View):
    """
    HTMX GET endpoint. Receives insurer_id, base_id and vehicle_type_id.
    Resolves optional concepts, loaded surcharge, always_apply_iva flag
    and base coordinates server-side. Returns the full wizard steps
    fragment (steps 3-9 + submit) via _wizard_steps_fragment.html.
    ---
    Endpoint GET HTMX. Recibe insurer_id, base_id y vehicle_type_id.
    Resuelve conceptos opcionales, recargo cargado, flag always_apply_iva
    y coordenadas de base en el servidor. Devuelve el fragmento completo
    de pasos del wizard (pasos 3-9 + submit) via _wizard_steps_fragment.html.
    """

    template_name = "budgets/_wizard_steps_fragment.html"

    def get(self, request):
        """
        Resolve all step context from insurer_id, base_id and vehicle_type_id
        query params and return the rendered steps fragment.
        ---
        Resuelve todo el contexto de pasos desde los parametros insurer_id,
        base_id y vehicle_type_id y devuelve el fragmento de pasos renderizado.
        """
        company_user = _get_company_user(request)
        insurer_id      = request.GET.get("insurer_id", "").strip()
        base_id         = request.GET.get("base_id", "").strip()
        vehicle_type_id = request.GET.get("vehicle_type_id", "").strip()

        if not insurer_id or not vehicle_type_id:
            return HttpResponseBadRequest("insurer_id y vehicle_type_id requeridos.")

        # Resolve optional concepts and loaded surcharge.
        # Resolver conceptos opcionales y recargo de cargado.
        available_codes = _get_optional_concepts(int(insurer_id), int(vehicle_type_id))
        concepts = [
            OPTIONAL_CONCEPT_META[code]
            for code in available_codes
            if code in OPTIONAL_CONCEPT_META
        ]
        has_loaded = _has_loaded_surcharge(int(insurer_id))

        # Resolve always_apply_iva for the selected insurer.
        # Resolver always_apply_iva para la aseguradora seleccionada.
        try:
            insurer = Insurer.objects.get(
                pk=int(insurer_id),
                company=company_user.company,
                is_active=True,
            )
            always_apply_iva = insurer.always_apply_iva
        except Insurer.DoesNotExist:
            always_apply_iva = False

        # Resolve loaded surcharge percent for display.
        # Resolver porcentaje de recargo de cargado para mostrar.
        loaded_percent = ""
        if has_loaded:
            try:
                tariff = InsurerTariff.objects.get(
                    insurer_id=int(insurer_id),
                    valid_to__isnull=True,
                )
                line = tariff.lines.filter(
                    concept__code=TariffConcept.CODE_LOADED_PERCENT
                ).first()
                if line:
                    loaded_percent = line.price
            except InsurerTariff.DoesNotExist:
                pass

        # Resolve whether the selected base has coordinates.
        # Resolver si la base seleccionada tiene coordenadas.
        base_has_coords = False
        base_lat = None
        base_lng = None
        base_name = ""
        base_id_int = None
        if base_id:
            try:
                base = Base.objects.get(
                    pk=int(base_id),
                    company=company_user.company,
                )
                base_has_coords = bool(base.latitude and base.longitude)
                if base_has_coords:
                    base_lat    = float(base.latitude)
                    base_lng    = float(base.longitude)
                    base_name   = base.name
                    base_id_int = base.pk
            except Base.DoesNotExist:
                pass

        google_maps_map_id = os.environ.get("GOOGLE_MAPS_MAP_ID", "")

        # Comprobar si el tipo de vehículo tiene líneas de tarifa para los
        # conceptos obligatorios del motor (DEPARTURE, KM_NORMAL, KM_LONG).
        # Sin ellas el motor produce un presupuesto sin tarifa base.
        # Check whether the vehicle type has tariff lines for the mandatory
        # motor concepts (DEPARTURE, KM_NORMAL, KM_LONG).
        # Without them the engine produces a budget without base tariff.
        no_tariff_lines = False
        _MANDATORY_CODES = [
            TariffConcept.CODE_DEPARTURE,
            TariffConcept.CODE_KM_NORMAL,
            TariffConcept.CODE_KM_LONG,
            TariffConcept.CODE_SERVICE_LOCAL,
        ]
        try:
            tariff = InsurerTariff.objects.get(
                insurer_id=int(insurer_id),
                valid_to__isnull=True,
            )
            has_mandatory = tariff.lines.filter(
                models_q_vehicle(int(vehicle_type_id)),
                concept__code__in=_MANDATORY_CODES,
            ).exists()
            no_tariff_lines = not has_mandatory
        except InsurerTariff.DoesNotExist:
            no_tariff_lines = True

        # Determine whether the operator has provided a service time (step 2b).
        # Used by the template for step 6 badge and pre-check logic.
        # Determinar si el operario ha introducido la hora del servicio (paso 2b).
        # El template lo usa para el badge y el pre-marcado del paso 6.
        service_time_str  = request.GET.get("service_time", "").strip()
        service_date_str  = request.GET.get("service_date", "").strip()
        insurer_id_str    = request.GET.get("insurer_id", "").strip()

        # Resolve toll segments for the Manual mode itemized toll table
        # (troncal/salida, AP-7 "Autopista del Sol" y AP-46 "Autopista de
        # las Pedrizas"). Price shown is informational — the real price
        # used at calculation time is re-resolved fresh in
        # calculate_budget() via _compute_manual_toll_cost().
        # ---
        # Resolver tramos de peaje para la tabla itemizada del modo Manual
        # (troncal/salida, AP-7 "Autopista del Sol" y AP-46 "Autopista de
        # las Pedrizas"). El precio mostrado es informativo — el precio
        # real usado en el cálculo se resuelve de nuevo en
        # calculate_budget() vía _compute_manual_toll_cost().
        price_field, _vehicle_label, _high_field = _resolve_toll_pricing(
            company_user.company
        )

        def _toll_sort_key(seg):
            """
            Orden pedido por Miguel Ángel: Casabermeja (AP-46) primero,
            luego todos los tramos troncal de AP-7, y las salidas al
            final. Se basa en el mismo patrón de nombres ya usado en
            _compute_toll_cost() para resolver puertas troncal/salida
            (origin_name = "<Lugar> Troncal" / "Salida <Lugar>").
            ---
            Order requested by Miguel Ángel: Casabermeja (AP-46) first,
            then all AP-7 troncal segments, salidas last. Based on the
            same naming pattern already used by _compute_toll_cost() to
            resolve troncal/salida gates (origin_name = "<Place>
            Troncal" / "Salida <Place>").
            """
            if seg.road_code == "AP-46":
                group = 0
            elif seg.origin_name.strip().lower().startswith("salida"):
                group = 2
            else:
                group = 1
            return (group, seg.origin_name)

        toll_segments = [
            {
                "id": seg.pk,
                "label": f"{seg.road_code} | {seg.origin_name} \u2192 {seg.dest_name}",
                "price": getattr(seg, price_field, None),
            }
            for seg in sorted(
                TollSegment.objects.filter(
                    road_code__in=["AP-7", "AP-46"],
                    is_active=True,
                ),
                key=_toll_sort_key,
            )
        ]

        return render(request, self.template_name, {
            "concepts":             concepts,
            "has_loaded_surcharge": has_loaded,
            "loaded_percent":       loaded_percent,
            "always_apply_iva":     always_apply_iva,
            "base_has_coords":      base_has_coords,
            "base_id":              base_id_int,
            "base_lat":             base_lat,
            "base_lng":             base_lng,
            "base_name":            base_name,
            "google_maps_map_id":   google_maps_map_id,
            "google_maps_api_key":  os.environ.get(
                "GOOGLE_MAPS_API_KEY", ""
            ),
            "no_tariff_lines":      no_tariff_lines,
            "toll_segments":        toll_segments,
        })


# ---------------------------------------------------------------------------
# Budget result — operator view (total only)
# ---------------------------------------------------------------------------

class BudgetResultView(AssistanceRequiredMixin, View):
    """
    Displays the budget result to the operator. Only the total amount is
    shown — no tariff lines, no unit prices, no breakdown.
    Provides ACCEPTED / REJECTED action buttons.
    ---
    Muestra el resultado del presupuesto al operario. Solo se muestra el
    importe total — sin lineas de tarifa, sin precios unitarios, sin desglose.
    Proporciona los botones de accion ACEPTADO / RECHAZADO.
    """

    template_name = "budgets/result.html"

    def get(self, request, pk):
        """
        Render the result page for the given budget pk.
        ---
        Renderiza la pagina de resultado para el pk de presupuesto dado.
        """
        company_user = _get_company_user(request)
        budget = get_object_or_404(
            Budget,
            pk=pk,
            company=company_user.company,
        )
        ctx = _build_base_context(request, {
            "budget": budget,
            "active_nav": "budgets_wizard",
        })
        return render(request, self.template_name, ctx)


# ---------------------------------------------------------------------------
# Budget status update — ACCEPTED / REJECTED
# ---------------------------------------------------------------------------

class BudgetStatusUpdateView(AssistanceRequiredMixin, View):
    """
    Updates the status of a budget to ACCEPTED or REJECTED.
    Called from the result view action buttons.
    ---
    Actualiza el estado de un presupuesto a ACCEPTED o REJECTED.
    Se llama desde los botones de accion de la vista de resultado.
    """

    def post(self, request, pk):
        """
        Set budget.status to the value provided in the POST body.
        Redirects back to the result view.
        ---
        Establece budget.status al valor proporcionado en el cuerpo POST.
        Redirige de vuelta a la vista de resultado.
        """
        company_user = _get_company_user(request)
        budget = get_object_or_404(
            Budget,
            pk=pk,
            company=company_user.company,
        )
        new_status = request.POST.get("status")
        if new_status not in (
            Budget.STATUS_ACCEPTED,
            Budget.STATUS_REJECTED,
        ):
            messages.error(request, "Estado no valido.")
            return redirect("budgets:result", pk=pk)
        budget.status = new_status
        budget.save(update_fields=["status"])
        return redirect("budgets:result", pk=pk)


# ---------------------------------------------------------------------------
# Budget history — ADMIN, or ASSISTANCE with can_view_budget_breakdown
# ---------------------------------------------------------------------------

class BudgetHistoryView(BudgetAuditAccessMixin, View):
    """
    Lists all budgets for the company with filters by insurer, date and status.
    Visible to ADMIN, and to ASSISTANCE users whose
    CompanyUser.can_view_budget_breakdown flag is True.
    ---
    Lista todos los presupuestos de la empresa con filtros por aseguradora,
    fecha y estado. Visible para ADMIN, y para usuarios ASSISTANCE cuyo
    flag CompanyUser.can_view_budget_breakdown sea True.
    """

    template_name = "budgets/history.html"

    def get(self, request):
        """
        Render the budget history list with optional filters.
        ---
        Renderiza la lista de historial de presupuestos con filtros opcionales.
        """
        company_user = _get_company_user(request)
        qs = Budget.objects.filter(
            company=company_user.company,
        ).select_related(
            "insurer",
            "vehicle_type",
            "operator__user",
            "tariff",
        ).annotate(
            iva_amount=ExpressionWrapper(
                F("total_amount_with_iva") - F("total_amount"),
                output_field=DecimalField(max_digits=10, decimal_places=2),
            )
        ).order_by("-created_at")

        # Optional filters from GET params.
        # Filtros opcionales desde parametros GET.
        insurer_id = request.GET.get("insurer")
        status = request.GET.get("status")
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")

        if insurer_id:
            qs = qs.filter(insurer_id=insurer_id)
        if status:
            qs = qs.filter(status=status)
        if date_from:
            qs = qs.filter(service_date__gte=date_from)
        if date_to:
            qs = qs.filter(service_date__lte=date_to)

        insurers = Insurer.objects.filter(
            company=company_user.company,
        ).order_by("name")

        ctx = _build_base_context(request, {
            "budgets": qs,
            "insurers": insurers,
            "status_choices": Budget.STATUS_CHOICES,
            "active_nav": "budgets_history",
            "filters": {
                "insurer": insurer_id,
                "status": status,
                "date_from": date_from,
                "date_to": date_to,
            },
        })
        return render(request, self.template_name, ctx)


# ---------------------------------------------------------------------------
# Budget detail — ADMIN, or ASSISTANCE with can_view_budget_breakdown
# ---------------------------------------------------------------------------

class BudgetDetailView(BudgetAuditAccessMixin, View):
    """
    Displays the full calculation breakdown of a budget for audit.
    Shows all BudgetLine records with concept, units, unit price and
    subtotal. Visible to ADMIN, and to ASSISTANCE users whose
    CompanyUser.can_view_budget_breakdown flag is True.
    ---
    Muestra el desglose completo del calculo de un presupuesto para
    auditoria. Muestra todos los registros BudgetLine con concepto,
    unidades, precio unitario y subtotal. Visible para ADMIN, y para
    usuarios ASSISTANCE cuyo flag CompanyUser.can_view_budget_breakdown
    sea True.
    """

    template_name = "budgets/detail.html"

    def get(self, request, pk):
        """
        Render the full breakdown for the given budget pk.
        ---
        Renderiza el desglose completo para el pk de presupuesto dado.
        """
        company_user = _get_company_user(request)
        budget = get_object_or_404(
            Budget,
            pk=pk,
            company=company_user.company,
        )
        lines = budget.lines.exclude(concept_code="IVA").order_by("sort_order")
        # Calculate iva_amount for the fiscal summary in the tfoot.
        # Only populated when apply_iva is True and total_amount_with_iva is set.
        # Calcula iva_amount para el resumen fiscal en el tfoot.
        # Solo se rellena cuando apply_iva es True y total_amount_with_iva esta establecido.
        from decimal import Decimal
        iva_amount = None
        if budget.apply_iva and budget.total_amount_with_iva is not None:
            iva_amount = (budget.total_amount_with_iva - budget.total_amount).quantize(
                Decimal("0.01")
            )
        ctx = _build_base_context(request, {
            "budget": budget,
            "lines": lines,
            "iva_amount": iva_amount,
            "active_nav": "budgets_history",
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
            "google_maps_map_id": os.environ.get("GOOGLE_MAPS_MAP_ID", ""),
        })
        return render(request, self.template_name, ctx)


# ---------------------------------------------------------------------------
# Insurer management views — ADMIN only
# Vistas de gestion de aseguradoras — solo ADMIN
# ---------------------------------------------------------------------------


class InsurerForm(django_forms.ModelForm):
    """
    ModelForm for the Insurer model. Used by InsurerCreateView and
    InsurerUpdateView. Scopes the queryset to the company of the
    authenticated CompanyUser — unique_together validation is handled
    by the model itself.
    ---
    ModelForm para el modelo Insurer. Usado por InsurerCreateView e
    InsurerUpdateView. No expone el campo company en el formulario —
    se asigna automaticamente en la vista desde el CompanyUser autenticado.
    La validacion unique_together la gestiona el propio modelo.
    """

    # Declare management_fee_percent with localize=True so Django uses
    # DECIMAL_SEPARATOR from settings (comma for ES locale).
    # Declarar management_fee_percent con localize=True para que Django use
    # DECIMAL_SEPARATOR de settings (coma para locale ES).
    management_fee_percent = django_forms.DecimalField(
        localize=True,
        required=False,
        initial=0,
        max_digits=5,
        decimal_places=2,
        label="Gastos de gestión (%)",
        help_text=(
            "Porcentaje de gastos de gestión aplicado sobre el total del presupuesto. "
            "0 si la aseguradora no aplica este concepto. Ejemplo: COVEI aplica 5%."
        ),
    )

    class Meta:
        model = Insurer
        fields = [
            "name",
            "insurer_company_name",
            "service_company_name",
            "code",
            "management_fee_percent",
            "surcharges_are_cumulative",
            "is_active",
            "is_insurance_company",
            "always_apply_iva",
            "special_night_holiday_tariff",
            "night_schedule",
            "notes",
        ]

    def clean_code(self):
        """
        Normalise the code field to uppercase and strip whitespace.
        ---
        Normaliza el campo code a mayusculas y elimina espacios.
        """
        return self.cleaned_data.get("code", "").strip().upper()



def _insurer_list_qs(company):
    """
    Return the annotated Insurer queryset for the given company,
    including vehicle_type_count and tariff_count annotations.
    ---
    Devuelve el queryset de Insurer anotado para la empresa dada,
    incluyendo las anotaciones vehicle_type_count y tariff_count.
    """
    return (
        Insurer.objects.filter(company=company)
        .annotate(
            vehicle_type_count=Count("vehicle_types", distinct=True),
            tariff_count=Count("tariffs", distinct=True),
        )
        .order_by("name")
    )


class InsurerListView(AdminRoleRequiredMixin, View):
    """
    Lists all insurers for the company. Supports live HTMX search by
    name/code and filter by active/inactive status.
    On HTMX request: returns only the table fragment partial.
    On full request: returns the full page.
    ---
    Lista todas las aseguradoras de la empresa. Soporta busqueda live
    HTMX por nombre/codigo y filtro por estado activa/inactiva.
    En peticion HTMX: devuelve solo el parcial de tabla.
    En peticion completa: devuelve la pagina completa.
    """

    template_name = "budgets/insurer_list.html"
    partial_template = "budgets/_insurer_table_fragment.html"

    def get(self, request):
        """
        Render the insurer list with optional search and status filters.
        ---
        Renderiza el listado de aseguradoras con filtros opcionales de
        busqueda y estado.
        """
        company_user = _get_company_user(request)
        qs = _insurer_list_qs(company_user.company)

        # Apply search filter.
        # Aplicar filtro de busqueda.
        q = request.GET.get("q", "").strip()
        if q:
            qs = qs.filter(
                Q(name__icontains=q) | Q(code__icontains=q)
            )

        # Apply status filter.
        # Aplicar filtro de estado.
        status = request.GET.get("status", "")
        if status == "active":
            qs = qs.filter(is_active=True)
        elif status == "inactive":
            qs = qs.filter(is_active=False)

        ctx = _build_base_context(request, {
            "insurers": qs,
            "active_nav": "budgets_insurers",
        })

        # HTMX partial swap — return only the table fragment.
        # Swap parcial HTMX — devolver solo el fragmento de tabla.
        if request.headers.get("HX-Request"):
            return render(request, self.partial_template, ctx)

        return render(request, self.template_name, ctx)


class InsurerCreateView(AdminRoleRequiredMixin, View):
    """
    Renders and processes the insurer creation form.
    On GET: renders an empty InsurerForm.
    On POST: validates, assigns company from the authenticated user
    and saves. Redirects to insurer list on success.
    ---
    Renderiza y procesa el formulario de creacion de aseguradora.
    En GET: renderiza un InsurerForm vacio.
    En POST: valida, asigna la company del usuario autenticado
    y guarda. Redirige al listado en caso de exito.
    """

    template_name = "budgets/insurer_form.html"

    def get(self, request):
        """
        Render the empty creation form with night schedules for the
        insurer night_schedule selector.
        ---
        Renderiza el formulario de creacion vacio con los horarios
        nocturnos para el selector night_schedule de la aseguradora.
        """
        company_user = _get_company_user(request)
        form = InsurerForm()
        night_schedules = NightSchedule.objects.filter(
            company=company_user.company,
            is_active=True,
        ).order_by("-is_default", "name")
        ctx = _build_base_context(request, {
            "form": form,
            "mode": "create",
            "active_nav": "budgets_insurers",
            "night_schedules": night_schedules,
        })
        return render(request, self.template_name, ctx)

    def post(self, request):
        """
        Validate and save the new insurer bound to the company.
        ---
        Valida y guarda la nueva aseguradora vinculada a la empresa.
        """
        company_user = _get_company_user(request)
        form = InsurerForm(request.POST)
        if form.is_valid():
            insurer = form.save(commit=False)
            insurer.company = company_user.company
            try:
                insurer.save()
                messages.success(
                    request,
                    f"Aseguradora '{insurer.name}' creada correctamente.",
                )
                return redirect("budgets:insurer_list")
            except Exception:
                messages.error(
                    request,
                    "El codigo introducido ya existe para esta empresa. "
                    "Usa un codigo diferente.",
                )
        ctx = _build_base_context(request, {
            "form": form,
            "mode": "create",
            "active_nav": "budgets_insurers",
        })
        return render(request, self.template_name, ctx)


class InsurerUpdateView(AdminRoleRequiredMixin, View):
    """
    Renders and processes the insurer edit form for an existing insurer.
    On GET: renders InsurerForm pre-filled with the insurer data.
    On POST: validates and saves. Redirects to insurer list on success.
    ---
    Renderiza y procesa el formulario de edicion de una aseguradora existente.
    En GET: renderiza el InsurerForm con los datos de la aseguradora.
    En POST: valida y guarda. Redirige al listado en caso de exito.
    """

    template_name = "budgets/insurer_form.html"

    def _build_edit_context(self, request, insurer, form):
        """
        Build the full accordion context for the edit view:
        active tariff, tariff history with line counts, tariff line groups
        and choice lists for the inline editing selects.
        ---
        Construye el contexto completo del acordeon para la vista de edicion:
        tarifa activa, historial de tarifas con conteo de lineas, grupos de
        lineas de tarifa y listas de choices para los selects de edicion inline.
        """
        from budgets.models import InsurerTariff
        from django.db.models import Count as _Count

        # Resolve active tariff (valid_to=None).
        # Resolver tarifa activa (valid_to=None).
        active_tariff = (
            InsurerTariff.objects
            .filter(insurer=insurer, valid_to__isnull=True)
            .first()
        )

        # Tariff version history (all non-active versions).
        # Historial de versiones de tarifa (todas las versiones no activas).
        tariff_history = (
            InsurerTariff.objects
            .filter(insurer=insurer, valid_to__isnull=False)
            .annotate(line_count=_Count('lines'))
            .order_by('-valid_from')
        )

        # Tariff lines grouped by vehicle type for accordion Panel 3.
        # Lineas de tarifa agrupadas por tipo de vehiculo para Panel 3.
        tariff_line_groups = []
        vehicle_types = []
        if active_tariff:
            lines = (
                active_tariff.lines
                .select_related('vehicle_type')
                .order_by('vehicle_type__sort_order', 'vehicle_type__name', 'concept')
            )
            vehicle_types = (
                insurer.vehicle_types
                .filter(is_active=True)
                .order_by('sort_order', 'name')
            )
            # Group lines by vehicle_type (None = general concepts).
            # Agrupar lineas por vehicle_type (None = conceptos generales).
            groups_dict = {}
            for line in lines:
                key = line.vehicle_type
                if key not in groups_dict:
                    groups_dict[key] = []
                groups_dict[key].append(line)
            # None (general) first, then sorted by vehicle type name.
            # None (general) primero, luego ordenado por nombre de tipo.
            if None in groups_dict:
                tariff_line_groups.append({
                    'vehicle_type': None,
                    'lines': groups_dict.pop(None),
                })
            for vt, vt_lines in groups_dict.items():
                tariff_line_groups.append({
                    'vehicle_type': vt,
                    'lines': vt_lines,
                })

        # Night schedules for the selector in the insurer form.
        # Horarios nocturnos para el selector en el formulario de aseguradora.
        night_schedules = NightSchedule.objects.filter(
            company=insurer.company,
            is_active=True,
        ).order_by("-is_default", "name")

        return _build_base_context(request, {
            'form': form,
            'insurer': insurer,
            'mode': 'edit',
            'active_nav': 'budgets_insurers',
            'active_tariff': active_tariff,
            'tariff_history': tariff_history,
            'tariff_line_groups': tariff_line_groups,
            'vehicle_types': vehicle_types,
            'vehicle_types_all': insurer.vehicle_types.order_by('sort_order', 'name'),
            'concept_choices': _get_concept_choices(insurer.company),
            'unit_choices': TariffLine.UNIT_CHOICES,
            'night_schedules': night_schedules,
        })

    def get(self, request, pk):
        """
        Render the edit form pre-filled with the insurer instance data.
        ---
        Renderiza el formulario de edicion con los datos de la instancia.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )
        form = InsurerForm(instance=insurer)
        ctx = self._build_edit_context(request, insurer, form)
        return render(request, self.template_name, ctx)

    def post(self, request, pk):
        """
        Validate and save the updated insurer data.
        On success redirects back to the edit view (not to the list)
        so the user can continue editing tariff and lines.
        ---
        Valida y guarda los datos actualizados de la aseguradora.
        En caso de exito redirige de vuelta a la vista de edicion (no al listado)
        para que el usuario pueda continuar editando tarifa y lineas.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )
        form = InsurerForm(request.POST, instance=insurer)
        if form.is_valid():
            try:
                form.save()
                messages.success(
                    request,
                    f"Aseguradora '{insurer.name}' actualizada correctamente.",
                )
                return redirect('budgets:insurer_update', pk=insurer.pk)
            except Exception:
                messages.error(
                    request,
                    "El codigo introducido ya existe para esta empresa. "
                    "Usa un codigo diferente.",
                )
        ctx = self._build_edit_context(request, insurer, form)
        return render(request, self.template_name, ctx)


# ---------------------------------------------------------------------------
# Night schedule management views — ADMIN only
# Vistas de gestión de horarios nocturnos — solo ADMIN
# ---------------------------------------------------------------------------


class NightScheduleForm(django_forms.ModelForm):
    """
    ModelForm for the NightSchedule model. Scoped to the company of the
    authenticated CompanyUser. company field is excluded and assigned
    automatically in the view.
    ---
    ModelForm para el modelo NightSchedule. Con ámbito de la empresa del
    CompanyUser autenticado. El campo company se excluye y se asigna
    automáticamente en la vista.
    """

    class Meta:
        model = NightSchedule
        fields = ["name", "night_start", "night_end", "is_default", "is_active"]
        widgets = {
            "night_start": django_forms.TimeInput(
                attrs={"type": "time"},
                format="%H:%M",
            ),
            "night_end": django_forms.TimeInput(
                attrs={"type": "time"},
                format="%H:%M",
            ),
        }


class TollSegmentForm(django_forms.ModelForm):
    """
    ModelForm for the TollSegment model. Exposes all editable fields
    for the CRUD panel. valid_from defaults to today if not provided.
    ---
    ModelForm para el modelo TollSegment. Expone todos los campos
    editables para el panel CRUD. valid_from toma hoy por defecto
    si no se proporciona.
    """

    class Meta:
        from budgets.models import TollSegment as _TollSegment
        model = _TollSegment
        fields = [
            "road_code",
            "section_name",
            "origin_name",
            "dest_name",
            "price_light",
            "price_heavy_1",
            "price_heavy_2",
            "price_light_high",
            "price_heavy_1_high",
            "price_heavy_2_high",
            "tariff_level",
            "has_free_night",
            "free_night_start",
            "free_night_end",
            "valid_from",
            "is_active",
            "season_high_start",
            "season_high_end",
        ]
        widgets = {
            "free_night_start": django_forms.TimeInput(
                attrs={"type": "time"},
                format="%H:%M",
            ),
            "free_night_end": django_forms.TimeInput(
                attrs={"type": "time"},
                format="%H:%M",
            ),
            "valid_from": django_forms.DateInput(
                attrs={"type": "date"},
                format="%Y-%m-%d",
            ),
        }


class NightScheduleListView(AdminRoleRequiredMixin, View):
    """
    Lists all NightSchedule records for the company and renders the
    creation form for inline modal creation. On POST creates a new
    record and redirects to the list.
    ---
    Lista todos los registros NightSchedule de la empresa y renderiza
    el formulario de creación. En POST crea un nuevo registro y redirige
    al listado.
    """

    template_name = "budgets/night_schedule_list.html"

    def _get_context(self, request, form=None):
        """
        Build the full context for the list view.
        ---
        Construye el contexto completo para la vista de lista.
        """
        cu = request.user.company_user
        schedules = NightSchedule.objects.filter(
            company=cu.company,
        ).order_by("-is_default", "name")
        return _build_base_context(request, {
            "schedules": schedules,
            "form": form or NightScheduleForm(),
            "active_nav": "night_schedules",
        })

    def get(self, request):
        """
        Render the full night schedule list page.
        ---
        Renderiza la página completa de lista de horarios nocturnos.
        """
        return render(
            request,
            self.template_name,
            self._get_context(request),
        )

    def post(self, request):
        """
        Create a new NightSchedule from the form. On success redirects
        to the list. On error re-renders with validation errors.
        ---
        Crea un nuevo NightSchedule desde el formulario. En caso de
        éxito redirige al listado. En error re-renderiza con errores.
        """
        cu = request.user.company_user
        form = NightScheduleForm(request.POST)
        if form.is_valid():
            schedule = form.save(commit=False)
            schedule.company = cu.company
            schedule.save()
            messages.success(
                request,
                f"Horario '{schedule.name}' creado correctamente.",
            )
            return redirect("budgets:night_schedule_list")
        return render(
            request,
            self.template_name,
            self._get_context(request, form=form),
        )


class NightScheduleUpdateView(AdminRoleRequiredMixin, View):
    """
    Renders and processes the edit form for an existing NightSchedule.
    On GET: renders the form pre-filled. On POST: validates, saves and
    redirects back to the list.
    ---
    Renderiza y procesa el formulario de edición de un NightSchedule.
    En GET: renderiza el formulario pre-rellenado. En POST: valida,
    guarda y redirige al listado.
    """

    template_name = "budgets/night_schedule_list.html"

    def get(self, request, pk):
        """
        Render the edit form for the given NightSchedule pk.
        ---
        Renderiza el formulario de edición para el pk dado.
        """
        cu = request.user.company_user
        schedule = get_object_or_404(
            NightSchedule, pk=pk, company=cu.company,
        )
        form = NightScheduleForm(instance=schedule)
        schedules = NightSchedule.objects.filter(
            company=cu.company,
        ).order_by("-is_default", "name")
        ctx = _build_base_context(request, {
            "schedules": schedules,
            "form": form,
            "edit_schedule": schedule,
            "active_nav": "budgets_insurers",
        })
        return render(request, self.template_name, ctx)

    def post(self, request, pk):
        """
        Validate and save the updated NightSchedule. Redirects to list.
        ---
        Valida y guarda el NightSchedule actualizado. Redirige al listado.
        """
        cu = request.user.company_user
        schedule = get_object_or_404(
            NightSchedule, pk=pk, company=cu.company,
        )
        form = NightScheduleForm(request.POST, instance=schedule)
        if form.is_valid():
            form.save()
            messages.success(
                request,
                f"Horario '{schedule.name}' actualizado correctamente.",
            )
            return redirect("budgets:night_schedule_list")
        schedules = NightSchedule.objects.filter(
            company=cu.company,
        ).order_by("-is_default", "name")
        ctx = _build_base_context(request, {
            "schedules": schedules,
            "form": form,
            "edit_schedule": schedule,
            "active_nav": "budgets_insurers",
        })
        return render(request, self.template_name, ctx)


class NightScheduleDeleteView(AdminRoleRequiredMixin, View):
    """
    Deletes a NightSchedule if it has no linked Insurer records.
    ---
    Elimina un NightSchedule si no tiene registros Insurer vinculados.
    """

    def post(self, request, pk):
        """
        Delete the NightSchedule if no insurers reference it.
        ---
        Elimina el NightSchedule si no hay aseguradoras que lo referencien.
        """
        cu = request.user.company_user
        schedule = get_object_or_404(
            NightSchedule, pk=pk, company=cu.company,
        )
        if schedule.insurers.exists():
            messages.error(
                request,
                f"No se puede eliminar '{schedule.name}': "
                "tiene aseguradoras vinculadas. Desvínculalas primero.",
            )
        else:
            name = schedule.name
            schedule.delete()
            messages.success(
                request,
                f"Horario '{name}' eliminado correctamente.",
            )
        return redirect("budgets:night_schedule_list")


class InsurerToggleView(AdminRoleRequiredMixin, View):
    """
    HTMX endpoint. Toggles the is_active flag of an insurer and returns
    the updated badge fragment for inline swap in the list table.
    ---
    Endpoint HTMX. Alterna el flag is_active de una aseguradora y devuelve
    el fragmento de badge actualizado para swap inline en la tabla del listado.
    """

    def post(self, request, pk):
        """
        Toggle is_active and return the badge partial.
        ---
        Alterna is_active y devuelve el parcial del badge.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )
        insurer.is_active = not insurer.is_active
        insurer.save(update_fields=["is_active", "updated_at"])
        return render(
            request,
            "budgets/_insurer_badge_fragment.html",
            {"insurer": insurer},
        )


class InsurerDeleteView(AdminRoleRequiredMixin, View):
    """
    Deletes an insurer and all its related data (vehicle types, tariffs,
    tariff lines) via CASCADE. Redirects to insurer list with a success
    or error message.

    SpecialRateTariff.insurer_tariff and SpecialRateLine.vehicle_type use
    on_delete=PROTECT (by design, to avoid accidental loss of night/holiday
    rate history via unrelated cascades). Because of this, a direct
    insurer.delete() fails with ProtectedError whenever the insurer has an
    active SpecialRateTariff. This view explicitly deletes the
    SpecialRateTariff row(s) for the insurer's tariffs first (which
    CASCADEs their SpecialRateLine rows), then deletes the insurer, so the
    remaining CASCADE chain (VehicleType, InsurerTariff, TariffLine,
    InsurerBase) completes cleanly.
    ---
    Elimina una aseguradora y todos sus datos relacionados (tipos de vehiculo,
    tarifas, lineas de tarifa) via CASCADE. Redirige al listado con un mensaje
    de exito o error.

    SpecialRateTariff.insurer_tariff y SpecialRateLine.vehicle_type usan
    on_delete=PROTECT (por diseño, para evitar pérdida accidental del
    histórico de tarifas nocturno/festivo por cascadas ajenas). Por esto,
    un insurer.delete() directo falla con ProtectedError siempre que la
    aseguradora tenga una SpecialRateTariff activa. Esta vista borra
    explícitamente la(s) SpecialRateTariff de las tarifas de la aseguradora
    primero (lo que arrastra sus SpecialRateLine vía CASCADE), y solo
    entonces borra la aseguradora, para que el resto de la cadena CASCADE
    (VehicleType, InsurerTariff, TariffLine, InsurerBase) se complete
    limpiamente.
    """

    def post(self, request, pk):
        """
        Delete the insurer instance. Redirect to insurer list.
        ---
        Elimina la instancia de aseguradora. Redirige al listado.
        """
        from budgets.models import InsurerTariff, SpecialRateTariff

        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )
        name = insurer.name
        try:
            with transaction.atomic():
                # Explicitly remove SpecialRateTariff rows first (PROTECT
                # FK blocks the automatic CASCADE otherwise).
                # Eliminar primero explícitamente las SpecialRateTariff
                # (la FK PROTECT bloquea el CASCADE automático si no).
                SpecialRateTariff.objects.filter(
                    insurer_tariff__insurer=insurer
                ).delete()
                insurer.delete()
            messages.success(
                request,
                f"Aseguradora '{name}' eliminada correctamente.",
            )
        except Exception as exc:
            logger.error(
                "# [budgets] Error eliminando aseguradora pk=%r name=%r: %s",
                insurer.pk, name, exc, exc_info=True,
            )
            messages.error(
                request,
                f"No se pudo eliminar la aseguradora '{name}'. Puede que "
                "tenga tarifas o presupuestos asociados.",
            )
        return redirect("budgets:insurer_list")


# ---------------------------------------------------------------------------
# Tariff line inline save — HTMX POST, saves a single field change and
# returns the updated row fragment.
# Guardado inline de linea de tarifa — HTMX POST, guarda un cambio de campo
# y devuelve el fragmento de fila actualizado.
# ---------------------------------------------------------------------------


class TariffLineSaveView(AdminRoleRequiredMixin, View):
    """
    HTMX endpoint. Receives a POST with the full row field values for a
    TariffLine, saves them and returns the updated row fragment for swap.
    ---
    Endpoint HTMX. Recibe un POST con todos los valores de campo de la fila
    de una TariffLine, los guarda y devuelve el fragmento de fila actualizado.
    """

    def post(self, request, pk):
        """
        Save TariffLine fields from POST and return updated row fragment.
        ---
        Guarda los campos de TariffLine desde POST y devuelve el fragmento
        de fila actualizado.
        """
        from budgets.models import InsurerTariff
        company_user = _get_company_user(request)
        line = get_object_or_404(
            TariffLine,
            pk=pk,
            tariff__insurer__company=company_user.company,
        )
        data = request.POST

        # Update all editable fields atomically.
        # Actualizar todos los campos editables de forma atomica.
        concept_code = data.get("concept", "")
        if concept_code:
            line.concept = get_object_or_404(TariffConcept, code=concept_code)
        line.unit = data.get("unit", line.unit)
        line.price = _parse_decimal(data.get("price", ""), default="0")
        km_threshold_raw = data.get("km_threshold", "").strip().replace(",", ".")
        line.km_threshold = km_threshold_raw if km_threshold_raw else None
        min_units_raw = data.get("min_units", "").strip().replace(",", ".")
        line.min_units = min_units_raw if min_units_raw else None
        # Checkbox: present in POST = checked, absent = unchecked.
        # Checkbox: presente en POST = marcado, ausente = desmarcado.
        line.requires_authorization = "requires_authorization" in data
        line.save()

        ctx = {
            "line": line,
            "concept_choices": _get_concept_choices(line.tariff.insurer.company),
            "unit_choices": TariffLine.UNIT_CHOICES,
            "csrf_token": request.META.get("CSRF_COOKIE", ""),
        }
        return render(
            request,
            "budgets/_tariff_line_row_fragment.html",
            ctx,
        )


# ---------------------------------------------------------------------------
# Tariff line delete — HTMX POST, deletes the line and returns empty string
# so the row is removed from the DOM.
# Eliminacion de linea de tarifa — HTMX POST, elimina la linea y devuelve
# cadena vacia para que la fila desaparezca del DOM.
# ---------------------------------------------------------------------------


class TariffLineDeleteView(AdminRoleRequiredMixin, View):
    """
    HTMX endpoint. Deletes a TariffLine and returns an empty response
    so HTMX replaces the row with nothing (outerHTML swap).
    ---
    Endpoint HTMX. Elimina una TariffLine y devuelve una respuesta vacia
    para que HTMX sustituya la fila por nada (swap outerHTML).
    """

    def post(self, request, pk):
        """
        Delete the TariffLine and return empty 200 response.
        ---
        Elimina la TariffLine y devuelve respuesta 200 vacia.
        """
        company_user = _get_company_user(request)
        line = get_object_or_404(
            TariffLine,
            pk=pk,
            tariff__insurer__company=company_user.company,
        )
        line.delete()
        from django.http import HttpResponse
        return HttpResponse("")


# ---------------------------------------------------------------------------
# Tariff line add form — HTMX GET, returns the add-line inline form fragment.
# Formulario de adicion de linea — HTMX GET, devuelve el fragmento del
# formulario inline de nueva linea.
# ---------------------------------------------------------------------------


class TariffLineAddFormView(AdminRoleRequiredMixin, View):
    """
    HTMX GET endpoint. Returns the inline add-line form fragment for the
    given active tariff. The form is injected into #new-line-row.
    ---
    Endpoint HTMX GET. Devuelve el fragmento de formulario inline de nueva
    linea para la tarifa activa dada. El formulario se inyecta en #new-line-row.
    """

    def get(self, request, pk):
        """
        Return the add-line form fragment for the given tariff pk.
        ---
        Devuelve el fragmento de formulario de nueva linea para el pk de tarifa.
        """
        from budgets.models import InsurerTariff
        company_user = _get_company_user(request)
        tariff = get_object_or_404(
            InsurerTariff,
            pk=pk,
            insurer__company=company_user.company,
        )
        vehicle_types = (
            tariff.insurer.vehicle_types
            .filter(is_active=True)
            .order_by("sort_order", "name")
        )
        ctx = {
            "tariff": tariff,
            "vehicle_types": vehicle_types,
            "concept_choices": _get_concept_choices(company_user.company),
            "unit_choices": TariffLine.UNIT_CHOICES,
        }
        return render(
            request,
            "budgets/_tariff_line_add_form_fragment.html",
            ctx,
        )


# ---------------------------------------------------------------------------
# Tariff line add — HTMX POST, creates a new TariffLine and returns the
# new row fragment inserted at the top of the table.
# Adicion de linea — HTMX POST, crea una nueva TariffLine y devuelve el
# fragmento de fila nuevo insertado en la tabla.
# ---------------------------------------------------------------------------


class TariffLineAddView(AdminRoleRequiredMixin, View):
    """
    HTMX POST endpoint. Creates a new TariffLine for the given active tariff
    and returns the new row fragment. Also clears the add form by returning
    an empty string swapped into #new-line-row.
    ---
    Endpoint HTMX POST. Crea una nueva TariffLine para la tarifa activa dada
    y devuelve el fragmento de fila nuevo. Tambien limpia el formulario de
    adicion devolviendo cadena vacia en #new-line-row.
    """

    def post(self, request, pk):
        """
        Create a new TariffLine from POST data and return updated row.
        ---
        Crea una nueva TariffLine desde los datos POST y devuelve la fila.
        """
        from budgets.models import InsurerTariff
        from django.http import HttpResponse
        company_user = _get_company_user(request)
        tariff = get_object_or_404(
            InsurerTariff,
            pk=pk,
            insurer__company=company_user.company,
        )
        data = request.POST
        vehicle_type_id = data.get("vehicle_type_id", "").strip()
        vehicle_type = None
        if vehicle_type_id:
            vehicle_type = get_object_or_404(
                VehicleType,
                pk=int(vehicle_type_id),
                insurer=tariff.insurer,
            )
        km_threshold = data.get("km_threshold", "").strip()
        min_units = data.get("min_units", "").strip()
        concept_code = data.get("concept", TariffConcept.CODE_DEPARTURE)
        concept_obj = get_object_or_404(
            TariffConcept,
            code=concept_code,
        )
        line = TariffLine.objects.create(
            tariff=tariff,
            vehicle_type=vehicle_type,
            concept=concept_obj,
            unit=data.get("unit", TariffLine.UNIT_FIXED),
            price=_parse_decimal(data.get("price", ""), default="0"),
            km_threshold=(
                _parse_decimal(km_threshold) if km_threshold else None
            ),
            min_units=(
                _parse_decimal(min_units) if min_units else None
            ),
            requires_authorization="requires_authorization" in data,
        )
        ctx = {
            "line": line,
            "concept_choices": _get_concept_choices(company_user.company),
            "unit_choices": TariffLine.UNIT_CHOICES,
            "csrf_token": request.META.get("CSRF_COOKIE", ""),
        }
        return render(
            request,
            "budgets/_tariff_line_row_fragment.html",
            ctx,
        )


# ---------------------------------------------------------------------------
# Insurer tariff create — creates a new tariff version, closes the active one.
# Creacion de nueva version de tarifa — crea una nueva version y cierra la activa.
# ---------------------------------------------------------------------------


class InsurerTariffCreateView(AdminRoleRequiredMixin, View):
    """
    Creates a new InsurerTariff version for the given insurer.
    Closes the currently active tariff (valid_to = valid_from - 1 day).
    Redirects back to the insurer edit view.
    ---
    Crea una nueva version de InsurerTariff para la aseguradora dada.
    Cierra la tarifa actualmente activa (valid_to = valid_from - 1 dia).
    Redirige de vuelta a la vista de edicion de la aseguradora.
    """

    def post(self, request, pk):
        """
        Create the new tariff version and close the previous active one.
        ---
        Crea la nueva version de tarifa y cierra la anterior activa.
        """
        from budgets.models import InsurerTariff
        from datetime import date, timedelta
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )
        year_raw = request.POST.get("year", "").strip()
        valid_from_raw = request.POST.get("valid_from", "").strip()
        if not year_raw or not valid_from_raw:
            messages.error(request, "El ano y la fecha de inicio son obligatorios.")
            return redirect("budgets:insurer_update", pk=pk)
        try:
            year = int(year_raw)
            valid_from = date.fromisoformat(valid_from_raw)
        except ValueError:
            messages.error(request, "Formato de ano o fecha invalido.")
            return redirect("budgets:insurer_update", pk=pk)

        with transaction.atomic():
            # Close the current active tariff.
            # Cerrar la tarifa activa actual.
            active = InsurerTariff.objects.filter(
                insurer=insurer,
                valid_to__isnull=True,
            ).first()
            if active:
                active.valid_to = valid_from - timedelta(days=1)
                active.save(update_fields=["valid_to"])
            # Create the new tariff version.
            # Crear la nueva version de tarifa.
            InsurerTariff.objects.create(
                insurer=insurer,
                year=year,
                valid_from=valid_from,
                valid_to=None,
                notes="",
            )
        messages.success(
            request,
            f"Nueva version de tarifa {year} creada correctamente.",
        )
        return redirect("budgets:insurer_update", pk=pk)


# ---------------------------------------------------------------------------
# Tariff notes save — saves the notes field of the active tariff.
# Guardado de notas de tarifa — guarda el campo notes de la tarifa activa.
# ---------------------------------------------------------------------------


class TariffSaveNotesView(AdminRoleRequiredMixin, View):
    """
    Saves the notes field of an InsurerTariff. Redirects back to the
    insurer edit view.
    ---
    Guarda el campo notes de un InsurerTariff. Redirige de vuelta a la
    vista de edicion de la aseguradora.
    """

    def post(self, request, pk):
        """
        Save tariff notes and redirect to insurer edit view.
        ---
        Guarda las notas de la tarifa y redirige a la vista de edicion.
        """
        from budgets.models import InsurerTariff
        company_user = _get_company_user(request)
        tariff = get_object_or_404(
            InsurerTariff,
            pk=pk,
            insurer__company=company_user.company,
        )
        tariff.notes = request.POST.get("notes", "").strip()
        tariff.save(update_fields=["notes"])
        messages.success(request, "Notas de la tarifa guardadas.")
        return redirect("budgets:insurer_update", pk=tariff.insurer.pk)


# ---------------------------------------------------------------------------
# Budget bulk delete — ADMIN only
# Eliminacion masiva de presupuestos — solo ADMIN
# ---------------------------------------------------------------------------

class BudgetBulkDeleteView(AdminRoleRequiredMixin, View):
    """
    Deletes multiple budgets selected from the history list.
    Only DRAFT and REJECTED budgets can be deleted.
    ACCEPTED budgets are silently skipped even if their pk is submitted.
    Redirects back to the history view with a summary message.
    ---
    Elimina multiples presupuestos seleccionados desde el listado de historial.
    Solo se pueden eliminar presupuestos en estado DRAFT o REJECTED.
    Los presupuestos ACCEPTED se omiten silenciosamente aunque su pk sea enviado.
    Redirige de vuelta al historial con un mensaje resumen.
    """

    def post(self, request):
        """
        Process the bulk delete POST request.
        ---
        Procesa la peticion POST de eliminacion masiva.
        """
        company_user = _get_company_user(request)
        raw_ids = request.POST.getlist("budget_ids")

        # Parse and validate submitted pk values.
        # Parsear y validar los valores pk enviados.
        try:
            pk_list = [int(pk) for pk in raw_ids if pk.strip().isdigit()]
        except (ValueError, AttributeError):
            pk_list = []

        if not pk_list:
            messages.warning(request, "No se ha seleccionado ningún presupuesto.")
            return redirect("budgets:history")

        # Filter to company scope and exclude ACCEPTED budgets.
        # Filtrar al ambito de la empresa y excluir presupuestos ACCEPTED.
        qs = Budget.objects.filter(
            pk__in=pk_list,
            company=company_user.company,
        ).exclude(status=Budget.STATUS_ACCEPTED)

        count = qs.count()
        qs.delete()

        if count == 0:
            messages.warning(
                request,
                "Ninguno de los presupuestos seleccionados puede eliminarse. "
                "Los presupuestos aceptados no se pueden borrar.",
            )
        elif count == 1:
            messages.success(request, "1 presupuesto eliminado correctamente.")
        else:
            messages.success(
                request,
                f"{count} presupuestos eliminados correctamente.",
            )

        return redirect("budgets:history")


# ---------------------------------------------------------------------------
# Insurer clone — clones an insurer with all its tariff data. ADMIN only.
# Clonado de aseguradora — clona una aseguradora con toda su tarifa. Solo ADMIN.
# ---------------------------------------------------------------------------


class InsurerCloneView(AdminRoleRequiredMixin, View):
    """
    Clones an existing Insurer together with its active InsurerTariff,
    all TariffLine records, VehicleType catalogue, SpecialRateTariff,
    SpecialRateLines and InsurerBase links.
    The operator provides insurer_company_name and service_company_name
    via a POST form. The combined name and code are derived automatically.
    All other fields are copied verbatim from the source insurer.
    The cloned insurer is created with is_active=True and the cloned
    tariff starts on today's date with valid_to=None.
    Redirects to the edit view of the new insurer on success.
    ---
    Clona una Insurer existente junto con su InsurerTariff activa,
    todos los registros TariffLine, el catálogo VehicleType,
    SpecialRateTariff, SpecialRateLine y los vínculos InsurerBase.
    El operario proporciona insurer_company_name y service_company_name
    mediante un formulario POST. El nombre combinado y el código se
    derivan automáticamente.
    El resto de campos se copian literalmente de la aseguradora origen.
    La aseguradora clonada se crea con is_active=True y la tarifa clonada
    arranca en la fecha de hoy con valid_to=None.
    Redirige a la vista de edición de la nueva aseguradora en éxito.
    """

    @staticmethod
    def _derive_code(insurer_company_name, service_company_name):
        """
        Derives a code from insurer_company_name and service_company_name
        following the project convention: INSURERNAME_SERVICENAME.
        Strips spaces, hyphens, dots and parentheses; uppercases the result.
        Example: 'Europ Assistance' + 'Transgrual' -> 'EUROP_TRANSGRUAL'
        ---
        Deriva el código a partir de insurer_company_name y service_company_name
        siguiendo la convención del proyecto: ASEGURADORA_PRESTADORA.
        Elimina espacios, guiones, puntos y paréntesis; convierte a mayúsculas.
        Ejemplo: 'Europ Assistance' + 'Transgrual' -> 'EUROP_TRANSGRUAL'
        """
        import re as _re
        def _normalize(s):
            s = s.upper()
            s = _re.sub(r"[\s\-\.\(\)\/]+", "_", s)
            s = _re.sub(r"_+", "_", s)
            return s.strip("_")

        insurer_part = _normalize(insurer_company_name)[:25]
        service_part = _normalize(service_company_name)[:20] if service_company_name else ""
        if service_part:
            return f"{insurer_part}_{service_part}"
        return insurer_part

    def post(self, request, pk):
        """
        Clone the insurer identified by pk.
        Requires 'insurer_company_name' and optionally 'service_company_name'
        in POST data.
        ---
        Clona la aseguradora identificada por pk.
        Requiere 'insurer_company_name' y opcionalmente 'service_company_name'
        en los datos POST.
        """
        import datetime as _dt
        from budgets.models import (
            InsurerTariff, TariffLine, VehicleType, InsurerBase,
            SpecialRateTariff, SpecialRateLine,
        )

        company_user = _get_company_user(request)
        source = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )

        new_insurer_name = request.POST.get("insurer_company_name", "").strip()
        new_service_name = request.POST.get("service_company_name", "").strip()

        if not new_insurer_name:
            messages.error(
                request,
                "El nombre de la aseguradora es obligatorio.",
            )
            return redirect("budgets:insurer_detail", pk=pk)

        # Build combined name following project convention.
        # Construir nombre combinado siguiendo la convención del proyecto.
        if new_service_name:
            new_name = f"{new_insurer_name} / {new_service_name}"
        else:
            new_name = new_insurer_name

        # Guard: combined name must be unique within the company.
        # Guardia: el nombre combinado debe ser único dentro de la empresa.
        if Insurer.objects.filter(
            company=company_user.company,
            name=new_name,
        ).exists():
            messages.error(
                request,
                f"Ya existe una aseguradora con el nombre '{new_name}' en esta empresa.",
            )
            return redirect("budgets:insurer_detail", pk=pk)

        # Derive code from insurer_company_name + service_company_name.
        # If collision, append numeric suffix until unique.
        # Derivar código de insurer_company_name + service_company_name.
        # Si hay colisión, añadir sufijo numérico hasta encontrar uno libre.
        new_code_base = self._derive_code(new_insurer_name, new_service_name)[:45]
        new_code = new_code_base
        suffix = 1
        while Insurer.objects.filter(
            company=company_user.company,
            code=new_code,
        ).exists():
            new_code = f"{new_code_base}_{suffix}"
            suffix += 1

        with transaction.atomic():
            # --- Clone the Insurer record ---
            # --- Clonar el registro Insurer ---
            clone = Insurer.objects.create(
                company=source.company,
                name=new_name,
                insurer_company_name=new_insurer_name,
                service_company_name=new_service_name,
                code=new_code,
                management_fee_percent=source.management_fee_percent,
                surcharges_are_cumulative=source.surcharges_are_cumulative,
                is_active=True,
                is_insurance_company=source.is_insurance_company,
                always_apply_iva=source.always_apply_iva,
                special_night_holiday_tariff=source.special_night_holiday_tariff,
                local_service_km_threshold=source.local_service_km_threshold,
                notes=source.notes,
            )

            # --- Clone VehicleType catalogue ---
            # Map source_vt_pk -> clone_vt for use when cloning TariffLine
            # and SpecialRateLine FKs.
            # --- Clonar catálogo VehicleType ---
            # Mapa source_vt_pk -> clone_vt para las FK de TariffLine
            # y SpecialRateLine.
            vt_map = {}
            for vt in source.vehicle_types.all().order_by("sort_order", "name"):
                clone_vt = VehicleType.objects.create(
                    insurer=clone,
                    name=vt.name,
                    sort_order=vt.sort_order,
                    is_active=vt.is_active,
                )
                vt_map[vt.pk] = clone_vt

            # --- Clone active InsurerTariff + TariffLines ---
            # --- Clonar InsurerTariff activa + TariffLines ---
            source_tariff = InsurerTariff.objects.filter(
                insurer=source,
                valid_to__isnull=True,
            ).prefetch_related("lines").first()

            clone_tariff = None
            if source_tariff:
                clone_tariff = InsurerTariff.objects.create(
                    insurer=clone,
                    year=source_tariff.year,
                    valid_from=_dt.date.today(),
                    valid_to=None,
                    notes=source_tariff.notes,
                )
                for line in source_tariff.lines.all():
                    TariffLine.objects.create(
                        tariff=clone_tariff,
                        vehicle_type=(
                            vt_map.get(line.vehicle_type_id)
                            if line.vehicle_type_id else None
                        ),
                        concept=line.concept,
                        unit=line.unit,
                        price=line.price,
                        km_threshold=line.km_threshold,
                        min_units=line.min_units,
                        requires_authorization=line.requires_authorization,
                    )

                # --- Clone SpecialRateTariff + SpecialRateLines (if present) ---
                # SpecialRateTariff is a OneToOneField on InsurerTariff
                # (related_name='special_rate'). Only exists when
                # insurer.special_night_holiday_tariff is True.
                # --- Clonar SpecialRateTariff + SpecialRateLines (si existe) ---
                # SpecialRateTariff es un OneToOneField sobre InsurerTariff
                # (related_name='special_rate'). Solo existe cuando
                # insurer.special_night_holiday_tariff es True.
                try:
                    source_srt = source_tariff.special_rate
                    clone_srt = SpecialRateTariff.objects.create(
                        insurer_tariff=clone_tariff,
                        notes=source_srt.notes,
                    )
                    for srl in source_srt.lines.all():
                        SpecialRateLine.objects.create(
                            special_rate_tariff=clone_srt,
                            vehicle_type=(
                                vt_map.get(srl.vehicle_type_id)
                                if srl.vehicle_type_id else None
                            ),
                            concept=srl.concept,
                            unit=srl.unit,
                            price=srl.price,
                            km_threshold=srl.km_threshold,
                            min_units=srl.min_units,
                        )
                except SpecialRateTariff.DoesNotExist:
                    pass

            # --- Clone InsurerBase links ---
            # --- Clonar vínculos InsurerBase ---
            for ib in InsurerBase.objects.filter(insurer=source):
                InsurerBase.objects.create(
                    insurer=clone,
                    base=ib.base,
                    is_active=ib.is_active,
                )

        messages.success(
            request,
            f"Aseguradora '{new_name}' creada correctamente como clon de '{source.name}'.",
        )
        return redirect("budgets:insurer_update", pk=clone.pk)


# ---------------------------------------------------------------------------
# Insurer copy tariff — POST, copia la tarifa activa de una aseguradora
# origen a una aseguradora destino existente. Solo ADMIN.
# Insurer copy tariff — POST, copies the active tariff from a source insurer
# to an existing target insurer. ADMIN only.
# ---------------------------------------------------------------------------

class InsurerCopyTariffView(AdminRoleRequiredMixin, View):
    """
    Copies the active InsurerTariff (with all TariffLine, VehicleType,
    SpecialRateTariff and SpecialRateLine records) from a source insurer
    to an existing target insurer. The target insurer's current active tariff
    is closed (valid_to = today - 1 day) before the new one is created.
    The target insurer's own VehicleType catalogue is merged with the
    source's catalogue by exact (trimmed) name: matching names are reused
    in place (sort_order/is_active updated); target names absent from the
    source are deactivated, never deleted. No VehicleType row is ever
    deleted by this view, so no protected relation (budgets,
    special_rate_lines, work_orders_assistance) can ever block the copy.
    Existing Budget records for the target insurer are not touched.
    Redirects to the target insurer detail view on success.
    ---
    Copia la InsurerTariff activa (con todos los registros TariffLine,
    VehicleType, SpecialRateTariff y SpecialRateLine) de una aseguradora origen
    a una aseguradora destino existente. La tarifa activa actual de la destino
    se cierra (valid_to = hoy - 1 dia) antes de crear la nueva.
    El catalogo VehicleType de la destino se fusiona con el de la origen por
    nombre exacto (recortado): los nombres coincidentes se reutilizan en
    sitio (se actualiza sort_order/is_active); los nombres de la destino
    ausentes en la origen se desactivan, nunca se borran. Esta vista jamas
    borra ninguna fila VehicleType, por lo que ninguna relacion protegida
    (budgets, special_rate_lines, work_orders_assistance) puede bloquear
    la copia.
    Los registros Budget de la aseguradora destino no se modifican.
    Redirige a la vista de detalle de la aseguradora destino en exito.
    """

    def post(self, request, pk):
        """
        Copy the active tariff of insurer pk to the target insurer.
        Requires 'target_insurer_pk' in POST data.
        ---
        Copia la tarifa activa de la aseguradora pk a la aseguradora destino.
        Requiere 'target_insurer_pk' en los datos POST.
        """
        import datetime as _dt
        from budgets.models import (
            InsurerTariff, TariffLine, VehicleType,
            SpecialRateTariff, SpecialRateLine, TariffConcept,
        )

        company_user = _get_company_user(request)
        source = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )

        target_pk = request.POST.get("target_insurer_pk", "").strip()
        if not target_pk:
            messages.error(request, "Debes seleccionar una aseguradora destino.")
            return redirect("budgets:insurer_detail", pk=pk)

        try:
            target_pk = int(target_pk)
        except (ValueError, TypeError):
            messages.error(request, "Aseguradora destino no válida.")
            return redirect("budgets:insurer_detail", pk=pk)

        if target_pk == pk:
            messages.error(
                request,
                "La aseguradora destino no puede ser la misma que la origen.",
            )
            return redirect("budgets:insurer_detail", pk=pk)

        target = get_object_or_404(
            Insurer,
            pk=target_pk,
            company=company_user.company,
        )

        # Locate the active source tariff.
        # Localizar la tarifa activa de la aseguradora origen.
        source_tariff = (
            InsurerTariff.objects
            .filter(insurer=source, valid_to__isnull=True)
            .order_by("-valid_from")
            .first()
        )
        if source_tariff is None:
            messages.error(
                request,
                f"La aseguradora '{source.name}' no tiene tarifa activa para copiar.",
            )
            return redirect("budgets:insurer_detail", pk=pk)

        today = _dt.date.today()

        with transaction.atomic():
            # Close the target's current active tariff, if any.
            # Cerrar la tarifa activa actual de la destino, si existe.
            InsurerTariff.objects.filter(
                insurer=target,
                valid_to__isnull=True,
            ).update(valid_to=today - _dt.timedelta(days=1))

            # Merge VehicleType catalogue for target from source, matched
            # by exact (trimmed) name. Matching names are reused in place
            # (sort_order/is_active updated) — no row is ever deleted, so
            # no protected relation (budgets, special_rate_lines,
            # work_orders_assistance) can ever block this operation.
            # Target VehicleType names absent from the source (e.g. a
            # renamed concept) are deactivated, never deleted, and
            # reported in the success message for manual review.
            # ---
            # Fusiona el catalogo VehicleType de la destino con la origen,
            # emparejando por nombre exacto (recortado). Los nombres
            # coincidentes se reutilizan en sitio (se actualiza
            # sort_order/is_active) — nunca se borra ninguna fila, por lo
            # que ninguna relacion protegida (budgets, special_rate_lines,
            # work_orders_assistance) puede bloquear esta operacion.
            # Los VehicleType de la destino cuyo nombre no aparece en la
            # origen (p.ej. un concepto renombrado) se desactivan, nunca
            # se borran, y se informan en el mensaje de exito para
            # revision manual.
            source_types = list(
                source.vehicle_types.order_by("sort_order", "name")
            )
            source_names = {vt.name.strip() for vt in source_types}

            target_types_by_name = {
                vt.name.strip(): vt
                for vt in VehicleType.objects.filter(insurer=target)
            }

            # Build VehicleType map: source_vt_pk -> target VehicleType
            # (reused if the name matches, created new otherwise).
            # Construir mapa VehicleType: source_vt_pk -> VehicleType
            # destino (reutilizado si el nombre coincide, creado nuevo
            # en caso contrario).
            vt_map = {}
            for vt in source_types:
                key = vt.name.strip()
                existing = target_types_by_name.get(key)
                if existing is not None:
                    existing.sort_order = vt.sort_order
                    existing.is_active = vt.is_active
                    existing.save(update_fields=["sort_order", "is_active"])
                    vt_map[vt.pk] = existing
                else:
                    new_vt = VehicleType.objects.create(
                        insurer=target,
                        name=vt.name,
                        sort_order=vt.sort_order,
                        is_active=vt.is_active,
                    )
                    vt_map[vt.pk] = new_vt

            # Deactivate target VehicleType names not present in the
            # source — never delete.
            # Desactivar los nombres VehicleType de la destino ausentes
            # en la origen — nunca borrar.
            deactivated_names = []
            for name, vt in target_types_by_name.items():
                if name not in source_names and vt.is_active:
                    vt.is_active = False
                    vt.save(update_fields=["is_active"])
                    deactivated_names.append(vt.name)

            # Clone the active tariff to target.
            # Clonar la tarifa activa a la destino.
            new_tariff = InsurerTariff.objects.create(
                insurer=target,
                year=source_tariff.year,
                valid_from=today,
                valid_to=None,
                notes=source_tariff.notes,
            )

            # Clone TariffLine records mapping VehicleType FKs.
            # Clonar los registros TariffLine mapeando las FK de VehicleType.
            for line in source_tariff.lines.select_related(
                "vehicle_type", "concept"
            ).all():
                new_vt = vt_map.get(line.vehicle_type_id) if line.vehicle_type_id else None
                TariffLine.objects.create(
                    tariff=new_tariff,
                    vehicle_type=new_vt,
                    concept=line.concept,
                    price=line.price,
                )

            # Clone SpecialRateTariff + SpecialRateLine if present.
            # SpecialRateTariff is OneToOne to InsurerTariff via related_name
            # "special_rate". Access via source_tariff.special_rate — raises
            # SpecialRateTariff.DoesNotExist if none exists for this tariff.
            # SpecialRateLine FK to SpecialRateTariff via related_name "lines".
            # Clonar SpecialRateTariff + SpecialRateLine si existe.
            # SpecialRateTariff es OneToOne a InsurerTariff via related_name
            # "special_rate". Acceso via source_tariff.special_rate — lanza
            # SpecialRateTariff.DoesNotExist si no existe para esta tarifa.
            # SpecialRateLine FK a SpecialRateTariff via related_name "lines".
            try:
                source_special = source_tariff.special_rate
            except SpecialRateTariff.DoesNotExist:
                source_special = None

            if source_special is not None:
                new_special = SpecialRateTariff.objects.create(
                    insurer_tariff=new_tariff,
                    notes=source_special.notes,
                )
                for sline in source_special.lines.select_related(
                    "vehicle_type", "concept"
                ).all():
                    new_vt = vt_map.get(sline.vehicle_type_id) if sline.vehicle_type_id else None
                    SpecialRateLine.objects.create(
                        special_rate_tariff=new_special,
                        vehicle_type=new_vt,
                        concept=sline.concept,
                        unit=sline.unit,
                        price=sline.price,
                    )

        if deactivated_names:
            names_str = ", ".join(sorted(deactivated_names))
            messages.success(
                request,
                f"Tarifa de '{source.name}' copiada correctamente a "
                f"'{target.name}'. Se han desactivado "
                f"{len(deactivated_names)} tipo(s) de vehículo no "
                f"presentes en la tarifa origen: {names_str}."
            )
        else:
            messages.success(
                request,
                f"Tarifa de '{source.name}' copiada correctamente a "
                f"'{target.name}'.",
            )
        return redirect("budgets:insurer_detail", pk=target.pk)


# ---------------------------------------------------------------------------
# Insurer detail — read-only tariff view. ADMIN only.
# Vista de detalle de aseguradora — solo lectura. Solo ADMIN.
# ---------------------------------------------------------------------------

class InsurerDetailView(AdminRoleRequiredMixin, View):
    """
    Read-only view showing the full tariff detail for an insurer.
    Displays general data, vehicle types and all tariff lines grouped
    by vehicle type. No editing allowed from this view.
    ---
    Vista de solo lectura que muestra el detalle completo de tarifa
    de una aseguradora. Muestra datos generales, tipos de vehiculo y
    todas las lineas de tarifa agrupadas por tipo de vehiculo.
    No se permite edicion desde esta vista.
    """

    template_name = "budgets/insurer_detail.html"

    def get(self, request, pk):
        """
        Render the insurer detail view for the given pk.
        ---
        Renderiza la vista de detalle de aseguradora para el pk dado.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer.objects.select_related('night_schedule'),
            pk=pk,
            company=company_user.company,
        )
        # Resolve active tariff.
        # Resolver tarifa activa.
        from budgets.models import InsurerTariff
        tariff = InsurerTariff.objects.filter(
            insurer=insurer,
            valid_to__isnull=True,
        ).prefetch_related(
            "lines__vehicle_type",
        ).first()

        # Build vehicle_type -> lines mapping for clean template rendering.
        # Construir mapa tipo_vehiculo -> lineas para renderizado limpio en template.
        vehicle_types = []
        generic_lines = []
        if tariff:
            vt_map = {}
            for line in tariff.lines.order_by(
                "vehicle_type__sort_order", "concept"
            ):
                if line.vehicle_type is None:
                    generic_lines.append(line)
                else:
                    vt_key = line.vehicle_type.pk
                    if vt_key not in vt_map:
                        vt_map[vt_key] = {
                            "vehicle_type": line.vehicle_type,
                            "lines": [],
                        }
                    vt_map[vt_key]["lines"].append(line)
            vehicle_types = list(vt_map.values())

        # Resolve SpecialRateTariff if the insurer has one.
        # Build the same vt->lines structure for the special rate table.
        # Resolver SpecialRateTariff si la aseguradora tiene una.
        # Construir la misma estructura vt->lineas para la tabla especial.
        from budgets.models import SpecialRateTariff
        special_rate = None
        special_vehicle_types = []
        special_generic_lines = []
        if tariff and insurer.special_night_holiday_tariff:
            try:
                special_rate = tariff.special_rate
                srt_vt_map = {}
                for line in special_rate.lines.order_by(
                    "vehicle_type__sort_order", "concept"
                ):
                    if line.vehicle_type is None:
                        special_generic_lines.append(line)
                    else:
                        vt_key = line.vehicle_type.pk
                        if vt_key not in srt_vt_map:
                            srt_vt_map[vt_key] = {
                                "vehicle_type": line.vehicle_type,
                                "lines": [],
                            }
                        srt_vt_map[vt_key]["lines"].append(line)
                special_vehicle_types = list(srt_vt_map.values())
            except SpecialRateTariff.DoesNotExist:
                pass

        insurer_bases = InsurerBase.objects.filter(
            insurer=insurer,
        ).select_related("base").order_by("base__name")
        bases = Base.objects.filter(
            company=company_user.company, is_active=True
        ).order_by("name")
        base_form = BaseForm()

        # Resolve night schedule display data for the template.
        # The resolution order follows the engine: insurer.night_schedule
        # -> company default NightSchedule -> not configured.
        # Resolver datos de horario nocturno para el template.
        # El orden sigue al motor: insurer.night_schedule
        # -> NightSchedule por defecto de empresa -> sin configurar.
        if insurer.night_schedule and insurer.night_schedule.is_active:
            # Insurer has an explicit active night schedule assigned.
            # La aseguradora tiene un horario nocturno activo asignado.
            night_schedule_display = {
                "source": "insurer",
                "name": insurer.night_schedule.name,
                "night_start": (
                    insurer.night_schedule.night_start
                    .strftime("%H:%M")
                ),
                "night_end": (
                    insurer.night_schedule.night_end
                    .strftime("%H:%M")
                ),
            }
        else:
            # Fall back to the company default NightSchedule.
            # Usar el NightSchedule por defecto de la empresa como fallback.
            default_ns = NightSchedule.objects.filter(
                company=company_user.company,
                is_default=True,
                is_active=True,
            ).first()
            if default_ns:
                night_schedule_display = {
                    "source": "company_default",
                    "name": default_ns.name,
                    "night_start": (
                        default_ns.night_start.strftime("%H:%M")
                    ),
                    "night_end": (
                        default_ns.night_end.strftime("%H:%M")
                    ),
                }
            else:
                # Neither insurer nor company has a configured schedule.
                # Ni la aseguradora ni la empresa tienen horario configurado.
                night_schedule_display = {"source": "none"}

        # Other insurers of the same company for the copy-tariff modal selector.
        # Excluye la aseguradora actual para no permitir copiar sobre sí misma.
        # Otras aseguradoras de la misma empresa para el selector del modal de
        # copia de tarifa. Excluye la aseguradora actual.
        other_insurers = (
            Insurer.objects
            .filter(company=company_user.company, is_active=True)
            .exclude(pk=insurer.pk)
            .order_by("name")
        )

        ctx = _build_base_context(request, {
            "insurer": insurer,
            "insurer_bases": insurer_bases,
            "insurer_vehicle_types": insurer.vehicle_types.order_by("sort_order", "name"),
            "tariff": tariff,
            "vehicle_types": vehicle_types,
            "generic_lines": generic_lines,
            "special_rate": special_rate,
            "special_vehicle_types": special_vehicle_types,
            "special_generic_lines": special_generic_lines,
            "bases": bases,
            "base_form": base_form,
            "night_schedule_display": night_schedule_display,
            "other_insurers": other_insurers,
            "active_nav": "insurer_list",
        })
        return render(request, self.template_name, ctx)

# ---------------------------------------------------------------------------
# Base management views — CRUD for service bases linked to insurers
# Vistas de gestion de bases — CRUD para bases de servicio vinculadas a aseguradoras
# ---------------------------------------------------------------------------


class BaseForm(django_forms.ModelForm):
    """
    ModelForm for creating and editing Base records.
    ---
    ModelForm para crear y editar registros Base.
    """

    class Meta:
        model = Base
        fields = [
            "name",
            "municipality",
            "latitude",
            "longitude",
            "labor_calendar",
            "is_active",
        ]
        widgets = {
            "name": django_forms.TextInput(attrs={"class": "form-control form-control-sm"}),
            "municipality": django_forms.TextInput(attrs={"class": "form-control form-control-sm"}),
            "latitude": django_forms.NumberInput(attrs={"class": "form-control form-control-sm", "step": "0.000001"}),
            "longitude": django_forms.NumberInput(attrs={"class": "form-control form-control-sm", "step": "0.000001"}),
            "labor_calendar": django_forms.Textarea(attrs={"class": "form-control form-control-sm", "rows": 4, "readonly": True}),
            "is_active": django_forms.CheckboxInput(attrs={"class": "form-check-input"}),
        }
        labels = {
            "name": "Nombre de la base",
            "municipality": "Municipio",
            "latitude": "Latitud",
            "longitude": "Longitud",
            "labor_calendar": "Calendario laboral (JSON — gestionado automaticamente)",
            "is_active": "Activa",
        }


class BaseManageView(AdminRoleRequiredMixin, View):
    """
    Dedicated base management view for a given insurer.
    Shows a read-only summary of active bases for the insurer at the top,
    and a full list of all company bases with per-insurer InsurerBase toggle
    at the bottom. Implements the design agreed in H18 S001 (annex section 2.7).
    ADMIN only.
    ---
    Vista dedicada de gestion de bases para una aseguradora dada.
    Muestra un resumen de solo lectura de las bases activas de la aseguradora
    arriba, y el listado completo de bases de la empresa con toggle InsurerBase
    por aseguradora abajo. Implementa el diseno acordado en H18 S001 (anexo 2.7).
    Solo ADMIN.
    """

    template_name = "budgets/bases.html"

    def get(self, request, pk):
        """
        Render the base management page for the given insurer pk.
        Builds two querysets:
        - active_insurer_bases: InsurerBase records active for this insurer.
        - all_company_bases: all company bases annotated with has_active_ib flag.
        ---
        Renderiza la pagina de gestion de bases para el pk de aseguradora dado.
        Construye dos querysets:
        - active_insurer_bases: registros InsurerBase activos para esta aseguradora.
        - all_company_bases: todas las bases de la empresa con flag has_active_ib.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )
        # Top section: active bases for this insurer (read-only summary).
        # Seccion superior: bases activas de esta aseguradora (resumen solo lectura).
        active_insurer_bases = (
            InsurerBase.objects.filter(
                insurer=insurer,
                is_active=True,
                base__is_active=True,
            )
            .select_related("base")
            .order_by("base__name")
        )
        # Bottom section: all company bases, annotated with per-insurer active flag.
        # Build a dict {base_pk: InsurerBase} for efficient annotation.
        # Seccion inferior: todas las bases de la empresa, anotadas con flag activa
        # por aseguradora. Construir dict {base_pk: InsurerBase} para anotacion eficiente.
        ib_map = {
            ib.base_id: ib
            for ib in InsurerBase.objects.filter(insurer=insurer).select_related("base")
        }
        all_company_bases_qs = (
            Base.objects.filter(
                company=company_user.company,
                is_active=True,
            )
            .order_by("name")
        )
        # Annotate each base with its InsurerBase record (or None) for this insurer.
        # Anotar cada base con su registro InsurerBase (o None) para esta aseguradora.
        all_company_bases = []
        for base in all_company_bases_qs:
            ib = ib_map.get(base.pk)
            base.insurer_base = ib
            base.has_active_ib = bool(ib and ib.is_active)
            all_company_bases.append(base)
        base_form = BaseForm()
        ctx = _build_base_context(request, {
            "insurer": insurer,
            "active_insurer_bases": active_insurer_bases,
            "all_company_bases": all_company_bases,
            "base_form": base_form,
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
            "active_nav": "budgets_insurers",
        })
        return render(request, self.template_name, ctx)


class InsurerBaseToggleView(AdminRoleRequiredMixin, View):
    """
    HTMX endpoint. Toggles InsurerBase.is_active for a (insurer_pk, base_pk) pair.
    Creates the InsurerBase record if it does not exist yet (first activation).
    Returns the updated toggle button fragment for HTMX outerHTML swap.
    ---
    Endpoint HTMX. Alterna InsurerBase.is_active para el par (insurer_pk, base_pk).
    Crea el registro InsurerBase si no existe (primera activacion).
    Devuelve el fragmento del boton toggle actualizado para swap HTMX outerHTML.
    """

    def post(self, request, insurer_pk, base_pk):
        """
        Toggle InsurerBase.is_active and return the updated toggle fragment.
        ---
        Alterna InsurerBase.is_active y devuelve el fragmento toggle actualizado.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer,
            pk=insurer_pk,
            company=company_user.company,
        )
        base = get_object_or_404(
            Base,
            pk=base_pk,
            company=company_user.company,
        )
        # Get or create the InsurerBase relation; toggle is_active.
        # Obtener o crear la relacion InsurerBase; alternar is_active.
        ib, created = InsurerBase.objects.get_or_create(
            insurer=insurer,
            base=base,
            defaults={"is_active": True},
        )
        if not created:
            ib.is_active = not ib.is_active
            ib.save(update_fields=["is_active", "updated_at"])
        # Annotate base with the updated insurer_base for template rendering.
        # Anotar base con insurer_base actualizado para el renderizado del template.
        base.insurer_base = ib
        base.has_active_ib = ib.is_active
        return render(
            request,
            "budgets/partials/insurerbase_toggle_fragment.html",
            {"base": base, "insurer": insurer},
        )


class BaseCreateView(AdminRoleRequiredMixin, View):
    """
    Creates a new Base for the given insurer.
    Accepts POST only. Returns HTMX fragment with updated base list on success.
    ---
    Crea una nueva Base para la aseguradora dada.
    Solo acepta POST. Devuelve fragmento HTMX con la lista actualizada en exito.
    """

    def post(self, request, pk):
        """
        Validate form and create Base. Returns HTMX fragment or error response.
        ---
        Valida el formulario y crea la Base. Devuelve fragmento HTMX o respuesta de error.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(Insurer, pk=pk, company=company_user.company)
        form = BaseForm(request.POST)
        if form.is_valid():
            base = form.save(commit=False)
            # Assign company from the insurer — Base is no longer insurer-scoped.
            # Asignar company desde la aseguradora — Base ya no tiene ambito de aseguradora.
            base.company = insurer.company
            base.save()
            # Create the InsurerBase relation linking base to this insurer.
            # Crear la relacion InsurerBase vinculando la base a esta aseguradora.
            InsurerBase.objects.get_or_create(
                insurer=insurer,
                base=base,
                defaults={"is_active": True},
            )
            # Return the new row fragment for HTMX beforeend swap.
            # Devolver el fragmento de fila nueva para swap HTMX beforeend.
            return render(
                request,
                "budgets/partials/base_row_fragment.html",
                {"base": base},
            )
        company_bases = Base.objects.filter(
            company=insurer.company, is_active=True
        ).order_by("name")
        return render(
            request,
            "budgets/partials/base_list_fragment.html",
            {"insurer": insurer, "bases": company_bases, "base_form": form, "form_errors": True},
        )


class BaseUpdateView(AdminRoleRequiredMixin, View):
    """
    Updates an existing Base record.
    GET renders a dedicated edit page with Google Maps geolocation.
    POST saves the form and redirects back to the global base list.
    ---
    Actualiza un registro Base existente.
    GET renderiza una pagina de edicion dedicada con geolocalizacion Google Maps.
    POST guarda el formulario y redirige al listado global de bases.
    """

    template_name = "budgets/base_edit_page.html"

    def get(self, request, pk):
        """
        Render the dedicated edit page for the given base pk.
        ---
        Renderiza la pagina de edicion dedicada para el pk de base dado.
        """
        company_user = _get_company_user(request)
        base = get_object_or_404(Base, pk=pk, company=company_user.company)
        form = BaseForm(instance=base)
        ctx = _build_base_context(request, {
            "base": base,
            "base_form": form,
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
            "active_nav": "budgets_bases",
        })
        return render(request, self.template_name, ctx)

    def post(self, request, pk):
        """
        Save the updated base and redirect to the global base list.
        ---
        Guarda la base actualizada y redirige al listado global de bases.
        """
        company_user = _get_company_user(request)
        base = get_object_or_404(Base, pk=pk, company=company_user.company)
        form = BaseForm(request.POST, instance=base)
        if form.is_valid():
            form.save()
            from django.contrib import messages as _msg
            _msg.success(
                request,
                f"Base '{base.name}' actualizada correctamente."
            )
            return redirect("budgets:base_global")
        ctx = _build_base_context(request, {
            "base": base,
            "base_form": form,
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
            "active_nav": "budgets_bases",
        })
        return render(request, self.template_name, ctx)


class BaseToggleView(AdminRoleRequiredMixin, View):
    """
    Toggles the is_active flag of a Base record via HTMX POST.
    Returns the updated row fragment.
    ---
    Alterna el flag is_active de un registro Base via HTMX POST.
    Devuelve el fragmento de fila actualizado.
    """

    def post(self, request, pk):
        """
        Toggle is_active and return the updated row fragment.
        ---
        Alterna is_active y devuelve el fragmento de fila actualizado.
        """
        company_user = _get_company_user(request)
        base = get_object_or_404(Base, pk=pk, company=company_user.company)
        # Toggle global Base.is_active flag.
        # Alternar flag global Base.is_active.
        base.is_active = not base.is_active
        base.save(update_fields=["is_active", "updated_at"])
        return render(
            request,
            "budgets/partials/base_row_fragment.html",
            {"base": base},
        )


class BaseDeleteView(AdminRoleRequiredMixin, View):
    """
    Deletes a Base record via POST with modal confirmation.
    Returns empty 200 response for HTMX swap-oob removal.
    ---
    Elimina un registro Base via POST con confirmacion modal.
    Devuelve respuesta 200 vacia para eliminacion HTMX swap-oob.
    """

    def post(self, request, pk):
        """
        Delete the base and return empty response.
        ---
        Elimina la base y devuelve respuesta vacia.
        """
        company_user = _get_company_user(request)
        base = get_object_or_404(Base, pk=pk, company=company_user.company)
        # Guard: do not delete bases linked to existing budgets.
        # Guardia: no eliminar bases vinculadas a presupuestos existentes.
        if base.budgets.exists():
            return HttpResponseBadRequest(
                "No se puede eliminar una base con presupuestos asociados."
            )
        base.delete()
        from django.http import HttpResponse
        return HttpResponse(status=200)


# ---------------------------------------------------------------------------
# Vehicle type management — HTMX CRUD inline in InsurerDetailView
# Gestión de tipos de vehículo — CRUD HTMX inline en InsurerDetailView
# ---------------------------------------------------------------------------


class VehicleTypeCreateView(AdminRoleRequiredMixin, View):
    """
    Creates a new VehicleType for the given insurer via HTMX POST.
    sort_order is assigned automatically as max(existing) + 1 so the new
    entry always appears last. The user never inputs an order value.
    Returns the updated vehicle type list fragment on success.
    ---
    Crea un nuevo VehicleType para la aseguradora dada via HTMX POST.
    sort_order se asigna automaticamente como max(existente) + 1, de modo
    que el nuevo entrada siempre aparece al final. El usuario nunca introduce
    un valor de orden.
    Devuelve el fragmento de lista actualizado en caso de exito.
    """

    def post(self, request, pk):
        """
        Validate and create VehicleType. sort_order assigned automatically.
        If 'globalize' flag is present in POST, creates the VehicleType
        for every active Insurer in the company (propagation mode).
        ---
        Valida y crea VehicleType. sort_order asignado automaticamente.
        Si el flag 'globalize' esta en el POST, crea el VehicleType
        para cada Insurer activo de la empresa (modo propagacion).
        """
        from django.http import HttpResponseBadRequest
        from django.db.models import Max
        company_user = _get_company_user(request)
        insurer = get_object_or_404(
            Insurer, pk=pk, company=company_user.company,
        )
        name = request.POST.get("name", "").strip()
        globalize = request.POST.get("globalize") == "1"
        if not name:
            return HttpResponseBadRequest(
                "El nombre del tipo de vehículo es obligatorio."
            )
        if globalize:
            # Propagate to every active insurer in the company.
            # Propagar a cada aseguradora activa de la empresa.
            insurers = Insurer.objects.filter(
                company=company_user.company,
                is_active=True,
            )
            for target_insurer in insurers:
                # Assign sort_order as max + 1 per insurer.
                # Asignar sort_order como max + 1 por aseguradora.
                max_order = (
                    target_insurer.vehicle_types
                    .aggregate(m=Max("sort_order"))["m"]
                ) or 0
                VehicleType.objects.create(
                    insurer=target_insurer,
                    name=name,
                    sort_order=max_order + 1,
                    is_active=True,
                )
        else:
            max_order = (
                insurer.vehicle_types
                .aggregate(m=Max("sort_order"))["m"]
            ) or 0
            VehicleType.objects.create(
                insurer=insurer,
                name=name,
                sort_order=max_order + 1,
                is_active=True,
            )
        vehicle_types = insurer.vehicle_types.order_by("sort_order", "name")
        return render(
            request,
            "budgets/partials/vehicle_type_list_fragment.html",
            {"insurer": insurer, "vehicle_types": vehicle_types},
        )


class VehicleTypeUpdateView(AdminRoleRequiredMixin, View):
    """
    Updates an existing VehicleType name via HTMX POST.
    sort_order is managed exclusively via VehicleTypeReorderView (drag & drop).
    Returns the updated row fragment on success.
    ---
    Actualiza el nombre de un VehicleType existente via HTMX POST.
    sort_order se gestiona exclusivamente via VehicleTypeReorderView
    (drag & drop). Devuelve el fragmento de fila actualizado en exito.
    """

    def get(self, request, pk):
        """
        Return the inline edit form fragment for the given VehicleType pk.
        ---
        Devuelve el fragmento de formulario de edicion inline para el pk dado.
        """
        company_user = _get_company_user(request)
        vt = get_object_or_404(
            VehicleType,
            pk=pk,
            insurer__company=company_user.company,
        )
        return render(
            request,
            "budgets/partials/vehicle_type_edit_fragment.html",
            {"vt": vt},
        )

    def post(self, request, pk):
        """
        Save updated VehicleType name and return the updated row fragment.
        sort_order is intentionally not updated here.
        ---
        Guarda el nombre actualizado del VehicleType y devuelve la fila.
        sort_order no se actualiza aqui intencionalmente.
        """
        company_user = _get_company_user(request)
        vt = get_object_or_404(
            VehicleType,
            pk=pk,
            insurer__company=company_user.company,
        )
        name = request.POST.get("name", "").strip()
        if not name:
            from django.http import HttpResponseBadRequest
            return HttpResponseBadRequest(
                "El nombre del tipo de vehículo es obligatorio."
            )
        vt.name = name
        vt.save(update_fields=["name"])
        return render(
            request,
            "budgets/partials/vehicle_type_row_fragment.html",
            {"vt": vt},
        )


class VehicleTypeToggleView(AdminRoleRequiredMixin, View):
    """
    Toggles the is_active flag of a VehicleType via HTMX POST.
    Returns the updated row fragment.
    ---
    Alterna el flag is_active de un VehicleType via HTMX POST.
    Devuelve el fragmento de fila actualizado.
    """

    def post(self, request, pk):
        """
        Toggle is_active and return updated row fragment.
        ---
        Alterna is_active y devuelve el fragmento de fila actualizado.
        """
        company_user = _get_company_user(request)
        vt = get_object_or_404(
            VehicleType,
            pk=pk,
            insurer__company=company_user.company,
        )
        vt.is_active = not vt.is_active
        vt.save(update_fields=["is_active"])
        return render(
            request,
            "budgets/partials/vehicle_type_row_fragment.html",
            {"vt": vt},
        )


class VehicleTypeDeleteView(AdminRoleRequiredMixin, View):
    """
    Deactivates a VehicleType via HTMX POST with modal confirmation.
    "Delete" no longer performs a real database deletion — it sets
    is_active=False. The record, and every reference to it (TariffLine,
    Budget, SpecialRateLine, WorkOrderAssistance), remains intact; the
    vehicle type simply disappears from operative dropdowns and listings.
    Returns empty 200 response for HTMX row removal.
    ---
    Desactiva un VehicleType via HTMX POST con confirmación modal.
    "Eliminar" ya no realiza un borrado real en base de datos — establece
    is_active=False. El registro, y toda referencia a él (TariffLine,
    Budget, SpecialRateLine, WorkOrderAssistance), permanece intacto; el
    tipo de vehículo simplemente desaparece de los desplegables operativos
    y del listado.
    Devuelve respuesta 200 vacía para eliminación HTMX de la fila.
    """

    def post(self, request, pk):
        """
        Deactivate VehicleType. Always safe — no referential integrity
        check is needed, since no row is ever deleted.
        ---
        Desactiva VehicleType. Siempre seguro — no requiere ninguna
        comprobación de integridad referencial, ya que ninguna fila se
        borra jamás.
        """
        from django.http import HttpResponse
        company_user = _get_company_user(request)
        vt = get_object_or_404(
            VehicleType,
            pk=pk,
            insurer__company=company_user.company,
        )
        vt.is_active = False
        vt.save(update_fields=["is_active"])
        return HttpResponse(status=200)


class VehicleTypeReorderView(AdminRoleRequiredMixin, View):
    """
    Persists a new sort_order for all VehicleType records of an insurer
    after a drag-and-drop reorder operation in the UI.

    Accepts a JSON POST body: {"insurer_pk": N, "pks": [pk1, pk2, ...]}
    where pks is the ordered list of VehicleType primary keys reflecting
    the new visual order. Each VehicleType receives sort_order = its index
    in the list (0-based). All updates are executed in a single atomic
    transaction. Returns JSON {"status": "ok"} on success or HTTP 400/404
    on validation failure.
    ---
    Persiste un nuevo sort_order para todos los registros VehicleType de
    una aseguradora tras una operacion de reordenacion drag-and-drop en
    la UI.

    Acepta un cuerpo JSON POST: {"insurer_pk": N, "pks": [pk1, pk2, ...]}
    donde pks es la lista ordenada de claves primarias de VehicleType que
    refleja el nuevo orden visual. Cada VehicleType recibe sort_order igual
    a su indice en la lista (base 0). Todas las actualizaciones se ejecutan
    en una unica transaccion atomica. Devuelve JSON {"status": "ok"} en
    exito o HTTP 400/404 en caso de fallo de validacion.
    """

    def post(self, request):
        """
        Parse JSON body, validate ownership, persist sort_order atomically.
        ---
        Parsea el cuerpo JSON, valida la propiedad, persiste sort_order
        de forma atomica.
        """
        import json as _json
        from django.db import transaction
        from django.http import JsonResponse, HttpResponseBadRequest

        company_user = _get_company_user(request)

        try:
            body = _json.loads(request.body)
        except (ValueError, KeyError):
            return HttpResponseBadRequest("JSON inválido.")

        insurer_pk = body.get("insurer_pk")
        pks = body.get("pks", [])

        if not insurer_pk or not isinstance(pks, list) or not pks:
            return HttpResponseBadRequest(
                "Se requieren insurer_pk y pks (lista no vacía)."
            )

        insurer = get_object_or_404(
            Insurer, pk=insurer_pk, company=company_user.company,
        )

        # Validate that all pks belong to this insurer.
        # Validar que todos los pks pertenecen a esta aseguradora.
        qs = VehicleType.objects.filter(
            pk__in=pks, insurer=insurer,
        )
        if qs.count() != len(pks):
            return HttpResponseBadRequest(
                "Algún tipo de vehículo no pertenece a esta aseguradora."
            )

        # Persist new order atomically.
        # Persistir el nuevo orden de forma atomica.
        with transaction.atomic():
            for idx, vt_pk in enumerate(pks):
                VehicleType.objects.filter(pk=vt_pk).update(sort_order=idx)

        return JsonResponse({"status": "ok"})


class TariffConceptCreateView(AdminRoleRequiredMixin, View):
    """
    Creates a new custom TariffConcept for the company owning the given insurer.
    Accepts HTMX POST only. Returns empty 200 on success — the caller reloads.
    Rejects creation if a concept with the same label already exists for the company.
    ---
    Crea un nuevo TariffConcept personalizado para la empresa de la aseguradora dada.
    Solo acepta HTMX POST. Devuelve 200 vacío en éxito — el llamador recarga la página.
    Rechaza la creación si ya existe un concepto con el mismo nombre para la empresa.
    """

    def post(self, request, pk):
        """
        Validate and create TariffConcept. Returns 200 or 400.
        ---
        Valida y crea TariffConcept. Devuelve 200 o 400.
        """
        from django.http import HttpResponse
        import re
        company_user = _get_company_user(request)
        insurer = get_object_or_404(Insurer, pk=pk, company=company_user.company)
        company = insurer.company

        label = request.POST.get("label", "").strip()
        default_unit = request.POST.get("default_unit", TariffLine.UNIT_FIXED).strip()

        if not label:
            return HttpResponseBadRequest("El nombre del concepto es obligatorio.")

        # Generate a unique code from the label.
        # Generar un código único a partir del nombre.
        base_code = re.sub(r"[^A-Z0-9]", "_", label.upper())[:40]
        code = base_code
        counter = 1
        while TariffConcept.objects.filter(code=code).exists():
            code = f"{base_code}_{counter}"
            counter += 1

        # Reject if a concept with the same label already exists for this company.
        # Rechazar si ya existe un concepto con el mismo nombre para esta empresa.
        if TariffConcept.objects.filter(company=company, label__iexact=label).exists():
            return HttpResponseBadRequest(
                f"Ya existe un concepto con el nombre '{label}' para esta empresa."
            )

        TariffConcept.objects.create(
            code=code,
            label=label,
            default_unit=default_unit,
            is_system=False,
            company=company,
            sort_order=100,
        )
        return HttpResponse(status=200)


class TariffConceptCreateGlobalView(AdminRoleRequiredMixin, View):
    """
    Creates a new custom TariffConcept for the authenticated user's company.
    Used from the dedicated concept list page (no insurer_pk needed).
    Returns the new row fragment for HTMX beforeend swap.
    ---
    Crea un nuevo TariffConcept personalizado para la empresa del usuario autenticado.
    Usado desde la página dedicada de conceptos (sin insurer_pk).
    Devuelve el fragmento de fila nuevo para swap HTMX beforeend.
    """

    def post(self, request):
        """
        Validate and create TariffConcept. Returns row fragment or 400.
        ---
        Valida y crea TariffConcept. Devuelve fragmento de fila o 400.
        """
        import re
        company_user = _get_company_user(request)
        company = company_user.company

        label = request.POST.get("label", "").strip()
        default_unit = request.POST.get("default_unit", TariffLine.UNIT_FIXED).strip()
        sort_order_raw = request.POST.get("sort_order", "100").strip()

        if not label:
            return HttpResponseBadRequest("El nombre del concepto es obligatorio.")

        try:
            sort_order = int(sort_order_raw)
        except ValueError:
            sort_order = 100

        # Generate unique code.
        # Generar código único.
        base_code = re.sub(r"[^A-Z0-9]", "_", label.upper())[:40]
        code = base_code
        counter = 1
        while TariffConcept.objects.filter(code=code).exists():
            code = f"{base_code}_{counter}"
            counter += 1

        if TariffConcept.objects.filter(company=company, label__iexact=label).exists():
            return HttpResponseBadRequest(
                f"Ya existe un concepto con el nombre '{label}' para esta empresa."
            )

        concept = TariffConcept.objects.create(
            code=code,
            label=label,
            default_unit=default_unit,
            is_system=False,
            company=company,
            sort_order=sort_order,
        )
        return render(
            request,
            "budgets/partials/tariff_concept_row_fragment.html",
            {"concept": concept, "unit_choices": TariffLine.UNIT_CHOICES},
        )


class TariffConceptListView(AdminRoleRequiredMixin, View):
    """
    Dedicated page listing all TariffConcept available for the company:
    system concepts (read-only) and custom concepts (editable/deletable).
    ---
    Página dedicada que lista todos los TariffConcept disponibles para la empresa:
    conceptos de sistema (solo lectura) y personalizados (editables/eliminables).
    """

    template_name = "budgets/tariff_concept_list.html"

    def get(self, request):
        """
        Render the concept list page.
        ---
        Renderiza la página de listado de conceptos.
        """
        company_user = _get_company_user(request)
        from django.db.models import Q
        concepts = TariffConcept.objects.filter(
            Q(company__isnull=True) | Q(company=company_user.company)
        ).order_by("sort_order", "label")
        ctx = _build_base_context(request, {
            "concepts": concepts,
            "unit_choices": TariffLine.UNIT_CHOICES,
            "active_nav": "budgets_concepts",
        })
        return render(request, self.template_name, ctx)


class TariffConceptUpdateView(AdminRoleRequiredMixin, View):
    """
    Inline HTMX edit for a custom TariffConcept.
    GET returns the edit row fragment. POST saves and returns the updated row.
    System concepts (is_system=True) cannot be edited.
    ---
    Edición inline HTMX de un TariffConcept personalizado.
    GET devuelve el fragmento de fila de edición. POST guarda y devuelve la fila.
    Los conceptos de sistema (is_system=True) no pueden editarse.
    """

    def get(self, request, pk):
        """
        Return the inline edit form fragment.
        ---
        Devuelve el fragmento de formulario de edición inline.
        """
        company_user = _get_company_user(request)
        concept = get_object_or_404(
            TariffConcept,
            pk=pk,
            company=company_user.company,
            is_system=False,
        )
        return render(
            request,
            "budgets/partials/tariff_concept_edit_fragment.html",
            {"concept": concept, "unit_choices": TariffLine.UNIT_CHOICES},
        )

    def post(self, request, pk):
        """
        Save updated TariffConcept and return the updated row fragment.
        ---
        Guarda el TariffConcept actualizado y devuelve el fragmento de fila.
        """
        company_user = _get_company_user(request)
        concept = get_object_or_404(
            TariffConcept,
            pk=pk,
            company=company_user.company,
            is_system=False,
        )
        label = request.POST.get("label", "").strip()
        default_unit = request.POST.get("default_unit", concept.default_unit).strip()
        sort_order_raw = request.POST.get("sort_order", "0").strip()
        if not label:
            return HttpResponseBadRequest("El nombre del concepto es obligatorio.")
        try:
            sort_order = int(sort_order_raw)
        except ValueError:
            sort_order = concept.sort_order
        concept.label = label
        concept.default_unit = default_unit
        concept.sort_order = sort_order
        concept.save(update_fields=["label", "default_unit", "sort_order"])
        return render(
            request,
            "budgets/partials/tariff_concept_row_fragment.html",
            {"concept": concept, "unit_choices": TariffLine.UNIT_CHOICES},
        )


class TariffConceptDeleteView(AdminRoleRequiredMixin, View):
    """
    Deletes a custom TariffConcept via HTMX POST.
    Rejects deletion if the concept has associated tariff lines.
    System concepts cannot be deleted.
    ---
    Elimina un TariffConcept personalizado via HTMX POST.
    Rechaza la eliminación si el concepto tiene líneas de tarifa asociadas.
    Los conceptos de sistema no pueden eliminarse.
    """

    def post(self, request, pk):
        """
        Delete concept if safe. Return empty response or error.
        ---
        Elimina el concepto si es seguro. Devuelve respuesta vacía o error.
        """
        from django.http import HttpResponse
        company_user = _get_company_user(request)
        concept = get_object_or_404(
            TariffConcept,
            pk=pk,
            company=company_user.company,
            is_system=False,
        )
        if concept.tariff_lines.exists() or concept.special_rate_lines.exists():
            return HttpResponseBadRequest(
                "No se puede eliminar un concepto con líneas de tarifa asociadas. "
                "Elimina primero las líneas que usan este concepto."
            )
        concept.delete()
        return HttpResponse(status=200)


# ---------------------------------------------------------------------------
# Export views — Insurer tariff and budget history exports (CSV, Excel, PDF, Word)
# Vistas de exportacion — tarifas de aseguradora e historial de presupuestos
# ---------------------------------------------------------------------------

from django.http import HttpResponse


def _insurer_tariff_rows(insurer, tariff):
    """
    Build a flat list of row dicts for insurer tariff export.
    Each row represents one TariffLine with its full context.
    ---
    Construye una lista plana de filas dict para exportacion de tarifas.
    Cada fila representa una TariffLine con su contexto completo.
    """
    rows = []
    if not tariff:
        return rows
    lines = tariff.lines.select_related("vehicle_type").order_by(
        "vehicle_type__sort_order", "vehicle_type__name", "concept"
    )
    for line in lines:
        rows.append({
            "Aseguradora": insurer.insurer_company_name or insurer.name,
            "Empresa prestadora": insurer.service_company_name,
            "Codigo": insurer.code,
            "Ano tarifa": tariff.year,
            "Valida desde": tariff.valid_from.strftime("%d/%m/%Y"),
            "Valida hasta": tariff.valid_to.strftime("%d/%m/%Y") if tariff.valid_to else "Activa",
            "Tipo de vehiculo": line.vehicle_type.name if line.vehicle_type else "General",
            "Concepto": line.concept.label,
            "Unidad": line.get_unit_display(),
            "Precio": str(line.price),
            "Umbral km": str(line.km_threshold) if line.km_threshold is not None else "",
            "Unidades minimas": str(line.min_units) if line.min_units is not None else "",
        })
    return rows


def _budget_rows(qs):
    """
    Build a flat list of row dicts for budget history export.
    Each row represents one Budget with its BudgetLine breakdown.
    ---
    Construye una lista plana de filas dict para exportacion de historial.
    Cada fila representa un Budget con su desglose BudgetLine.
    """
    rows = []
    for budget in qs.prefetch_related("lines"):
        base = {
            "ID": budget.pk,
            "Fecha servicio": budget.service_date.strftime("%d/%m/%Y") if budget.service_date else "",
            "Fecha creacion": budget.created_at.strftime("%d/%m/%Y %H:%M"),
            "Aseguradora": budget.insurer.name,
            "Tipo de vehiculo": budget.vehicle_type.name,
            "Km totales": str(budget.km_total),
            "Nocturno/Festivo": "Si" if budget.is_night_or_holiday else "No",
            "Cargado": "Si" if budget.is_loaded else "No",
            "Total": str(budget.total_amount),
            "IVA aplicado": "Si" if budget.apply_iva else "No",
            "Total con IVA": str(budget.total_amount_with_iva) if budget.total_amount_with_iva else "",
            "Estado": budget.get_status_display(),
        }
        lines = budget.lines.exclude(concept_code="IVA").order_by("sort_order")
        if lines.exists():
            for line in lines:
                row = dict(base)
                row["Concepto desglose"] = line.concept_label
                row["Unidades desglose"] = str(line.units)
                row["Precio unitario desglose"] = str(line.unit_price)
                row["Subtotal desglose"] = str(line.subtotal)
                rows.append(row)
        else:
            base["Concepto desglose"] = ""
            base["Unidades desglose"] = ""
            base["Precio unitario desglose"] = ""
            base["Subtotal desglose"] = ""
            rows.append(base)
    return rows


def _build_tariff_qs(insurer):
    """
    Return the active InsurerTariff for the given insurer, or None.
    ---
    Devuelve la InsurerTariff activa para la aseguradora dada, o None.
    """
    from budgets.models import InsurerTariff
    return InsurerTariff.objects.filter(
        insurer=insurer,
        valid_to__isnull=True,
    ).prefetch_related("lines__vehicle_type").first()


def _build_budget_qs(request, company):
    """
    Replicate BudgetHistoryView filters from GET params and return queryset.
    ---
    Replica los filtros de BudgetHistoryView desde GET params y devuelve el queryset.
    """
    qs = Budget.objects.filter(company=company).select_related(
        "insurer", "vehicle_type", "operator__user", "tariff",
    ).order_by("-created_at")
    insurer_id = request.GET.get("insurer")
    status = request.GET.get("status")
    date_from = request.GET.get("date_from")
    date_to = request.GET.get("date_to")
    if insurer_id:
        qs = qs.filter(insurer_id=insurer_id)
    if status:
        qs = qs.filter(status=status)
    if date_from:
        qs = qs.filter(service_date__gte=date_from)
    if date_to:
        qs = qs.filter(service_date__lte=date_to)
    return qs


class InsurerTariffExportCsvView(AdminRoleRequiredMixin, View):
    """
    Exports the active tariff of an insurer as a CSV file.
    ---
    Exporta la tarifa activa de una aseguradora como archivo CSV.
    """

    def get(self, request, pk):
        """
        Build and return the CSV response for the insurer tariff.
        ---
        Construye y devuelve la respuesta CSV para la tarifa de la aseguradora.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(Insurer, pk=pk, company=company_user.company)
        tariff = _build_tariff_qs(insurer)
        rows = _insurer_tariff_rows(insurer, tariff)

        response = HttpResponse(content_type="text/csv; charset=utf-8-sig")
        filename = f"tarifa_{insurer.code}_{tariff.year if tariff else 'sin_tarifa'}.csv"
        response["Content-Disposition"] = f'attachment; filename="{filename}"'

        if not rows:
            response.write("Sin datos de tarifa disponibles.")
            return response

        writer = csv.DictWriter(response, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)
        return response


class InsurerTariffExportExcelView(AdminRoleRequiredMixin, View):
    """
    Exports the active tariff of an insurer as an Excel file.
    ---
    Exporta la tarifa activa de una aseguradora como archivo Excel.
    """

    def get(self, request, pk):
        """
        Build and return the Excel response for the insurer tariff.
        ---
        Construye y devuelve la respuesta Excel para la tarifa de la aseguradora.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(Insurer, pk=pk, company=company_user.company)
        tariff = _build_tariff_qs(insurer)
        rows = _insurer_tariff_rows(insurer, tariff)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Tarifa"

        header_font = Font(bold=True, color="FFFFFF", size=10)
        header_fill = PatternFill(fill_type="solid", fgColor="1F4E79")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        if not rows:
            ws.append(["Sin datos de tarifa disponibles."])
        else:
            headers = list(rows[0].keys())
            ws.append(headers)
            for col_idx, _ in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col_idx)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            for row in rows:
                ws.append(list(row.values()))
            for col in ws.columns:
                max_len = max(len(str(c.value or "")) for c in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        filename = f"tarifa_{insurer.code}_{tariff.year if tariff else 'sin_tarifa'}.xlsx"
        response = HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

class InsurerTariffExportPdfView(AdminRoleRequiredMixin, View):
    """
    Exports the active tariff of an insurer as a PDF file via WeasyPrint.
    ---
    Exporta la tarifa activa de una aseguradora como PDF via WeasyPrint.
    """

    def get(self, request, pk):
        """
        Build and return the PDF response for the insurer tariff.
        ---
        Construye y devuelve la respuesta PDF para la tarifa de la aseguradora.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(Insurer, pk=pk, company=company_user.company)
        tariff = _build_tariff_qs(insurer)
        rows = _insurer_tariff_rows(insurer, tariff)

        title = f"Tarifa {insurer.name} - {tariff.year if tariff else 'Sin tarifa activa'}"
        if rows:
            headers = list(rows[0].keys())
            thead = "".join(f"<th>{h}</th>" for h in headers)
            tbody = ""
            for i, row in enumerate(rows):
                bg = "#f0f4f8" if i % 2 == 0 else "#ffffff"
                cells = "".join(f"<td>{v}</td>" for v in row.values())
                tbody += f"<tr style='background:{bg}'>{cells}</tr>"
            table_html = f"<table><thead><tr>{thead}</tr></thead><tbody>{tbody}</tbody></table>"
        else:
            table_html = "<p>Sin datos de tarifa disponibles.</p>"

        html_content = (
            "<!DOCTYPE html><html lang='es'><head><meta charset='utf-8'>"
            "<style>"
            "body{font-family:Arial,sans-serif;font-size:8pt;margin:20px}"
            "h1{color:#1F4E79;font-size:12pt;margin-bottom:10px}"
            "table{width:100%;border-collapse:collapse}"
            "th{background:#1F4E79;color:white;padding:4px 6px;font-size:7pt;text-align:left}"
            "td{padding:3px 6px;font-size:7pt;border-bottom:1px solid #e0e0e0}"
            "</style></head><body>"
            f"<h1>{title}</h1>{table_html}"
            "</body></html>"
        )

        pdf_file = weasyprint.HTML(string=html_content).write_pdf()
        filename = f"tarifa_{insurer.code}_{tariff.year if tariff else 'sin_tarifa'}.pdf"
        response = HttpResponse(pdf_file, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class InsurerTariffExportWordView(AdminRoleRequiredMixin, View):
    """
    Exports the active tariff of an insurer as a Word (.docx) file.
    ---
    Exporta la tarifa activa de una aseguradora como archivo Word (.docx).
    """

    def get(self, request, pk):
        """
        Build and return the Word response for the insurer tariff.
        ---
        Construye y devuelve la respuesta Word para la tarifa de la aseguradora.
        """
        company_user = _get_company_user(request)
        insurer = get_object_or_404(Insurer, pk=pk, company=company_user.company)
        tariff = _build_tariff_qs(insurer)
        rows = _insurer_tariff_rows(insurer, tariff)

        doc = Document()
        title_para = doc.add_heading(level=1)
        title_run = title_para.add_run(
            f"Tarifa {insurer.name} - {tariff.year if tariff else 'Sin tarifa activa'}"
        )
        title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        if not rows:
            doc.add_paragraph("Sin datos de tarifa disponibles.")
        else:
            headers = list(rows[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            for row in rows:
                row_cells = table.add_row().cells
                for i, v in enumerate(row.values()):
                    row_cells[i].text = str(v)
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"tarifa_{insurer.code}_{tariff.year if tariff else 'sin_tarifa'}.docx"
        response = HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class BudgetExportCsvView(AdminRoleRequiredMixin, View):
    """
    Exports the filtered budget history as a CSV file.
    ---
    Exporta el historial de presupuestos filtrado como archivo CSV.
    """

    def get(self, request):
        """
        Build and return the CSV response for the budget history.
        ---
        Construye y devuelve la respuesta CSV para el historial de presupuestos.
        """
        company_user = _get_company_user(request)
        qs = _build_budget_qs(request, company_user.company)
        rows = _budget_rows(qs)

        response = HttpResponse(content_type="text/csv; charset=utf-8-sig")
        response["Content-Disposition"] = 'attachment; filename="presupuestos.csv"'

        if not rows:
            response.write("Sin presupuestos para los filtros seleccionados.")
            return response

        writer = csv.DictWriter(response, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)
        return response


class BudgetExportExcelView(AdminRoleRequiredMixin, View):
    """
    Exports the filtered budget history as an Excel file.
    ---
    Exporta el historial de presupuestos filtrado como archivo Excel.
    """

    def get(self, request):
        """
        Build and return the Excel response for the budget history.
        ---
        Construye y devuelve la respuesta Excel para el historial de presupuestos.
        """
        company_user = _get_company_user(request)
        qs = _build_budget_qs(request, company_user.company)
        rows = _budget_rows(qs)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Presupuestos"

        header_font = Font(bold=True, color="FFFFFF", size=10)
        header_fill = PatternFill(fill_type="solid", fgColor="1F4E79")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        if not rows:
            ws.append(["Sin presupuestos para los filtros seleccionados."])
        else:
            headers = list(rows[0].keys())
            ws.append(headers)
            for col_idx, _ in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col_idx)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            for row in rows:
                ws.append(list(row.values()))
            for col in ws.columns:
                max_len = max(len(str(c.value or "")) for c in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="presupuestos.xlsx"'
        return response


class BudgetExportPdfView(AdminRoleRequiredMixin, View):
    """
    Exports the filtered budget history as a PDF file via WeasyPrint.
    ---
    Exporta el historial de presupuestos filtrado como PDF via WeasyPrint.
    """

    def get(self, request):
        """
        Build and return the PDF response for the budget history.
        ---
        Construye y devuelve la respuesta PDF para el historial de presupuestos.
        """
        company_user = _get_company_user(request)
        qs = _build_budget_qs(request, company_user.company)
        rows = _budget_rows(qs)

        if rows:
            headers = list(rows[0].keys())
            thead = "".join(f"<th>{h}</th>" for h in headers)
            tbody = ""
            for i, row in enumerate(rows):
                bg = "#f0f4f8" if i % 2 == 0 else "#ffffff"
                cells = "".join(f"<td>{v}</td>" for v in row.values())
                tbody += f"<tr style='background:{bg}'>{cells}</tr>"
            table_html = f"<table><thead><tr>{thead}</tr></thead><tbody>{tbody}</tbody></table>"
        else:
            table_html = "<p>Sin presupuestos para los filtros seleccionados.</p>"

        html_content = (
            "<!DOCTYPE html><html lang='es'><head><meta charset='utf-8'>"
            "<style>"
            "body{font-family:Arial,sans-serif;font-size:8pt;margin:20px}"
            "h1{color:#1F4E79;font-size:12pt;margin-bottom:10px}"
            "table{width:100%;border-collapse:collapse}"
            "th{background:#1F4E79;color:white;padding:4px 6px;font-size:7pt;text-align:left}"
            "td{padding:3px 6px;font-size:7pt;border-bottom:1px solid #e0e0e0}"
            "</style></head><body>"
            f"<h1>Historial de presupuestos</h1>{table_html}"
            "</body></html>"
        )

        pdf_file = weasyprint.HTML(string=html_content).write_pdf()
        response = HttpResponse(pdf_file, content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="presupuestos.pdf"'
        return response


class BudgetExportWordView(AdminRoleRequiredMixin, View):
    """
    Exports the filtered budget history as a Word (.docx) file.
    ---
    Exporta el historial de presupuestos filtrado como archivo Word (.docx).
    """

    def get(self, request):
        """
        Build and return the Word response for the budget history.
        ---
        Construye y devuelve la respuesta Word para el historial de presupuestos.
        """
        company_user = _get_company_user(request)
        qs = _build_budget_qs(request, company_user.company)
        rows = _budget_rows(qs)

        doc = Document()
        title_para = doc.add_heading(level=1)
        title_run = title_para.add_run("Historial de presupuestos")
        title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        if not rows:
            doc.add_paragraph("Sin presupuestos para los filtros seleccionados.")
        else:
            headers = list(rows[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            for row in rows:
                row_cells = table.add_row().cells
                for i, v in enumerate(row.values()):
                    row_cells[i].text = str(v)
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="presupuestos.docx"'
        return response


# ---------------------------------------------------------------------------
# Base global view — full company base list with filter by insurer. ADMIN only.
# Vista global de bases — listado completo de bases de empresa con filtro por
# aseguradora. Solo ADMIN.
# ---------------------------------------------------------------------------


class BaseClearCoordsView(AdminRoleRequiredMixin, View):
    """
    HTMX POST endpoint. Receives a list of base PKs and sets their
    latitude and longitude to null. Returns an HTML fragment with
    the result summary suitable for HTMX swap.
    ---
    Endpoint POST HTMX. Recibe una lista de PKs de bases y pone su
    latitud y longitud a null. Devuelve un fragmento HTML con el
    resumen del resultado para el swap HTMX.
    """

    def post(self, request):
        """
        Clear coordinates for the given base PKs.
        ---
        Limpia las coordenadas de las bases con los PKs indicados.
        """
        company_user = _get_company_user(request)
        company      = company_user.company

        # Parse PKs from POST — sent as repeated 'base_pks' values.
        # Parsear PKs del POST — enviados como valores repetidos 'base_pks'.
        raw_pks = request.POST.getlist("base_pks")
        try:
            pks = [int(pk) for pk in raw_pks if pk.strip().isdigit()]
        except (ValueError, AttributeError):
            pks = []

        if not pks:
            from django.http import HttpResponse as _HR
            return _HR(
                '<div class="alert alert-warning py-2 small mb-2">'
                '<i class="bi bi-exclamation-triangle me-1"></i>'
                'No se han seleccionado bases.'
                '</div>'
            )

        # Only allow clearing bases that belong to this company.
        # Solo permitir limpiar bases que pertenecen a esta empresa.
        updated = Base.objects.filter(
            pk__in=pks,
            company=company,
        ).update(latitude=None, longitude=None)

        from django.http import HttpResponse as _HR
        return _HR(
            f'<div class="alert alert-success alert-dismissible fade show py-2 small mb-2" role="alert">'
            f'<i class="bi bi-check2-circle me-1"></i>'
            f'<strong>Coordenadas limpiadas en {updated} base(s).</strong> '
            f'Recarga la página para ver los cambios reflejados en la tabla.'
            f'<button type="button" class="btn-close btn-sm" data-bs-dismiss="alert"></button>'
            f'</div>'
        )


class BaseSyncCalendarsView(AdminRoleRequiredMixin, View):
    """
    HTMX POST endpoint. Syncs the labour calendars of all active company
    bases via the calendariosnacionales.com public API. Returns an inline
    HTML fragment with the sync results (ok count, error details).
    Reuses _fetch_holidays() and MUNICIPALITY_MAP from the management command.
    ---
    Endpoint POST HTMX. Sincroniza los calendarios laborales de todas las
    bases activas de la empresa via la API publica de calendariosnacionales.com.
    Devuelve un fragmento HTML inline con los resultados de la sincronizacion.
    Reutiliza _fetch_holidays() del management command sync_base_calendars.
    """

    def post(self, request):
        """
        Sync labour calendars for all active company bases.
        Returns an HTML fragment suitable for HTMX swap.
        ---
        Sincroniza los calendarios laborales de todas las bases activas.
        Devuelve un fragmento HTML para el swap HTMX en base_global.
        """
        import json as _json
        import datetime as _dt
        from django.utils import timezone as _tz
        from budgets.management.commands.sync_base_calendars import _fetch_holidays

        company_user = _get_company_user(request)
        company      = company_user.company

        # Target year: next year if Q4, otherwise current year.
        # Año objetivo: siguiente si Q4, si no el actual.
        today = _dt.date.today()
        year  = today.year + 1 if today.month >= 10 else today.year

        bases         = list(Base.objects.filter(company=company, is_active=True))
        ok_results    = []
        error_results = []

        for base in bases:
            try:
                holidays = _fetch_holidays(base.municipality, year)
                base.labor_calendar     = _json.dumps(holidays, ensure_ascii=False)
                base.calendar_synced_at = _tz.now()
                base.save(update_fields=["labor_calendar", "calendar_synced_at"])
                ok_results.append({"name": base.name, "count": len(holidays), "year": year})
            except KeyError as exc:
                logger.warning(
                    "# [budgets] sync_base_calendars: falta clave en la "
                    "respuesta para base pk=%r (%s): %s",
                    base.pk, base.municipality, exc, exc_info=True,
                )
                error_results.append({
                    "name": base.name,
                    "error": "Respuesta incompleta del servicio de festivos.",
                })
            except Exception as exc:
                logger.error(
                    "# [budgets] sync_base_calendars: error inesperado "
                    "para base pk=%r (%s): %s",
                    base.pk, base.municipality, exc, exc_info=True,
                )
                error_results.append({
                    "name": base.name,
                    "error": "Error inesperado. Revisa el registro del servidor.",
                })

        html_parts = []
        if ok_results:
            html_parts.append(
                '<div class="alert alert-success alert-dismissible fade show py-2 small mb-2" role="alert">'
                f'<i class="bi bi-check2-circle me-1"></i>'
                f'<strong>{len(ok_results)} base(s) sincronizada(s)</strong> para {year}:'
                '<ul class="mb-0 mt-1">'
            )
            for r in ok_results:
                html_parts.append(f'<li>{r["name"]} &mdash; {r["count"]} festivos</li>')
            html_parts.append('</ul><button type="button" class="btn-close btn-sm" data-bs-dismiss="alert"></button></div>')
        if error_results:
            html_parts.append(
                '<div class="alert alert-warning alert-dismissible fade show py-2 small mb-2" role="alert">'
                f'<i class="bi bi-exclamation-triangle me-1"></i>'
                f'<strong>{len(error_results)} base(s) con error:</strong>'
                '<ul class="mb-0 mt-1">'
            )
            for r in error_results:
                html_parts.append(f'<li>{r["name"]}: {r["error"]}</li>')
            html_parts.append('</ul><button type="button" class="btn-close btn-sm" data-bs-dismiss="alert"></button></div>')
        if not ok_results and not error_results:
            html_parts.append(
                '<div class="alert alert-info py-2 small mb-2">'
                '<i class="bi bi-info-circle me-1"></i>'
                'No hay bases activas para sincronizar.'
                '</div>'
            )

        from django.http import HttpResponse as _HR
        return _HR("".join(html_parts))


class BaseGlobalView(AdminRoleRequiredMixin, View):
    """
    Lists all service bases for the company with optional filter by insurer.
    Shows per-base coordinates, labor calendar sync status and global is_active
    flag. Provides inline Google Maps edit, global Base.is_active toggle,
    delete and a new-base creation form. ADMIN only.
    ---
    Lista todas las bases de servicio de la empresa con filtro opcional por
    aseguradora. Muestra coordenadas, estado de sincronizacion del calendario
    laboral y flag global is_active por base. Proporciona edicion inline con
    Google Maps, toggle global Base.is_active, baja y formulario de alta.
    Solo ADMIN.
    """

    template_name = "budgets/base_global.html"

    def get(self, request):
        """
        Render the global base list for the company.
        Accepts optional GET param insurer_id to filter by insurer.
        ---
        Renderiza el listado global de bases de la empresa.
        Acepta parametro GET opcional insurer_id para filtrar por aseguradora.
        """
        company_user = _get_company_user(request)
        company = company_user.company

        # Optional insurer filter from GET param.
        # Filtro opcional por aseguradora desde parametro GET.
        insurer_id = request.GET.get("insurer_id", "").strip()
        selected_insurer = None
        bases_qs = Base.objects.filter(company=company).order_by("name")
        if insurer_id:
            try:
                selected_insurer = Insurer.objects.get(
                    pk=int(insurer_id), company=company
                )
                # Filter to bases linked to this insurer via InsurerBase.
                # Filtrar a bases vinculadas a esta aseguradora via InsurerBase.
                linked_base_ids = InsurerBase.objects.filter(
                    insurer=selected_insurer
                ).values_list("base_id", flat=True)
                bases_qs = bases_qs.filter(pk__in=linked_base_ids)
            except (ValueError, Insurer.DoesNotExist):
                selected_insurer = None

        # Annotate each base with its linked insurers and insurer count.
        # Anotar cada base con sus aseguradoras vinculadas y el conteo.
        ib_by_base = {}
        for ib in InsurerBase.objects.filter(
            base__company=company
        ).select_related("insurer").order_by("insurer__name"):
            ib_by_base.setdefault(ib.base_id, []).append(ib)

        bases = []
        for base in bases_qs:
            base.linked_insurer_bases = ib_by_base.get(base.pk, [])
            bases.append(base)

        # All insurers for the filter dropdown.
        # Todas las aseguradoras para el desplegable de filtro.
        insurers = Insurer.objects.filter(
            company=company
        ).order_by("name")

        base_form = BaseForm()
        ctx = _build_base_context(request, {
            "bases": bases,
            "insurers": insurers,
            "selected_insurer": selected_insurer,
            "base_form": base_form,
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
            "active_nav": "budgets_bases",
        })
        return render(request, self.template_name, ctx)

    def post(self, request):
        """
        Create a new Base for the company (no insurer association at creation).
        Redirects to the global base list on success.
        Re-renders the form with errors on failure.
        ---
        Crea una nueva Base para la empresa (sin aseguradora asociada en el alta).
        Redirige al listado global en exito. Vuelve a renderizar el formulario
        con errores en caso de fallo.
        """
        company_user = _get_company_user(request)
        company = company_user.company
        form = BaseForm(request.POST)
        if form.is_valid():
            base = form.save(commit=False)
            base.company = company
            base.save()
            from django.contrib import messages as _messages
            _messages.success(
                request,
                f"Base '{base.name}' creada correctamente. "
                "Asignala a aseguradoras desde la vista de gestion de bases."
            )
            return redirect("budgets:base_global")
        # Re-render with form errors.
        # Volver a renderizar con errores de formulario.
        bases_qs = Base.objects.filter(company=company).order_by("name")
        ib_by_base = {}
        for ib in InsurerBase.objects.filter(
            base__company=company
        ).select_related("insurer").order_by("insurer__name"):
            ib_by_base.setdefault(ib.base_id, []).append(ib)
        bases = []
        for base in bases_qs:
            base.linked_insurer_bases = ib_by_base.get(base.pk, [])
            bases.append(base)
        insurers = Insurer.objects.filter(company=company).order_by("name")
        ctx = _build_base_context(request, {
            "bases": bases,
            "insurers": insurers,
            "selected_insurer": None,
            "base_form": form,
            "google_maps_api_key": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
            "active_nav": "budgets_bases",
        })
        return render(request, self.template_name, ctx)


# ---------------------------------------------------------------------------
# ALBARÁN DEMO — H17 offline prototype view.
# Vista prototipo offline H17 — formulario de albarán con firma canvas.
# ---------------------------------------------------------------------------

class AlbaranDemoView(AssistanceRequiredMixin, View):
    """
    Prototype view for the H17 offline delivery-note flow.
    Renders a mobile-first form with five fields (name, DNI, phone,
    vehicle plate, canvas signature) and handles the POST submission,
    storing the signed data in the session for demo purposes only.
    ---
    Vista prototipo para el flujo offline de albarán H17.
    Renderiza un formulario mobile-first con cinco campos (nombre, DNI,
    teléfono, matrícula vehículo, firma canvas) y gestiona el POST,
    almacenando los datos firmados en la sesión solo con fines de demo.
    """

    template_name = "budgets/albaran_demo.html"

    def get(self, request, *args, **kwargs):
        """
        Renders the blank demo delivery-note form.
        Renderiza el formulario de albarán demo en blanco.
        """
        company_user = get_object_or_404(CompanyUser, user=request.user)
        submitted = request.session.pop("albaran_demo_submitted", False)
        return render(request, self.template_name, {
            "active_nav":   "albaran_demo",
            "company_user": company_user,
            "submitted":    submitted,
        })

    def post(self, request, *args, **kwargs):
        """
        Handles the signed form submission. Stores a success flag in the
        session and redirects to GET to avoid double-submit on refresh.
        ---
        Gestiona el envío del formulario firmado. Almacena un flag de éxito
        en la sesión y redirige a GET para evitar doble envío al recargar.
        """
        # In the demo, we simply acknowledge receipt. No model persistence yet.
        # En el demo, solo acusamos recibo. Sin persistencia en modelo aún.
        request.session["albaran_demo_submitted"] = True
        return redirect("budgets:albaran_demo")


# ---------------------------------------------------------------------------
# H17 — WORK ORDER ASSISTANCE VIEWS
# Vistas de órdenes de trabajo de asistencia (Hito 17)
# ---------------------------------------------------------------------------

class WorkOrderCreateFromBudgetView(AssistanceRequiredMixin, View):
    """
    Creates a WorkOrderAssistance from an accepted Budget.
    Accessible via the 'Generar orden de trabajo' button on BudgetResultView
    when budget.status == ACCEPTED. Inherits insurer, vehicle_type and base
    from the source Budget. Redirects to WorkOrderDetailView on success.
    ---
    Crea una WorkOrderAssistance a partir de un Budget aceptado.
    Accesible desde el botón 'Generar orden de trabajo' en BudgetResultView
    cuando budget.status == ACCEPTED. Hereda aseguradora, tipo de vehículo
    y base del Budget de origen. Redirige a WorkOrderDetailView en éxito.
    """

    def post(self, request, pk):
        """
        POST handler. Resolves the Budget, validates its status is ACCEPTED,
        creates the WorkOrderAssistance and redirects to its detail view.
        ---
        Manejador POST. Resuelve el Budget, valida que su estado sea ACCEPTED,
        crea la WorkOrderAssistance y redirige a su vista de detalle.
        """
        from budgets.models import WorkOrderAssistance
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        budget = get_object_or_404(
            Budget,
            pk=pk,
            company=company_user.company,
        )
        if budget.status != Budget.STATUS_ACCEPTED:
            messages.error(
                request,
                "Solo se puede generar una orden de trabajo desde un presupuesto aceptado.",
            )
            return redirect("budgets:result", pk=budget.pk)
        # Guard: avoid creating duplicate work orders for the same budget.
        # Guardia: evitar crear órdenes duplicadas para el mismo presupuesto.
        existing = WorkOrderAssistance.objects.filter(budget=budget).first()
        if existing:
            messages.warning(
                request,
                "Este presupuesto ya tiene una orden de trabajo asociada.",
            )
            return redirect("budgets:work_order_detail", pk=existing.pk)
        with transaction.atomic():
            work_order = WorkOrderAssistance.objects.create(
                company=company_user.company,
                budget=budget,
                insurer=budget.insurer,
                vehicle_type=budget.vehicle_type,
                base=budget.base,
                operator=company_user,
                created_by=company_user,
                service_date=budget.service_date,
            )
        messages.success(
            request,
            f"Orden de trabajo {work_order.work_order_number} creada correctamente.",
        )
        return redirect("budgets:work_order_detail", pk=work_order.pk)


class WorkOrderCreateDirectView(AssistanceRequiredMixin, View):
    """
    Creates a WorkOrderAssistance directly, without a prior Budget.
    Renders a form pre-populated with the company's active insurers.
    On valid POST, creates the order and redirects to WorkOrderDetailView.
    ---
    Crea una WorkOrderAssistance directamente, sin presupuesto previo.
    Renderiza un formulario pre-poblado con las aseguradoras activas de
    la empresa. En POST válido, crea la orden y redirige a WorkOrderDetailView.
    """

    template_name = "budgets/work_order_create_direct.html"

    def get(self, request):
        """
        Renders the direct work order creation form.
        ---
        Renderiza el formulario de creación directa de orden de trabajo.
        """
        from budgets.models import WorkOrderAssistance
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        insurers = Insurer.objects.filter(
            company=company_user.company,
            is_active=True,
        ).order_by("name")
        ctx = _build_base_context(request, {
            "insurers": insurers,
            "active_nav": "work_orders",
        })
        return render(request, self.template_name, ctx)

    def post(self, request):
        """
        Handles the direct work order creation form submission.
        Validates required fields, creates the WorkOrderAssistance and
        redirects to its detail view.
        ---
        Gestiona el envío del formulario de creación directa.
        Valida los campos obligatorios, crea la WorkOrderAssistance y
        redirige a su vista de detalle.
        """
        from budgets.models import WorkOrderAssistance
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        insurer_id  = request.POST.get("insurer_id", "").strip()
        service_date = request.POST.get("service_date", "").strip()
        if not insurer_id or not service_date:
            messages.error(request, "La aseguradora y la fecha del servicio son obligatorias.")
            return redirect("budgets:work_order_create_direct")
        insurer = get_object_or_404(
            Insurer,
            pk=int(insurer_id),
            company=company_user.company,
            is_active=True,
        )
        with transaction.atomic():
            work_order = WorkOrderAssistance.objects.create(
                company=company_user.company,
                budget=None,
                insurer=insurer,
                vehicle_type=insurer.vehicle_types.first(),
                base=None,
                operator=company_user,
                created_by=company_user,
                service_date=service_date,
            )
        messages.success(
            request,
            f"Orden de trabajo {work_order.work_order_number} creada correctamente.",
        )
        return redirect("budgets:work_order_detail", pk=work_order.pk)


class WorkOrderDetailView(AssistanceRequiredMixin, View):
    """
    Displays the full detail of a WorkOrderAssistance record.
    Shows all TIREA fields, linked units and their status.
    Accessible to ASSISTANCE and ADMIN roles scoped to the company.
    ---
    Muestra el detalle completo de un registro WorkOrderAssistance.
    Muestra todos los campos TIREA, unidades vinculadas y su estado.
    Accesible para los roles ASSISTANCE y ADMIN con ámbito de empresa.
    """

    template_name = "budgets/work_order_detail.html"

    def get(self, request, pk):
        """
        Renders the work order detail page.
        ---
        Renderiza la página de detalle de la orden de trabajo.
        """
        from budgets.models import WorkOrderAssistance
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        work_order = get_object_or_404(
            WorkOrderAssistance,
            pk=pk,
            company=company_user.company,
        )
        units = work_order.units.select_related(
            "operator", "operator_2"
        ).prefetch_related(
            "signature", "incidence"
        ).order_by("unit_number")
        ctx = _build_base_context(request, {
            "work_order": work_order,
            "units": units,
            "active_nav": "work_orders",
        })
        return render(request, self.template_name, ctx)


class WorkOrderAlbaranView(AssistanceRequiredMixin, View):
    """
    Renders the mobile-first albarán form for a WorkOrderAssistanceUnit.
    Used by the operator on the Android app or mobile browser to fill in
    service data and capture the client signature. Handles GET (render)
    and POST (save data + signature).
    ---
    Renderiza el formulario de albarán mobile-first para una
    WorkOrderAssistanceUnit. Usado por el operario en la app Android o
    navegador móvil para rellenar datos del servicio y capturar la firma
    del cliente. Gestiona GET (renderizar) y POST (guardar datos + firma).
    """

    template_name = "budgets/albaran_operario.html"

    def get(self, request, pk):
        """
        Renders the albarán form for the given unit.
        ---
        Renderiza el formulario de albarán para la unidad indicada.
        """
        from budgets.models import WorkOrderAssistanceUnit
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        unit = get_object_or_404(
            WorkOrderAssistanceUnit,
            pk=pk,
            work_order__company=company_user.company,
        )
        ctx = _build_base_context(request, {
            "unit": unit,
            "work_order": unit.work_order,
            "active_nav": "work_orders",
        })
        return render(request, self.template_name, ctx)

    def post(self, request, pk):
        """
        Saves the albarán data and client signature for the given unit.
        Marks the unit as COMPLETED and creates the signature record.
        Redirects to WorkOrderDetailView on success.
        ---
        Guarda los datos del albarán y la firma del cliente para la unidad.
        Marca la unidad como COMPLETED y crea el registro de firma.
        Redirige a WorkOrderDetailView en éxito.
        """
        from budgets.models import WorkOrderAssistanceUnit, WorkOrderAssistanceSignature
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        unit = get_object_or_404(
            WorkOrderAssistanceUnit,
            pk=pk,
            work_order__company=company_user.company,
        )
        phase1_km     = request.POST.get("phase1_km", "").strip() or None
        phase2_km     = request.POST.get("phase2_km", "").strip() or None
        departure_fee = request.POST.get("departure_fee", "").strip() or None
        rescue_hours  = request.POST.get("rescue_hours", "").strip() or None
        wait_hours    = request.POST.get("wait_hours", "").strip() or None
        signature_data = request.POST.get("signature_data", "").strip()
        signer_name    = request.POST.get("signer_name", "").strip()
        signed_offline = request.POST.get("signed_offline", "false") == "true"
        with transaction.atomic():
            unit.phase1_km     = phase1_km
            unit.phase2_km     = phase2_km
            unit.departure_fee = departure_fee
            unit.rescue_hours  = rescue_hours
            unit.wait_hours    = wait_hours
            unit.status        = WorkOrderAssistanceUnit.STATUS_COMPLETED
            unit.save()
            if signature_data:
                WorkOrderAssistanceSignature.objects.update_or_create(
                    unit=unit,
                    defaults={
                        "signature_data": signature_data,
                        "signer_name":    signer_name,
                        "signed_offline": signed_offline,
                    },
                )
        messages.success(request, "Albarán guardado y firmado correctamente.")
        return redirect("budgets:work_order_detail", pk=unit.work_order.pk)


class TollSegmentListView(AdminRoleRequiredMixin, View):
    """
    Lists all TollSegment records with filtering by road_code, tariff_level
    and active status. Creation is handled by TollSegmentCreateView.
    Accessible to ADMIN role only.
    ---
    Lista todos los registros TollSegment con filtrado por road_code,
    tariff_level y estado. La creación se gestiona en TollSegmentCreateView.
    Accesible solo para el rol ADMIN.
    """

    template_name = "budgets/toll_segment_list.html"

    def get(self, request):
        """
        Render the toll segment list with optional filters.
        ---
        Renderiza el listado de tramos de peaje con filtros opcionales.
        """
        from budgets.models import TollSegment
        road_filter = request.GET.get("road_code", "").strip().upper()
        level_filter = request.GET.get("tariff_level", "").strip()
        active_filter = request.GET.get("is_active", "1").strip()

        qs = TollSegment.objects.all().order_by(
            "road_code", "section_name", "origin_name", "dest_name",
        )
        if road_filter:
            qs = qs.filter(road_code=road_filter)
        if level_filter:
            qs = qs.filter(tariff_level=level_filter)
        if active_filter == "1":
            qs = qs.filter(is_active=True)
        elif active_filter == "0":
            qs = qs.filter(is_active=False)

        road_codes = (
            TollSegment.objects
            .values_list("road_code", flat=True)
            .distinct()
            .order_by("road_code")
        )

        ctx = _build_base_context(request, {
            "segments":             qs,
            "road_codes":           road_codes,
            "road_filter":          road_filter,
            "level_filter":         level_filter,
            "active_filter":        active_filter,
            "tariff_level_choices": TollSegment.TARIFF_LEVEL_CHOICES,
            "active_nav":           "budgets_toll_segments",
            "total":                qs.count(),
            "toll_vehicle_type":    _get_company_user(request).company.toll_vehicle_type,
            "toll_markup_percent":  _get_company_user(request).company.toll_markup_percent,
            "toll_vehicle_choices": [
                ("LIGHT",   "Ligero"),
                ("HEAVY_1", "Pesado 1"),
                ("HEAVY_2", "Pesado 2"),
            ],
        })
        return render(request, self.template_name, ctx)


class TollSegmentConfigView(AdminRoleRequiredMixin, View):
    """
    POST — saves toll configuration (vehicle type and markup percent) to
    the company record. Redirects back to the toll segment list.
    Accessible to ADMIN role only.
    ---
    POST — guarda la configuración de peajes (tipo de vehículo y recargo)
    en el registro de empresa. Redirige al listado de tramos de peaje.
    Accesible solo para el rol ADMIN.
    """

    def post(self, request):
        from ivr_config.models import Company
        company_user = _get_company_user(request)
        company = company_user.company

        toll_vehicle_type = request.POST.get(
            "toll_vehicle_type", Company.TOLL_VEHICLE_HEAVY_1
        ).strip()
        if toll_vehicle_type not in dict(Company.TOLL_VEHICLE_CHOICES):
            toll_vehicle_type = Company.TOLL_VEHICLE_HEAVY_1

        try:
            from decimal import Decimal as _Decimal
            toll_markup_percent = _Decimal(
                request.POST.get("toll_markup_percent", "0").strip() or "0"
            )
            if toll_markup_percent < 0:
                toll_markup_percent = _Decimal("0")
        except Exception:
            from decimal import Decimal as _Decimal
            toll_markup_percent = _Decimal("0")

        Company.objects.filter(pk=company.pk).update(
            toll_vehicle_type=toll_vehicle_type,
            toll_markup_percent=toll_markup_percent,
        )
        messages.success(
            request,
            "Configuración de peajes actualizada.",
        )
        return redirect("budgets:toll_segment_list")


class TollSegmentCreateView(AdminRoleRequiredMixin, View):
    """
    Renders and processes the creation form for a new TollSegment.
    GET: empty form. POST: validates and saves, redirects to list.
    Accessible to ADMIN role only.
    ---
    Renderiza y procesa el formulario de alta de un nuevo TollSegment.
    GET: formulario vacío. POST: valida y guarda, redirige al listado.
    Accesible solo para el rol ADMIN.
    """

    template_name = "budgets/toll_segment_form.html"

    def get(self, request):
        form = TollSegmentForm(
            initial={"is_active": True, "valid_from": datetime.date.today()}
        )
        ctx = _build_base_context(request, {
            "form":                 form,
            "editing":              None,
            "tariff_level_choices": TollSegment.TARIFF_LEVEL_CHOICES,
            "active_nav":           "budgets_toll_segments",
        })
        return render(request, self.template_name, ctx)

    def post(self, request):
        from budgets.models import TollSegment
        form = TollSegmentForm(request.POST)
        if form.is_valid():
            segment = form.save()
            messages.success(
                request,
                f"Tramo {segment.road_code} "
                f"{segment.origin_name} → {segment.dest_name} "
                f"creado correctamente.",
            )
            return redirect("budgets:toll_segment_list")
        ctx = _build_base_context(request, {
            "form":                 form,
            "editing":              None,
            "tariff_level_choices": TollSegment.TARIFF_LEVEL_CHOICES,
            "active_nav":           "budgets_toll_segments",
        })
        return render(request, self.template_name, ctx)


class TollSegmentUpdateView(AdminRoleRequiredMixin, View):
    """
    Renders and processes the edit form for an existing TollSegment.
    Accessible to ADMIN role only.
    ---
    Renderiza y procesa el formulario de edicion de un TollSegment existente.
    Accesible solo para el rol ADMIN.
    """

    template_name = "budgets/toll_segment_form.html"

    def get(self, request, pk):
        """
        Render the edit form pre-filled with the TollSegment instance.
        ---
        Renderiza el formulario de edicion con los datos del TollSegment.
        """
        from budgets.models import TollSegment
        segment = get_object_or_404(TollSegment, pk=pk)
        form = TollSegmentForm(instance=segment)
        road_codes = (
            TollSegment.objects
            .values_list("road_code", flat=True)
            .distinct()
            .order_by("road_code")
        )
        qs = TollSegment.objects.filter(is_active=True).order_by(
            "road_code", "section_name", "origin_name", "dest_name",
        )
        ctx = _build_base_context(request, {
            "segments": qs,
            "form": form,
            "editing": segment,
            "road_codes": road_codes,
            "road_filter": "",
            "level_filter": "",
            "active_filter": "1",
            "tariff_level_choices": TollSegment.TARIFF_LEVEL_CHOICES,
            "active_nav": "budgets_toll_segments",
            "total": qs.count(),
        })
        return render(request, self.template_name, ctx)

    def post(self, request, pk):
        """
        Validate and save the updated TollSegment. Redirects to list.
        ---
        Valida y guarda el TollSegment actualizado. Redirige al listado.
        """
        from budgets.models import TollSegment
        segment = get_object_or_404(TollSegment, pk=pk)
        form = TollSegmentForm(request.POST, instance=segment)
        if form.is_valid():
            form.save()
            messages.success(
                request,
                f"Tramo {segment.road_code} "
                f"{segment.origin_name} → {segment.dest_name} "
                f"actualizado correctamente.",
            )
            return redirect("budgets:toll_segment_list")

        road_codes = (
            TollSegment.objects
            .values_list("road_code", flat=True)
            .distinct()
            .order_by("road_code")
        )
        qs = TollSegment.objects.filter(is_active=True).order_by(
            "road_code", "section_name", "origin_name", "dest_name",
        )
        ctx = _build_base_context(request, {
            "segments": qs,
            "form": form,
            "editing": segment,
            "road_codes": road_codes,
            "road_filter": "",
            "level_filter": "",
            "active_filter": "1",
            "tariff_level_choices": TollSegment.TARIFF_LEVEL_CHOICES,
            "active_nav": "budgets_toll_segments",
            "total": qs.count(),
        })
        return render(request, self.template_name, ctx)


class TollSegmentDeleteView(AdminRoleRequiredMixin, View):
    """
    Soft-deletes (is_active=False) or hard-deletes a TollSegment.
    POST with action=deactivate sets is_active=False.
    POST with action=delete performs hard delete.
    Accessible to ADMIN role only.
    ---
    Borra logicamente (is_active=False) o fisicamente un TollSegment.
    POST con action=deactivate pone is_active=False.
    POST con action=delete realiza borrado fisico.
    Accesible solo para el rol ADMIN.
    """

    def post(self, request, pk):
        """
        Process deactivation or deletion of the TollSegment.
        ---
        Procesa la desactivacion o eliminacion del TollSegment.
        """
        from budgets.models import TollSegment
        segment = get_object_or_404(TollSegment, pk=pk)
        action = request.POST.get("action", "deactivate")

        if action == "delete":
            label = (
                f"{segment.road_code} "
                f"{segment.origin_name} → {segment.dest_name}"
            )
            segment.delete()
            messages.success(
                request,
                f"Tramo {label} eliminado permanentemente.",
            )
        else:
            segment.is_active = False
            segment.save()
            messages.success(
                request,
                f"Tramo {segment.road_code} "
                f"{segment.origin_name} → {segment.dest_name} "
                f"desactivado correctamente.",
            )
        return redirect("budgets:toll_segment_list")


class TollSegmentBulkToggleView(AdminRoleRequiredMixin, View):
    """
    Bulk activate or deactivate a set of TollSegment records.
    POST receives a list of PKs (pks[]) and an action ('activate' or
    'deactivate'). Sets is_active accordingly on all matched records.
    Redirects back to the toll segment list with a summary message.
    Accessible to ADMIN role only.
    ---
    Activa o desactiva en bloque un conjunto de registros TollSegment.
    POST recibe una lista de PKs (pks[]) y una accion ('activate' o
    'deactivate'). Establece is_active en todos los registros coincidentes.
    Redirige al listado con un mensaje resumen.
    Accesible solo para el rol ADMIN.
    """

    def post(self, request):
        """
        Process bulk activation or deactivation.
        ---
        Procesa la activacion o desactivacion en bloque.
        """
        from budgets.models import TollSegment
        raw_pks = request.POST.getlist("pks")
        action = request.POST.get("action", "")

        # Sanitise: only numeric PKs accepted.
        # Sanitizar: solo PKs numericos aceptados.
        valid_pks = [
            int(pk) for pk in raw_pks
            if pk.isdigit()
        ]

        if not valid_pks:
            messages.warning(
                request,
                "No se ha seleccionado ningún tramo de peaje.",
            )
            return redirect("budgets:toll_segment_list")

        if action not in ("activate", "deactivate"):
            messages.error(request, "Acción no válida.")
            return redirect("budgets:toll_segment_list")

        target_state = action == "activate"
        updated = TollSegment.objects.filter(
            pk__in=valid_pks,
        ).update(is_active=target_state)

        verb = "activados" if target_state else "desactivados"
        messages.success(
            request,
            f"{updated} tramo{'s' if updated != 1 else ''} {verb} "
            f"correctamente.",
        )
        return redirect("budgets:toll_segment_list")


class WorkOrderPdfView(AssistanceRequiredMixin, View):
    """
    Generates and streams a PDF export of a WorkOrderAssistanceUnit albarán.
    Uses weasyprint to render albaran_pdf.html as a PDF response.
    Accessible to ASSISTANCE and ADMIN roles scoped to the company.
    ---
    Genera y envía como respuesta un PDF del albarán de una
    WorkOrderAssistanceUnit. Usa weasyprint para renderizar albaran_pdf.html
    como respuesta PDF. Accesible para ASSISTANCE y ADMIN con ámbito empresa.
    """

    template_name = "budgets/albaran_pdf.html"

    def get(self, request, pk):
        """
        Renders the unit albarán as a downloadable PDF.
        ---
        Renderiza el albarán de la unidad como PDF descargable.
        """
        from budgets.models import WorkOrderAssistanceUnit
        from django.http import HttpResponse
        company_user = _get_company_user(request)
        if not company_user:
            return redirect("panel:login")
        unit = get_object_or_404(
            WorkOrderAssistanceUnit,
            pk=pk,
            work_order__company=company_user.company,
        )
        html_string = render(
            request,
            self.template_name,
            {
                "unit":        unit,
                "work_order":  unit.work_order,
                "company":     company_user.company,
            },
        ).content.decode("utf-8")
        pdf_file = weasyprint.HTML(string=html_string).write_pdf()
        filename = (
            f"albaran_"
            f"{unit.work_order.work_order_number}-"
            f"{unit.unit_number:02d}.pdf"
        )
        response = HttpResponse(pdf_file, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


# ---------------------------------------------------------------------------
# Base calendar management — ADMIN only
# Gestión de calendarios laborales de bases — solo ADMIN
# ---------------------------------------------------------------------------

class BaseCalendarView(AdminRoleRequiredMixin, View):
    """
    List all active bases for the company with their labour calendar summary.
    Provides an entry point to manage each base's holiday dates manually.
    Also exposes the sync-all action to refresh calendars from the external
    source (calendariosnacionales.com).
    ---
    Lista todas las bases activas de la empresa con un resumen de su
    calendario laboral. Punto de entrada para gestionar manualmente las
    fechas festivas de cada base. También expone la acción de sincronización
    global para refrescar los calendarios desde la fuente externa
    (calendariosnacionales.com).
    """

    template_name = "budgets/base_calendar_list.html"

    def get(self, request):
        """
        Render the base calendar list for the operator's company.
        Annotates each base with the number of holiday dates in its
        labor_calendar JSON field.
        ---
        Renderiza el listado de calendarios de bases de la empresa del
        operario. Anota cada base con el número de fechas festivas en
        su campo JSON labor_calendar.
        """
        company_user = _get_company_user(request)
        bases = (
            Base.objects
            .filter(
                company=company_user.company,
                is_active=True,
            )
            .order_by("municipality", "name")
        )
        # Build summary: count of dates per base (labor_calendar is a JSON
        # string stored as TextField — must be deserialised before use).
        # Construye resumen: número de fechas por base (labor_calendar es
        # un string JSON almacenado como TextField — debe deserializarse).
        import json as _json
        base_data = []
        for base in bases:
            raw = (base.labor_calendar or "").strip()
            try:
                cal = _json.loads(raw) if raw else []
            except (ValueError, TypeError):
                cal = []
            base_data.append({
                "base": base,
                "date_count": len(cal),
                "last_synced": base.calendar_synced_at,
            })
        ctx = _build_base_context(request, {
            "base_data": base_data,
            "active_nav": "budgets_calendars",
        })
        return render(request, self.template_name, ctx)


class BaseCalendarDetailView(AdminRoleRequiredMixin, View):
    """
    Manage the labour calendar of a single Base: view, add and remove
    individual holiday dates. Also provides an HTMX endpoint for adding
    and deleting dates without full page reload.
    ---
    Gestiona el calendario laboral de una Base individual: ver, añadir
    y eliminar fechas festivas individuales. También provee un endpoint
    HTMX para añadir y eliminar fechas sin recarga completa de página.
    """

    template_name = "budgets/base_calendar_detail.html"
    FRAGMENT_TEMPLATE = (
        "budgets/partials/calendar_dates_fragment.html"
    )

    def _get_base(self, request, pk):
        """
        Return the Base scoped to the company of the authenticated user.
        ---
        Devuelve la Base acotada a la empresa del usuario autenticado.
        """
        company_user = _get_company_user(request)
        return get_object_or_404(
            Base,
            pk=pk,
            company=company_user.company,
        )

    def get(self, request, pk):
        """
        Render the detail page with the sorted list of holiday dates for
        the base, and an add-date form.
        ---
        Renderiza la página de detalle con la lista ordenada de fechas
        festivas de la base y un formulario de añadir fecha.
        """
        base = self._get_base(request, pk)
        import json as _json
        company_user = _get_company_user(request)
        raw = (base.labor_calendar or "").strip()
        try:
            dates = sorted(_json.loads(raw) if raw else [])
        except (ValueError, TypeError):
            dates = []
        other_bases = (
            Base.objects
            .filter(company=company_user.company, is_active=True)
            .exclude(pk=pk)
            .order_by("municipality", "name")
        )
        ctx = _build_base_context(request, {
            "base": base,
            "dates": dates,
            "other_bases": other_bases,
            "active_nav": "budgets_calendars",
        })
        return render(request, self.template_name, ctx)

    def post(self, request, pk):
        """
        Handle add-date (action=add) and remove-date (action=remove) POST
        requests. Returns an HTMX fragment with the updated date list when
        the request is HTMX, otherwise redirects to the detail page.
        ---
        Gestiona las peticiones POST de añadir fecha (action=add) y
        eliminar fecha (action=remove). Devuelve un fragmento HTMX con
        la lista de fechas actualizada cuando la petición es HTMX,
        si no redirige a la página de detalle.
        """
        import datetime as _dt
        import json as _json
        base = self._get_base(request, pk)
        action = request.POST.get("action", "add")
        date_str = request.POST.get("date", "").strip()
        is_htmx = request.headers.get("HX-Request") == "true"

        # Deserialise TextField JSON string to list.
        # Deserializar el string JSON del TextField a lista.
        raw = (base.labor_calendar or "").strip()
        try:
            cal = _json.loads(raw) if raw else []
        except (ValueError, TypeError):
            cal = []
        error = None

        if action == "add":
            # Validate ISO date format.
            # Validar formato de fecha ISO.
            try:
                _dt.date.fromisoformat(date_str)
            except ValueError:
                error = "Fecha inválida. Usa el formato YYYY-MM-DD."
            else:
                if date_str not in cal:
                    cal.append(date_str)
                    base.labor_calendar = _json.dumps(
                        sorted(cal), ensure_ascii=False
                    )
                    base.save(update_fields=["labor_calendar"])
                    messages.success(
                        request,
                        f"Fecha {date_str} añadida al calendario.",
                    )
                else:
                    error = (
                        f"La fecha {date_str} ya existe en el calendario."
                    )

        elif action == "remove":
            if date_str in cal:
                cal.remove(date_str)
                base.labor_calendar = _json.dumps(
                    sorted(cal), ensure_ascii=False
                )
                base.save(update_fields=["labor_calendar"])
                messages.success(
                    request,
                    f"Fecha {date_str} eliminada del calendario.",
                )
            else:
                error = (
                    f"La fecha {date_str} no existe en el calendario."
                )

        raw2 = (base.labor_calendar or "").strip()
        try:
            dates = sorted(_json.loads(raw2) if raw2 else [])
        except (ValueError, TypeError):
            dates = []

        if is_htmx:
            ctx = _build_base_context(request, {
                "base": base,
                "dates": dates,
                "error": error,
                "active_nav": "budgets_calendars",
            })
            return render(request, self.FRAGMENT_TEMPLATE, ctx)

class BaseCalendarCopyView(AdminRoleRequiredMixin, View):
    """
    Copy the labour calendar of one Base to another Base of the same
    company. Only the target base's labor_calendar is overwritten —
    the source base is not modified. Only ADMIN role is allowed.
    ---
    Copia el calendario laboral de una Base a otra Base de la misma
    empresa. Solo se sobreescribe el labor_calendar de la base destino
    — la base origen no se modifica. Solo rol ADMIN.
    """

    def post(self, request, pk):
        """
        POST bases/<pk>/copy-calendar/
        Body param: target_base_id (PK of the destination base).
        Copies source base labor_calendar JSON to target base and
        redirects to the target base calendar detail page.
        ---
        Param body: target_base_id (PK de la base destino).
        Copia el labor_calendar JSON de la base origen a la base
        destino y redirige al detalle del calendario de la base destino.
        """
        import json as _json
        company_user = _get_company_user(request)
        company = company_user.company

        source_base = get_object_or_404(Base, pk=pk, company=company)
        target_base_id = request.POST.get("target_base_id", "").strip()

        if not target_base_id:
            messages.error(request, "Selecciona una base destino.")
            return redirect("budgets:base_calendar_detail", pk=pk)

        try:
            target_base_id = int(target_base_id)
        except (ValueError, TypeError):
            messages.error(request, "Base destino no válida.")
            return redirect("budgets:base_calendar_detail", pk=pk)

        if target_base_id == source_base.pk:
            messages.error(
                request,
                "La base destino no puede ser la misma que la base origen.",
            )
            return redirect("budgets:base_calendar_detail", pk=pk)

        target_base = get_object_or_404(
            Base, pk=target_base_id, company=company
        )

        target_base.labor_calendar = source_base.labor_calendar or "[]"
        target_base.save(update_fields=["labor_calendar"])

        # Count dates for the success message.
        # Contar fechas para el mensaje de éxito.
        try:
            n_dates = len(
                _json.loads(target_base.labor_calendar or "[]")
            )
        except (ValueError, TypeError):
            n_dates = 0

        messages.success(
            request,
            f"Calendario de \"{source_base.name}\" copiado a "
            f"\"{target_base.name}\" ({n_dates} fechas).",
        )
        return redirect(
            "budgets:base_calendar_detail", pk=target_base_id
        )










# Bloque a AÑADIR al final de budgets/views.py

# ---------------------------------------------------------------------------
# Insurer tariff PDF import — upload + Gemini extraction + clone + review
# ---------------------------------------------------------------------------

class InsurerTariffPdfUploadView(AdminRoleRequiredMixin, View):
    """
    Handles the upload of a tariff PDF for an existing Insurer.
    On POST:
      1. Validates the uploaded file (PDF/image, max 20 MB).
      2. Clones the source Insurer as '<name> copia' (atomic, using
         the same logic as InsurerCloneView but with a fixed suffix).
      3. Runs TariffPdfExtractionService.extract() passing
         insurer.insurer_company_name so Gemini selects the right column.
      4. Saves the raw extraction as TariffPdfImport.extraction_raw (JSON).
      5. Redirects to InsurerTariffPdfReviewView for human validation.
    On error: sets TariffPdfImport.status = ERROR, shows error message,
    redirects back to the insurer edit form.
    ---
    Gestiona la subida de un PDF de tarifa para una Insurer existente.
    En POST:
      1. Valida el archivo subido (PDF/imagen, máx. 20 MB).
      2. Clona la Insurer origen como '<nombre> copia' (atómico, usando
         la misma lógica que InsurerCloneView pero con sufijo fijo).
      3. Ejecuta TariffPdfExtractionService.extract() pasando
         insurer.insurer_company_name para que Gemini seleccione la columna.
      4. Guarda la extracción bruta como TariffPdfImport.extraction_raw (JSON).
      5. Redirige a InsurerTariffPdfReviewView para validación humana.
    En error: establece TariffPdfImport.status = ERROR, muestra mensaje
    de error, redirige al formulario de edición de la aseguradora.
    """

    _MAX_PDF_BYTES = 20 * 1024 * 1024  # 20 MB
    _ALLOWED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".webp"}

    def post(self, request, pk):
        import datetime as _dt
        import json as _json
        from budgets.models import (
            TariffPdfImport, InsurerTariff, TariffLine, VehicleType,
            InsurerBase, SpecialRateTariff, SpecialRateLine,
        )
        from budgets.services import TariffPdfExtractionService

        company_user = _get_company_user(request)
        source = get_object_or_404(
            Insurer,
            pk=pk,
            company=company_user.company,
        )

        # --- Validate uploaded file ---
        # --- Validar archivo subido ---
        pdf_file = request.FILES.get("tariff_pdf")
        if not pdf_file:
            messages.error(request, "No se ha seleccionado ningún archivo.")
            return redirect("budgets:insurer_update", pk=pk)

        import pathlib as _pl
        ext = _pl.Path(pdf_file.name).suffix.lower()
        if ext not in self._ALLOWED_EXTENSIONS:
            messages.error(
                request,
                f"Tipo de archivo no válido ({ext}). "
                f"Se aceptan: PDF, JPG, PNG, WEBP.",
            )
            return redirect("budgets:insurer_update", pk=pk)

        if pdf_file.size > self._MAX_PDF_BYTES:
            messages.error(
                request,
                "El archivo supera el tamaño máximo permitido (20 MB).",
            )
            return redirect("budgets:insurer_update", pk=pk)

        # --- Build the copy name: '<name> copia' (with numeric suffix if
        # collision exists) ---
        # --- Construir el nombre de copia: '<nombre> copia' (con sufijo
        # numérico si ya existe) ---
        base_copy_name = f"{source.name} copia"
        copy_name = base_copy_name
        copy_suffix = 1
        while Insurer.objects.filter(
            company=company_user.company,
            name=copy_name,
        ).exists():
            copy_name = f"{base_copy_name} {copy_suffix}"
            copy_suffix += 1

        # Derive code for the copy using same logic as InsurerCloneView.
        # Derivar código de la copia usando la misma lógica que InsurerCloneView.
        import re as _re

        def _norm(s):
            s = s.upper()
            s = _re.sub(r"[\s\-\.\(\)\/]+", "_", s)
            s = _re.sub(r"_+", "_", s)
            return s.strip("_")

        copy_code_base = (
            _norm(copy_name)[:45]
        )
        copy_code = copy_code_base
        code_suffix = 1
        while Insurer.objects.filter(
            company=company_user.company,
            code=copy_code,
        ).exists():
            copy_code = f"{copy_code_base}_{code_suffix}"
            code_suffix += 1

        pdf_import = None

        try:
            with transaction.atomic():
                # --- Clone Insurer as copy ---
                # --- Clonar Insurer como copia ---
                copy_insurer = Insurer.objects.create(
                    company=source.company,
                    name=copy_name,
                    insurer_company_name=source.insurer_company_name,
                    service_company_name=source.service_company_name,
                    code=copy_code,
                    management_fee_percent=source.management_fee_percent,
                    surcharges_are_cumulative=source.surcharges_are_cumulative,
                    is_active=False,
                    is_insurance_company=source.is_insurance_company,
                    always_apply_iva=source.always_apply_iva,
                    special_night_holiday_tariff=source.special_night_holiday_tariff,
                    local_service_km_threshold=source.local_service_km_threshold,
                    notes=source.notes,
                )

                # Clone VehicleType catalogue, building pk map for TariffLine FKs.
                # Clonar catálogo VehicleType, construyendo mapa pk para FK de TariffLine.
                vt_map = {}
                for vt in source.vehicle_types.all().order_by("sort_order", "name"):
                    clone_vt = VehicleType.objects.create(
                        insurer=copy_insurer,
                        name=vt.name,
                        sort_order=vt.sort_order,
                        is_active=vt.is_active,
                    )
                    vt_map[vt.pk] = clone_vt

                # Clone active tariff + lines.
                # Clonar tarifa activa + líneas.
                source_tariff = InsurerTariff.objects.filter(
                    insurer=source,
                    valid_to__isnull=True,
                ).prefetch_related("lines").first()

                clone_tariff = None
                clone_srt = None

                if source_tariff:
                    clone_tariff = InsurerTariff.objects.create(
                        insurer=copy_insurer,
                        year=source_tariff.year,
                        valid_from=_dt.date.today(),
                        valid_to=None,
                        notes=source_tariff.notes,
                    )
                    for line in source_tariff.lines.all():
                        TariffLine.objects.create(
                            tariff=clone_tariff,
                            vehicle_type=(
                                vt_map.get(line.vehicle_type_id)
                                if line.vehicle_type_id else None
                            ),
                            concept=line.concept,
                            unit=line.unit,
                            price=line.price,
                            km_threshold=line.km_threshold,
                            min_units=line.min_units,
                            requires_authorization=line.requires_authorization,
                        )

                    try:
                        source_srt = source_tariff.special_rate
                        clone_srt = SpecialRateTariff.objects.create(
                            insurer_tariff=clone_tariff,
                            notes=source_srt.notes,
                        )
                        for srl in source_srt.lines.all():
                            SpecialRateLine.objects.create(
                                special_rate_tariff=clone_srt,
                                vehicle_type=(
                                    vt_map.get(srl.vehicle_type_id)
                                    if srl.vehicle_type_id else None
                                ),
                                concept=srl.concept,
                                unit=srl.unit,
                                price=srl.price,
                                km_threshold=srl.km_threshold,
                                min_units=srl.min_units,
                            )
                    except SpecialRateTariff.DoesNotExist:
                        pass

                # Clone InsurerBase links.
                # Clonar vínculos InsurerBase.
                for ib in InsurerBase.objects.filter(insurer=source):
                    InsurerBase.objects.create(
                        insurer=copy_insurer,
                        base=ib.base,
                        is_active=ib.is_active,
                    )

                # Create TariffPdfImport record (PDF saved here).
                # Crear registro TariffPdfImport (PDF guardado aquí).
                pdf_import = TariffPdfImport.objects.create(
                    insurer=source,
                    insurer_copy=copy_insurer,
                    pdf_file=pdf_file,
                    status=TariffPdfImport.STATUS_PENDING,
                )

            # --- Run Gemini extraction OUTSIDE the atomic block ---
            # --- Ejecutar extracción Gemini FUERA del bloque atómico ---
            svc = TariffPdfExtractionService()
            extraction = svc.extract(
                file_path=pdf_import.pdf_file.path,
                insurer=source,
            )

            # Persist raw extraction as JSON for audit.
            # Persistir extracción bruta como JSON para auditoría.
            pdf_import.extraction_raw = extraction.model_dump()
            pdf_import.save(update_fields=["extraction_raw", "updated_at"])

        except Exception as exc:
            _logger_upload = __import__("logging").getLogger(__name__)
            _logger_upload.exception(
                "# Error en TariffPdfUploadView para insurer pk=%s: %s",
                pk, exc,
            )
            if pdf_import is not None:
                pdf_import.status = TariffPdfImport.STATUS_ERROR
                pdf_import.error_message = str(exc)
                pdf_import.save(
                    update_fields=["status", "error_message", "updated_at"]
                )
            messages.error(
                request,
                "Error durante la extracción del PDF. Revisa el archivo "
                "o inténtalo de nuevo.",
            )
            return redirect("budgets:insurer_update", pk=pk)

        return redirect("budgets:insurer_tariff_pdf_review", pk=pdf_import.pk)


class InsurerTariffPdfReviewView(AdminRoleRequiredMixin, View):
    """
    GET: Displays two groups of extracted data for human review:
      1. Matched lines — extracted prices that Gemini paired with an
         EXISTING TariffConcept, shown side-by-side against the copy
         insurer's current price for that concept/vehicle_type/period.
      2. New concept proposals — extracted lines Gemini could not match
         to any existing concept, grouped by proposed label, with an
         editable name/unit and their price rows.
    All values are editable before applying.

    POST 'apply': updates matched TariffLine/SpecialRateLine prices,
    creates any new TariffConcept the operator confirms (company-scoped,
    auto-generated code) together with their TariffLine/SpecialRateLine
    rows, and — for insurers using a percent night/holiday surcharge
    instead of a full table — updates the NYF_PERCENT line if a value
    was extracted. Sets TariffPdfImport.status = APPLIED, keeps the copy.

    POST 'discard' (button labelled 'Cancelar'): permanently deletes the
    copy insurer (SpecialRateTariff is removed explicitly first, since
    SpecialRateTariff.insurer_tariff and SpecialRateLine.vehicle_type use
    on_delete=PROTECT and would otherwise block the cascade), marks the
    import DISCARDED, redirects to the source insurer. No trace of the
    copy remains.

    ---

    GET: Muestra dos grupos de datos extraídos para revisión humana:
      1. Líneas coincidentes — precios extraídos que Gemini emparejó con
         un TariffConcept YA EXISTENTE, mostrados junto al precio actual
         de la aseguradora copia para ese concepto/tipo de vehículo/
         periodo.
      2. Propuestas de concepto nuevo — líneas extraídas que Gemini no
         pudo emparejar con ningún concepto existente, agrupadas por
         nombre propuesto, con nombre/unidad editables y sus filas de
         precio.
    Todos los valores son editables antes de aplicar.

    POST 'apply': actualiza los precios de TariffLine/SpecialRateLine
    coincidentes, crea cualquier TariffConcept nuevo que el operario
    confirme (propio de la empresa, código auto-generado) junto con sus
    filas TariffLine/SpecialRateLine, y — para aseguradoras que usan
    recargo porcentual nocturno/festivo en vez de tabla completa —
    actualiza la línea NYF_PERCENT si se extrajo un valor. Marca
    TariffPdfImport.status = APPLIED, conserva la copia.

    POST 'discard' (botón 'Cancelar'): elimina permanentemente la
    aseguradora copia (se borra primero explícitamente SpecialRateTariff,
    ya que SpecialRateTariff.insurer_tariff y SpecialRateLine.vehicle_type
    usan on_delete=PROTECT y bloquearían la cascada si no), marca la
    importación como DISCARDED, redirige a la aseguradora origen. No
    queda ningún rastro de la copia.
    """

    @staticmethod
    def _norm(name):
        return (name or "").strip().upper()

    def _build_review_context(self, pdf_import):
        """
        Builds the context dict for the review template: matched_rows
        (existing concept, grouped laborable/festivo) and
        new_concept_groups (proposed concepts with their price rows).
        ---
        Construye el dict de contexto para el template de revisión:
        matched_rows (concepto existente, agrupado laborable/festivo) y
        new_concept_groups (conceptos propuestos con sus filas de precio).
        """
        from budgets.models import (
            InsurerTariff, TariffLine, TariffConcept,
            SpecialRateTariff, SpecialRateLine, VehicleType,
        )

        copy = pdf_import.insurer_copy
        extraction = pdf_import.extraction_raw or {}

        copy_tariff = InsurerTariff.objects.filter(
            insurer=copy,
            valid_to__isnull=True,
        ).prefetch_related(
            "lines__vehicle_type", "lines__concept"
        ).first()

        # Index existing lines by (vt_name_norm, concept_code) for
        # matched-concept comparison, and VehicleType by normalised name
        # for new-concept line creation.
        # Indexar líneas existentes por (vt_name_norm, concept_code) para
        # comparación de conceptos coincidentes, y VehicleType por nombre
        # normalizado para creación de líneas de concepto nuevo.
        lab_line_index = {}
        spe_line_index = {}
        vt_index = {}
        srt = None

        if copy_tariff:
            for vt in copy.vehicle_types.all():
                vt_index[self._norm(vt.name)] = vt

            for ln in copy_tariff.lines.all():
                vt_norm = self._norm(
                    ln.vehicle_type.name if ln.vehicle_type else "GENERAL"
                )
                lab_line_index[(vt_norm, ln.concept.code)] = ln

            try:
                srt = copy_tariff.special_rate
                for srl in srt.lines.all():
                    vt_norm = self._norm(
                        srl.vehicle_type.name if srl.vehicle_type else "GENERAL"
                    )
                    spe_line_index[(vt_norm, srl.concept.code)] = srl
            except SpecialRateTariff.DoesNotExist:
                srt = None

        matched_rows = {"laborable": [], "festivo": []}
        # new_concept_groups keyed by normalised proposed label to dedupe
        # multiple lines proposing the "same" new concept.
        # new_concept_groups indexado por etiqueta propuesta normalizada
        # para deduplicar varias líneas que proponen el "mismo" concepto.
        new_concept_groups = {}

        def _process(raw_lines, period, line_index):
            for item in raw_lines:
                vt_name = item.get("vehicle_type_name") or ""
                vt_norm = self._norm(vt_name) if vt_name else self._norm("GENERAL")
                price = item.get("price") or ""
                code_match = item.get("concept_code_match")

                if code_match:
                    current_line = line_index.get((vt_norm, code_match))
                    concept_label = code_match
                    try:
                        concept_label = TariffConcept.objects.get(
                            code=code_match
                        ).label
                    except TariffConcept.DoesNotExist:
                        pass
                    field_name = f"match_{period}_{vt_norm}_{code_match}"
                    matched_rows[period].append({
                        "vt_name": vt_name or "General",
                        "concept_code": code_match,
                        "concept_label": concept_label,
                        "current_price": (
                            str(current_line.price) if current_line else ""
                        ),
                        "extracted_price": price,
                        "line_pk": current_line.pk if current_line else "",
                        "field_name": field_name,
                    })
                else:
                    new_label = (item.get("concept_new_label") or "").strip()
                    if not new_label:
                        continue
                    key = self._norm(new_label)
                    group = new_concept_groups.setdefault(key, {
                        "label": new_label,
                        "unit": item.get("concept_new_unit") or "FIXED",
                        "rows": [],
                    })
                    row_idx = len(group["rows"])
                    field_name = f"newconcept_{key}_{period}_{vt_norm}_{row_idx}"
                    group["rows"].append({
                        "vt_name": vt_name or "General",
                        "vt_norm": vt_norm,
                        "period": period,
                        "extracted_price": price,
                        "field_name": field_name,
                    })

        _process(
            extraction.get("laborable_lines", []), "laborable", lab_line_index
        )
        _process(
            extraction.get("festivo_lines", []), "festivo", spe_line_index
        )

        # Night/holiday percent surcharge (percent-surcharge insurers only).
        # Recargo porcentual nocturno/festivo (solo aseguradoras sin tabla completa).
        nyf_percent_extracted = extraction.get("night_holiday_surcharge_percent")
        nyf_current_line = lab_line_index.get(
            (self._norm("GENERAL"), "NYF_PERCENT")
        )

        return {
            "pdf_import": pdf_import,
            "source_insurer": pdf_import.insurer,
            "copy_insurer": copy,
            "detected_insurer_name": extraction.get("detected_insurer_name", ""),
            "uses_full_night_holiday_table": bool(
                pdf_import.insurer.special_night_holiday_tariff
            ),
            "matched_laborable_rows": matched_rows["laborable"],
            "matched_festivo_rows": matched_rows["festivo"],
            "new_concept_groups": list(new_concept_groups.values()),
            "nyf_percent_extracted": nyf_percent_extracted or "",
            "nyf_percent_current": (
                str(nyf_current_line.price) if nyf_current_line else ""
            ),
            "nyf_percent_line_pk": (
                nyf_current_line.pk if nyf_current_line else ""
            ),
            "copy_tariff": copy_tariff,
        }

    def get(self, request, pk):
        from budgets.models import TariffPdfImport
        pdf_import = get_object_or_404(TariffPdfImport, pk=pk)
        company_user = _get_company_user(request)
        if pdf_import.insurer.company != company_user.company:
            return redirect("budgets:insurer_list")

        if pdf_import.status == TariffPdfImport.STATUS_ERROR:
            messages.error(
                request,
                f"La extracción falló: {pdf_import.error_message}",
            )
            return redirect("budgets:insurer_update", pk=pdf_import.insurer.pk)

        ctx = self._build_review_context(pdf_import)
        return render(request, "budgets/tariff_pdf_review.html", ctx)

    def post(self, request, pk):
        from decimal import Decimal, InvalidOperation
        import re as _re
        from budgets.models import (
            TariffPdfImport, TariffLine, SpecialRateLine, TariffConcept,
            InsurerTariff, SpecialRateTariff, VehicleType,
        )

        pdf_import = get_object_or_404(TariffPdfImport, pk=pk)
        company_user = _get_company_user(request)
        if pdf_import.insurer.company != company_user.company:
            return redirect("budgets:insurer_list")

        # 'Cancelar' action: deletes the copy insurer. SpecialRateTariff
        # must be removed explicitly first (PROTECT FK blocks the
        # automatic CASCADE otherwise). Marks the import DISCARDED.
        # ---
        # Acción 'Cancelar': elimina la aseguradora copia. Hay que borrar
        # SpecialRateTariff explícitamente primero (la FK PROTECT bloquea
        # el CASCADE automático si no). Marca la importación DISCARDED.
        if request.POST.get("action") == "discard":
            source_pk = pdf_import.insurer.pk
            copy = pdf_import.insurer_copy
            with transaction.atomic():
                pdf_import.status = TariffPdfImport.STATUS_DISCARDED
                pdf_import.insurer_copy = None
                pdf_import.save(
                    update_fields=["status", "insurer_copy", "updated_at"]
                )
                if copy is not None:
                    SpecialRateTariff.objects.filter(
                        insurer_tariff__insurer=copy
                    ).delete()
                    copy.delete()
            messages.info(
                request,
                "Importación cancelada. La aseguradora copia ha sido eliminada.",
            )
            return redirect("budgets:insurer_update", pk=source_pk)

        # 'Aceptar' action.
        # Acción 'Aceptar'.
        copy = pdf_import.insurer_copy
        if copy is None:
            messages.error(request, "La aseguradora copia ya no existe.")
            return redirect("budgets:insurer_list")

        def _safe_decimal(raw):
            if not raw:
                return None
            try:
                return Decimal(str(raw).replace(",", "."))
            except (InvalidOperation, ValueError):
                return None

        errors = []
        updated_count = 0
        created_concepts = 0

        with transaction.atomic():
            # --- 1. Matched lines: update existing TariffLine/SpecialRateLine ---
            # --- 1. Líneas coincidentes: actualizar TariffLine/SpecialRateLine existentes ---
            for key, raw_val in request.POST.items():
                if not raw_val:
                    continue
                new_price = _safe_decimal(raw_val)
                if new_price is None:
                    continue

                if key.startswith("match_laborable_"):
                    line_pk = request.POST.get(f"line_pk_{key}")
                    if not line_pk:
                        continue
                    try:
                        line = TariffLine.objects.get(pk=int(line_pk))
                        line.price = new_price
                        line.save(update_fields=["price"])
                        updated_count += 1
                    except (TariffLine.DoesNotExist, ValueError):
                        errors.append(f"TariffLine pk={line_pk} no encontrada.")

                elif key.startswith("match_festivo_"):
                    srl_pk = request.POST.get(f"srl_pk_{key}")
                    if not srl_pk:
                        continue
                    try:
                        srl = SpecialRateLine.objects.get(pk=int(srl_pk))
                        srl.price = new_price
                        srl.save(update_fields=["price"])
                        updated_count += 1
                    except (SpecialRateLine.DoesNotExist, ValueError):
                        errors.append(f"SpecialRateLine pk={srl_pk} no encontrada.")

            # --- 2. Night/holiday percent surcharge (if applicable) ---
            # --- 2. Recargo porcentual nocturno/festivo (si aplica) ---
            nyf_raw = request.POST.get("nyf_percent_value")
            nyf_line_pk = request.POST.get("nyf_percent_line_pk")
            nyf_price = _safe_decimal(nyf_raw)
            if nyf_price is not None:
                copy_tariff = InsurerTariff.objects.filter(
                    insurer=copy, valid_to__isnull=True
                ).first()
                if nyf_line_pk:
                    try:
                        line = TariffLine.objects.get(pk=int(nyf_line_pk))
                        line.price = nyf_price
                        line.save(update_fields=["price"])
                        updated_count += 1
                    except (TariffLine.DoesNotExist, ValueError):
                        errors.append("Línea NYF_PERCENT no encontrada.")
                elif copy_tariff:
                    try:
                        nyf_concept = TariffConcept.objects.get(
                            code="NYF_PERCENT"
                        )
                        TariffLine.objects.create(
                            tariff=copy_tariff,
                            vehicle_type=None,
                            concept=nyf_concept,
                            unit=TariffLine.UNIT_PERCENT,
                            price=nyf_price,
                        )
                        updated_count += 1
                    except TariffConcept.DoesNotExist:
                        errors.append(
                            "Concepto de sistema NYF_PERCENT no encontrado."
                        )

            # --- 3. New concept proposals: create TariffConcept + lines ---
            # --- 3. Propuestas de concepto nuevo: crear TariffConcept + líneas ---
            new_group_keys = set()
            for key in request.POST:
                if key.startswith("newconcept_label_"):
                    new_group_keys.add(key[len("newconcept_label_"):])

            copy_tariff = InsurerTariff.objects.filter(
                insurer=copy, valid_to__isnull=True
            ).first()
            srt = None
            if copy_tariff:
                try:
                    srt = copy_tariff.special_rate
                except SpecialRateTariff.DoesNotExist:
                    srt = None

            vt_index = {
                self._norm(vt.name): vt for vt in copy.vehicle_types.all()
            }

            for group_key in new_group_keys:
                create_flag = request.POST.get(f"newconcept_create_{group_key}")
                if not create_flag:
                    # Operator unchecked / did not confirm creation for
                    # this proposed concept — skip entirely.
                    # El operario no marcó / confirmó la creación de este
                    # concepto propuesto — se omite por completo.
                    continue

                label = (
                    request.POST.get(f"newconcept_label_{group_key}") or ""
                ).strip()
                unit = request.POST.get(f"newconcept_unit_{group_key}") or "FIXED"
                if not label:
                    continue

                # Auto-generate a unique code from the label.
                # Auto-generar un código único a partir del nombre.
                base_code = _re.sub(r"[^A-Z0-9]+", "_", label.upper()).strip("_")[:40]
                code = base_code
                suffix = 1
                while TariffConcept.objects.filter(code=code).exists():
                    code = f"{base_code}_{suffix}"
                    suffix += 1

                max_sort = TariffConcept.objects.filter(
                    company=copy.company
                ).count()
                concept = TariffConcept.objects.create(
                    code=code,
                    label=label,
                    default_unit=unit,
                    is_system=False,
                    company=copy.company,
                    sort_order=100 + max_sort,
                )
                created_concepts += 1

                # Create the line(s) for this new concept.
                # Crear la(s) línea(s) de este concepto nuevo.
                row_idx = 0
                while True:
                    price_key = f"newconcept_price_{group_key}_{row_idx}"
                    if price_key not in request.POST:
                        break
                    price_val = _safe_decimal(request.POST.get(price_key))
                    vt_norm_key = f"newconcept_vt_{group_key}_{row_idx}"
                    period_key = f"newconcept_period_{group_key}_{row_idx}"
                    vt_norm = request.POST.get(vt_norm_key, "")
                    period = request.POST.get(period_key, "laborable")
                    row_idx += 1

                    if price_val is None:
                        continue

                    vt_obj = (
                        vt_index.get(vt_norm)
                        if vt_norm and vt_norm != self._norm("GENERAL")
                        else None
                    )

                    if period == "laborable" and copy_tariff:
                        TariffLine.objects.create(
                            tariff=copy_tariff,
                            vehicle_type=vt_obj,
                            concept=concept,
                            unit=unit,
                            price=price_val,
                        )
                        updated_count += 1
                    elif period == "festivo" and srt:
                        SpecialRateLine.objects.create(
                            special_rate_tariff=srt,
                            vehicle_type=vt_obj,
                            concept=concept,
                            unit=unit,
                            price=price_val,
                        )
                        updated_count += 1

            pdf_import.status = TariffPdfImport.STATUS_APPLIED
            pdf_import.save(update_fields=["status", "updated_at"])

        if errors:
            messages.warning(
                request,
                f"Importación aplicada con {len(errors)} advertencia(s): "
                + "; ".join(errors),
            )
        else:
            extra = (
                f", {created_concepts} concepto(s) nuevo(s) creado(s)"
                if created_concepts else ""
            )
            messages.success(
                request,
                f"Tarifa importada correctamente ({updated_count} "
                f"precio(s) aplicado(s){extra}) en la aseguradora copia "
                f"'{copy.name}'.",
            )

        return redirect("budgets:insurer_update", pk=copy.pk)

